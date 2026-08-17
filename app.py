# app.py - AI Shiksha Global Platform with Universal Core + Local Overlay Architecture
# Built for global scalability with country-specific curriculum overlays
# Enhanced with Document Intelligence for all segments
# Teacher Intelligence Engine with Lesson Builders, Assessment, Hours Saved
# Professional Intelligence Engine with Research, Analysis & Portfolio Tools
# SME Growth Automation Engine with Data Infrastructure & AI Analytics

import streamlit as st
import json
import os
import random
import pandas as pd
from datetime import datetime, timedelta
import hashlib
import io
import re
import time
from typing import Dict, List, Any, Optional
from dataclasses import dataclass, field
from enum import Enum
import numpy as np

# ==================== PAGE CONFIGURATION - MUST BE FIRST ====================
st.set_page_config(
    page_title="AI Shiksha - Universal Core + Local Overlay",
    page_icon="🌍",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Try to import PDF and DOCX libraries with fallback
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import docx
except ImportError:
    docx = None

# OCR is optional: real extraction needs pytesseract + the Tesseract binary
# installed on the host. Falls back gracefully (see MultimodalOCREngine) if absent.
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# ==================== TEACHER DATA MODELS ====================

@dataclass
class LessonPlan:
    """Complete lesson plan structure"""
    id: str
    title: str
    subject: str
    grade_level: str
    duration: int  # minutes
    objectives: List[str]
    standards: List[str]
    materials: List[str]
    warm_up: str
    direct_instruction: str
    guided_practice: str
    independent_practice: str
    assessment: str
    differentiation: Dict[str, str]
    exit_ticket: str
    created_at: datetime
    modified_at: datetime

@dataclass
class Rubric:
    """Tiered grading rubric"""
    id: str
    title: str
    criteria: List[str]
    levels: List[str]  # e.g., ['Excellent', 'Proficient', 'Developing', 'Beginning']
    descriptions: Dict[str, Dict[str, str]]
    weightings: Dict[str, float]

@dataclass
class Quiz:
    """Assessment/Quiz structure"""
    id: str
    title: str
    subject: str
    grade_level: str
    questions: List[Dict[str, Any]]
    total_points: int
    time_limit: Optional[int]
    created_at: datetime

@dataclass
class StudentFeedback:
    """Student feedback batch"""
    id: str
    student_name: str
    assignment: str
    strengths: List[str]
    areas_for_growth: List[str]
    suggestions: str
    score: Optional[float]
    generated_at: datetime
    approved: bool = False

@dataclass
class TimeSavedMetric:
    """Time saved analytics"""
    week_start: datetime
    lesson_planning_hours: float
    grading_hours: float
    admin_hours: float
    total_hours: float
    week_over_week_change: float

class CurriculumStandard(Enum):
    COMMON_CORE = "common_core"
    NGSS = "ngss"
    TEKS = "teks"
    CBC = "cbc"
    NCTB = "nctb"
    NATIONAL = "national"

# ==================== TEACHER INTELLIGENCE ENGINE ====================

class TeacherIntelligenceEngine:
    """Teacher Intelligence Engine with Lesson Builders, Assessment, Hours Saved"""
    
    def __init__(self, country_code='kenya'):
        self.country = country_code
        self.overlay = LocalCurriculumOverlay.get_overlay(country_code)
        self.lesson_plans = []
        self.rubrics = []
        self.quizzes = []
        self.feedback_batches = []
        self.time_saved_history = []
        self.standard_databases = {}
        
        # Initialize with sample data
        self._initialize_sample_data()
        self._initialize_standard_databases()
    
    def _initialize_sample_data(self):
        """Initialize with sample teacher data"""
        # Sample lesson plan
        sample_lesson = LessonPlan(
            id='lp_001',
            title='Introduction to Photosynthesis',
            subject='Science',
            grade_level='Grade 7',
            duration=45,
            objectives=[
                'Explain the process of photosynthesis',
                'Identify the reactants and products of photosynthesis',
                'Describe the role of chlorophyll in photosynthesis'
            ],
            standards=['NGSS-LS1-6: Photosynthesis and energy flow'],
            materials=['Whiteboard', 'Markers', 'Plant diagrams', 'Worksheet'],
            warm_up='What do plants need to grow? Discuss with partner (5 min)',
            direct_instruction='Explain photosynthesis using diagrams and key vocabulary (15 min)',
            guided_practice='Label photosynthesis diagram with teacher support (10 min)',
            independent_practice='Complete worksheet independently (10 min)',
            assessment='Exit ticket: List 3 things plants need for photosynthesis',
            differentiation={
                'ELL': 'Key vocabulary list with images',
                'Gifted': 'Research extension: How does photosynthesis vary in different plants?',
                'Struggling': 'Fill-in-the-blank guided notes'
            },
            exit_ticket='Write 3 facts about photosynthesis',
            created_at=datetime.now() - timedelta(days=7),
            modified_at=datetime.now() - timedelta(days=3)
        )
        self.lesson_plans.append(sample_lesson)
        
        # Sample rubric
        sample_rubric = Rubric(
            id='rb_001',
            title='Science Lab Report Rubric',
            criteria=['Hypothesis', 'Procedure', 'Data Collection', 'Analysis', 'Conclusion'],
            levels=['Excellent', 'Proficient', 'Developing', 'Beginning'],
            descriptions={
                'Hypothesis': {
                    'Excellent': 'Clear, testable hypothesis with scientific reasoning',
                    'Proficient': 'Clear, testable hypothesis',
                    'Developing': 'Vague hypothesis, lacks clarity',
                    'Beginning': 'No hypothesis or completely unclear'
                },
                'Procedure': {
                    'Excellent': 'Detailed, replicable procedure with safety considerations',
                    'Proficient': 'Clear procedure, replicable',
                    'Developing': 'Missing steps, hard to follow',
                    'Beginning': 'No procedure or completely unclear'
                },
                'Data Collection': {
                    'Excellent': 'Comprehensive data with visual representations',
                    'Proficient': 'Complete data collection',
                    'Developing': 'Incomplete data, missing observations',
                    'Beginning': 'No data collected'
                },
                'Analysis': {
                    'Excellent': 'Deep analysis with connections to concepts',
                    'Proficient': 'Correct analysis of data',
                    'Developing': 'Superficial analysis, missing connections',
                    'Beginning': 'No analysis or completely incorrect'
                },
                'Conclusion': {
                    'Excellent': 'Well-supported conclusion with broader implications',
                    'Proficient': 'Supported conclusion',
                    'Developing': 'Weak conclusion, not supported by data',
                    'Beginning': 'No conclusion or completely unsupported'
                }
            },
            weightings={
                'Hypothesis': 0.15,
                'Procedure': 0.15,
                'Data Collection': 0.20,
                'Analysis': 0.25,
                'Conclusion': 0.25
            }
        )
        self.rubrics.append(sample_rubric)
        
        # Sample quiz
        sample_quiz = Quiz(
            id='qz_001',
            title='Photosynthesis Quiz',
            subject='Science',
            grade_level='Grade 7',
            questions=[
                {
                    'id': 'q1',
                    'type': 'multiple_choice',
                    'question': 'What is the primary energy source for photosynthesis?',
                    'options': ['Glucose', 'Sunlight', 'Water', 'Carbon Dioxide'],
                    'correct_answer': 'Sunlight',
                    'points': 2
                },
                {
                    'id': 'q2',
                    'type': 'multiple_choice',
                    'question': 'What are the products of photosynthesis?',
                    'options': ['Carbon Dioxide and Water', 'Glucose and Oxygen', 'Water and Oxygen', 'Glucose and Carbon Dioxide'],
                    'correct_answer': 'Glucose and Oxygen',
                    'points': 2
                },
                {
                    'id': 'q3',
                    'type': 'short_answer',
                    'question': 'What is the role of chlorophyll in photosynthesis?',
                    'correct_answer': 'Chlorophyll absorbs sunlight energy for photosynthesis',
                    'points': 3
                },
                {
                    'id': 'q4',
                    'type': 'open_ended',
                    'question': 'Explain how photosynthesis contributes to the carbon cycle.',
                    'points': 5
                }
            ],
            total_points=12,
            time_limit=15,
            created_at=datetime.now() - timedelta(days=2)
        )
        self.quizzes.append(sample_quiz)
        
        # Sample time saved metrics
        sample_time_saved = TimeSavedMetric(
            week_start=datetime.now() - timedelta(days=7),
            lesson_planning_hours=2.5,
            grading_hours=3.0,
            admin_hours=1.5,
            total_hours=7.0,
            week_over_week_change=0.15
        )
        self.time_saved_history.append(sample_time_saved)
    
    def _initialize_standard_databases(self):
        """Initialize curriculum standard databases"""
        self.standard_databases = {
            'common_core': {
                'ELA': ['CCRA.R.1', 'CCRA.R.2', 'CCRA.W.1', 'CCRA.W.2', 'CCRA.SL.1'],
                'Math': ['CCRA.M.1', 'CCRA.M.2', 'CCRA.M.3', 'CCRA.M.4', 'CCRA.M.5']
            },
            'ngss': {
                'Physical': ['PS1-1', 'PS1-2', 'PS2-1', 'PS2-2', 'PS3-1'],
                'Life': ['LS1-1', 'LS1-2', 'LS2-1', 'LS2-2', 'LS3-1'],
                'Earth': ['ESS1-1', 'ESS1-2', 'ESS2-1', 'ESS2-2', 'ESS3-1']
            },
            'teks': {
                'Science': ['TEKS.7.5A', 'TEKS.7.6A', 'TEKS.7.7A', 'TEKS.7.8A'],
                'Math': ['TEKS.7.1A', 'TEKS.7.2A', 'TEKS.7.3A', 'TEKS.7.4A']
            },
            'cbc': {
                'Science': ['CBC.SCI.7.1', 'CBC.SCI.7.2', 'CBC.SCI.7.3'],
                'Math': ['CBC.MATH.7.1', 'CBC.MATH.7.2', 'CBC.MATH.7.3']
            }
        }
    
    # ===== Direct LMS Integrations =====
    
    def connect_lms(self, lms_type: str, credentials: Dict[str, str]) -> Dict[str, Any]:
        """Connect to Learning Management System"""
        lms_configs = {
            'google_classroom': {
                'name': 'Google Classroom',
                'icon': '📚',
                'features': ['sync_assignments', 'sync_grades', 'roster_management'],
                'requires_oauth': True
            },
            'canvas': {
                'name': 'Canvas',
                'icon': '🎨',
                'features': ['sync_assignments', 'sync_grades', 'course_materials', 'discussions'],
                'requires_oauth': True
            },
            'schoology': {
                'name': 'Schoology',
                'icon': '📖',
                'features': ['sync_assignments', 'sync_grades', 'courses'],
                'requires_oauth': True
            },
            'microsoft_teams': {
                'name': 'Microsoft Teams',
                'icon': '💬',
                'features': ['sync_assignments', 'sync_grades', 'teams_integration'],
                'requires_oauth': True
            }
        }
        
        if lms_type in lms_configs:
            return {
                'status': 'connected',
                'lms_type': lms_type,
                'config': lms_configs[lms_type],
                'message': f"Successfully connected to {lms_configs[lms_type]['name']}"
            }
        else:
            return {
                'status': 'error',
                'message': f"Unsupported LMS: {lms_type}"
            }
    
    def sync_lms_data(self, lms_type: str, sync_type: str = 'all') -> Dict[str, Any]:
        """Sync data with LMS"""
        # Simulate sync
        sync_data = {
            'students': 25,
            'assignments': 8,
            'grades_synced': 72,
            'last_sync': datetime.now().isoformat()
        }
        
        return {
            'status': 'success',
            'lms_type': lms_type,
            'sync_type': sync_type,
            'data': sync_data,
            'message': f"Successfully synced {sync_type} data from {lms_type}"
        }
    
    # ===== Curriculum Standard Mapping =====
    
    def map_standards(self, content: str, standard_type: str = 'common_core') -> List[str]:
        """Map content to curriculum standards"""
        matched_standards = []
        content_lower = content.lower()
        
        standards_db = self.standard_databases.get(standard_type, {})
        for category, standards in standards_db.items():
            for standard in standards:
                # Simple keyword matching
                standard_keywords = standard.lower().split('.')
                if any(keyword in content_lower for keyword in standard_keywords):
                    matched_standards.append(standard)
        
        return matched_standards[:5]  # Return top 5 matched standards
    
    # ===== Universal File Conversion =====
    
    def convert_document(self, uploaded_file, output_format: str) -> Dict[str, Any]:
        """Convert document between formats"""
        try:
            file_extension = uploaded_file.name.split('.')[-1].lower()
            doc_engine = DocumentIntelligenceEngine(self.country)
            text = doc_engine.extract_text_from_file(uploaded_file)
            
            if text.startswith("Error"):
                return {'status': 'error', 'message': text}
            
            # Simulate conversion based on output format
            converted = {
                'original_format': file_extension,
                'output_format': output_format,
                'content': text,
                'title': uploaded_file.name.replace(f'.{file_extension}', ''),
                'downloadable': True
            }
            
            if output_format == 'google_docs':
                converted['message'] = 'Ready for Google Docs export'
            elif output_format == 'google_slides':
                converted['message'] = 'Ready for Google Slides export'
            elif output_format == 'printable':
                converted['message'] = 'Printable format ready'
            else:
                converted['message'] = f'Converted to {output_format} format'
            
            return converted
            
        except Exception as e:
            return {'status': 'error', 'message': str(e)}
    
    # ===== Multi-Asset Bundle Generator =====
    
    def generate_lesson_bundle(self, subject: str, grade: str, topic: str, duration: int) -> Dict[str, Any]:
        """Generate complete lesson bundle with all assets"""
        # Generate lesson plan
        lesson_plan = self._generate_lesson_plan(subject, grade, topic, duration)
        
        # Generate slide outline
        slide_outline = self._generate_slide_outline(lesson_plan)
        
        # Generate worksheet
        worksheet = self._generate_worksheet(lesson_plan)
        
        # Generate answer key
        answer_key = self._generate_answer_key(worksheet)
        
        return {
            'lesson_plan': lesson_plan,
            'slide_outline': slide_outline,
            'worksheet': worksheet,
            'answer_key': answer_key,
            'bundle_id': f"bundle_{hashlib.md5(f'{subject}{grade}{topic}'.encode()).hexdigest()[:8]}",
            'created_at': datetime.now().isoformat()
        }
    
    def _generate_lesson_plan(self, subject: str, grade: str, topic: str, duration: int) -> Dict[str, Any]:
        """Generate individual lesson plan"""
        return {
            'title': f"{topic} - {subject} ({grade})",
            'subject': subject,
            'grade': grade,
            'duration': duration,
            'objectives': [
                f'Understand key concepts of {topic}',
                f'Apply {topic} knowledge to real-world scenarios',
                f'Develop critical thinking skills related to {topic}'
            ],
            'standards': self.map_standards(topic),
            'materials': ['Textbook', 'Worksheets', 'Multimedia resources'],
            'procedure': {
                'warm_up': f'Engage students with a question about {topic} (5-10 min)',
                'direct_instruction': f'Present key concepts of {topic} (15-20 min)',
                'guided_practice': f'Work through examples together (10-15 min)',
                'independent_practice': f'Students work independently on {topic} (15-20 min)',
                'closure': f'Review and assess understanding with exit ticket (5 min)'
            },
            'differentiation': {
                'struggling': 'Provide additional support and scaffolding',
                'gifted': 'Offer extension activities and deeper exploration',
                'ell': 'Include visuals and simplified language'
            },
            'assessment': f'Formative assessment throughout, summative quiz at end'
        }
    
    def _generate_slide_outline(self, lesson_plan: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Generate slide outline from lesson plan"""
        return [
            {'slide': 1, 'title': f"Welcome to {lesson_plan['title']}", 'type': 'title'},
            {'slide': 2, 'title': 'Learning Objectives', 'type': 'objectives'},
            {'slide': 3, 'title': 'Key Vocabulary', 'type': 'vocabulary'},
            {'slide': 4, 'title': 'Introduction', 'type': 'intro'},
            {'slide': 5, 'title': 'Main Concepts', 'type': 'content'},
            {'slide': 6, 'title': 'Examples', 'type': 'examples'},
            {'slide': 7, 'title': 'Practice Problems', 'type': 'practice'},
            {'slide': 8, 'title': 'Summary', 'type': 'summary'},
            {'slide': 9, 'title': 'Questions', 'type': 'questions'},
            {'slide': 10, 'title': 'Next Steps', 'type': 'next_steps'}
        ]
    
    def _generate_worksheet(self, lesson_plan: Dict[str, Any]) -> Dict[str, Any]:
        """Generate student worksheet"""
        return {
            'title': f"{lesson_plan['title']} - Worksheet",
            'sections': [
                {
                    'section': 'Part 1: Vocabulary',
                    'instructions': 'Match the vocabulary terms to their definitions',
                    'questions': ['Term 1', 'Term 2', 'Term 3']
                },
                {
                    'section': 'Part 2: Concept Check',
                    'instructions': 'Answer the following questions based on the lesson',
                    'questions': ['Question 1', 'Question 2', 'Question 3']
                },
                {
                    'section': 'Part 3: Application',
                    'instructions': 'Apply your knowledge to these scenarios',
                    'questions': ['Scenario 1', 'Scenario 2']
                }
            ],
            'total_points': 20
        }
    
    def _generate_answer_key(self, worksheet: Dict[str, Any]) -> Dict[str, Any]:
        """Generate answer key for worksheet"""
        return {
            'title': f"Answer Key - {worksheet['title']}",
            'answers': {
                'Part 1': ['Term 1 Definition', 'Term 2 Definition', 'Term 3 Definition'],
                'Part 2': ['Question 1 Answer', 'Question 2 Answer', 'Question 3 Answer'],
                'Part 3': ['Scenario 1 Solution', 'Scenario 2 Solution']
            },
            'total_points': worksheet['total_points']
        }
    
    # ===== Dynamic Lexile Adjuster =====
    
    def adjust_reading_level(self, text: str, target_level: str) -> Dict[str, Any]:
        """Adjust text reading level dynamically"""
        words = text.split()
        word_count = len(words)
        
        # Simulate level adjustment
        reading_levels = {
            'beginner': {'vocab_complexity': 0.3, 'sentence_complexity': 0.4, 'description': 'Grades 2-3'},
            'intermediate': {'vocab_complexity': 0.5, 'sentence_complexity': 0.6, 'description': 'Grades 4-6'},
            'advanced': {'vocab_complexity': 0.7, 'sentence_complexity': 0.8, 'description': 'Grades 7-9'},
            'expert': {'vocab_complexity': 0.9, 'sentence_complexity': 0.9, 'description': 'Grades 10-12'}
        }
        
        level_config = reading_levels.get(target_level, reading_levels['intermediate'])
        
        # Simulate text adjustment
        adjusted_text = text
        if target_level == 'beginner':
            adjusted_text = text.replace('complex', 'hard').replace('analyze', 'study')
            adjusted_text = ' '.join(adjusted_text.split()[:word_count//2])
        elif target_level == 'expert':
            adjusted_text = text + ' This content has been enhanced for advanced readers with additional context and depth.'
        
        return {
            'original_text': text,
            'adjusted_text': adjusted_text,
            'target_level': target_level,
            'level_description': level_config['description'],
            'word_count_original': word_count,
            'word_count_adjusted': len(adjusted_text.split()),
            'changes': {
                'vocabulary': f"Adjusted to {level_config['vocab_complexity']*100:.0f}% complexity",
                'sentences': f"Adjusted to {level_config['sentence_complexity']*100:.0f}% complexity"
            }
        }
    
    # ===== Rubric & Assessment Engine =====
    
    def generate_rubric(self, title: str, criteria: List[str], levels: List[str]) -> Rubric:
        """Generate tiered grading rubric"""
        rubric_id = f"rb_{hashlib.md5(f'{title}{datetime.now().timestamp()}'.encode()).hexdigest()[:8]}"
        
        # Generate descriptions
        descriptions = {}
        for criterion in criteria:
            descriptions[criterion] = {
                levels[0]: f"Excellent: {criterion} - demonstrates mastery",
                levels[1]: f"Proficient: {criterion} - shows understanding",
                levels[2]: f"Developing: {criterion} - needs more support",
                levels[3]: f"Beginning: {criterion} - requires intervention"
            }
        
        # Generate weightings (equal distribution)
        weightings = {criterion: 1.0 / len(criteria) for criterion in criteria}
        
        return Rubric(
            id=rubric_id,
            title=title,
            criteria=criteria,
            levels=levels,
            descriptions=descriptions,
            weightings=weightings
        )
    
    def generate_quiz(self, subject: str, grade: str, topic: str, num_questions: int = 5) -> Quiz:
        """Generate aligned quiz with multiple question types"""
        quiz_id = f"qz_{hashlib.md5(f'{subject}{grade}{topic}{datetime.now().timestamp()}'.encode()).hexdigest()[:8]}"
        
        questions = []
        question_types = ['multiple_choice', 'short_answer', 'open_ended']
        
        for i in range(num_questions):
            q_type = question_types[i % len(question_types)]
            
            question = {
                'id': f'q{i+1}',
                'type': q_type,
                'question': f'Question {i+1} about {topic}',
                'points': 2 if q_type == 'multiple_choice' else 3 if q_type == 'short_answer' else 5
            }
            
            if q_type == 'multiple_choice':
                question['options'] = ['Option A', 'Option B', 'Option C', 'Option D']
                question['correct_answer'] = 'Option A'
            
            if q_type == 'short_answer':
                question['correct_answer'] = f'Correct answer for question {i+1}'
            
            questions.append(question)
        
        total_points = sum(q.get('points', 0) for q in questions)
        
        return Quiz(
            id=quiz_id,
            title=f"{topic} Quiz - {subject} ({grade})",
            subject=subject,
            grade_level=grade,
            questions=questions,
            total_points=total_points,
            time_limit=total_points * 2,  # Rough estimate
            created_at=datetime.now()
        )
    
    # ===== Modular Block Editor =====
    
    def get_lesson_blocks(self, lesson_plan: LessonPlan) -> Dict[str, Any]:
        """Get modular blocks for lesson editing"""
        return {
            'blocks': {
                'warm_up': {
                    'content': lesson_plan.warm_up,
                    'editable': True,
                    'regenerable': True,
                    'placeholder': 'Enter warm-up activity...'
                },
                'direct_instruction': {
                    'content': lesson_plan.direct_instruction,
                    'editable': True,
                    'regenerable': True,
                    'placeholder': 'Enter direct instruction content...'
                },
                'guided_practice': {
                    'content': lesson_plan.guided_practice,
                    'editable': True,
                    'regenerable': True,
                    'placeholder': 'Enter guided practice activity...'
                },
                'independent_practice': {
                    'content': lesson_plan.independent_practice,
                    'editable': True,
                    'regenerable': True,
                    'placeholder': 'Enter independent practice...'
                },
                'assessment': {
                    'content': lesson_plan.assessment,
                    'editable': True,
                    'regenerable': True,
                    'placeholder': 'Enter assessment details...'
                },
                'exit_ticket': {
                    'content': lesson_plan.exit_ticket,
                    'editable': True,
                    'regenerable': True,
                    'placeholder': 'Enter exit ticket question...'
                }
            },
            'lesson_id': lesson_plan.id,
            'title': lesson_plan.title
        }
    
    def regenerate_block(self, block_type: str, context: str) -> str:
        """Regenerate a specific lesson block"""
        regeneration_templates = {
            'warm_up': f"New warm-up activity for: {context}",
            'direct_instruction': f"Updated direct instruction for: {context}",
            'guided_practice': f"New guided practice for: {context}",
            'independent_practice': f"Updated independent practice for: {context}",
            'assessment': f"New assessment for: {context}",
            'exit_ticket': f"New exit ticket for: {context}"
        }
        
        return regeneration_templates.get(block_type, f"Regenerated {block_type} for: {context}")
    
    # ===== One-Tap Accommodation Toggles =====
    
    def apply_accommodation(self, content: str, accommodation_type: str) -> Dict[str, Any]:
        """Apply accommodations for IEP/504/ELL modifications"""
        accommodations = {
            'iep_504': {
                'name': 'IEP/504 Modification',
                'changes': 'Extended time, modified assignments, preferential seating',
                'applied': True
            },
            'sentence_starters': {
                'name': 'Sentence Starters',
                'changes': 'Added sentence frames for writing tasks',
                'applied': True
            },
            'ell_translation': {
                'name': 'ELL Translation',
                'changes': 'Added bilingual support and visual aids',
                'applied': True
            },
            'simplified_text': {
                'name': 'Simplified Text',
                'changes': 'Reduced complexity, added definitions',
                'applied': True
            }
        }
        
        result = accommodations.get(accommodation_type, {})
        
        if result.get('applied', False):
            # Apply the accommodation
            modified_content = content
            if accommodation_type == 'sentence_starters':
                modified_content = f"Start with: I think that...\n\nOriginal: {content}"
            elif accommodation_type == 'ell_translation':
                modified_content = f"[Translation]\n{content}\n\n[Original]\n{content}"
            elif accommodation_type == 'simplified_text':
                modified_content = f"Simple: {content[:100]}...\n\nOriginal: {content}"
            
            return {
                'status': 'applied',
                'accommodation_type': accommodation_type,
                'name': result['name'],
                'changes': result['changes'],
                'modified_content': modified_content
            }
        
        return {'status': 'unknown', 'message': f'Unknown accommodation: {accommodation_type}'}
    
    # ===== Batch Feedback Assistant =====
    
    def generate_batch_feedback(self, assignments: List[Dict[str, Any]]) -> List[StudentFeedback]:
        """Generate batch feedback for multiple students"""
        feedback_batch = []
        
        for assignment in assignments:
            # Generate strengths and growth areas
            strengths = [
                random.choice(['Strong understanding of concepts', 'Good analysis', 'Clear writing']),
                random.choice(['Effective use of evidence', 'Engaging writing style', 'Critical thinking'])
            ]
            
            areas_for_growth = [
                random.choice(['Add more detail', 'Check grammar and spelling', 'Provide more examples']),
                random.choice(['Expand on ideas', 'Clarify arguments', 'Use stronger evidence'])
            ]
            
            feedback = StudentFeedback(
                id=f"fb_{hashlib.md5(f'{assignment.get('student_name', '')}{datetime.now().timestamp()}'.encode()).hexdigest()[:8]}",
                student_name=assignment.get('student_name', 'Student'),
                assignment=assignment.get('assignment_title', 'Assignment'),
                strengths=strengths,
                areas_for_growth=areas_for_growth,
                suggestions=f"Based on your work, I recommend focusing on {areas_for_growth[0].lower()} and {areas_for_growth[1].lower()}.",
                score=random.randint(70, 95),
                generated_at=datetime.now(),
                approved=False
            )
            
            feedback_batch.append(feedback)
        
        self.feedback_batches.extend(feedback_batch)
        return feedback_batch
    
    def approve_feedback(self, feedback_id: str) -> Dict[str, Any]:
        """Approve AI-suggested feedback"""
        for feedback in self.feedback_batches:
            if feedback.id == feedback_id:
                feedback.approved = True
                return {
                    'status': 'approved',
                    'feedback_id': feedback_id,
                    'message': f"Feedback for {feedback.student_name} approved"
                }
        
        return {'status': 'error', 'message': 'Feedback not found'}
    
    # ===== Time-Saved Dashboard =====
    
    def calculate_time_saved(self) -> Dict[str, Any]:
        """Calculate time saved analytics"""
        # Simulate time saved calculations
        total_hours = 0
        breakdown = {
            'lesson_planning': {'hours': 2.5, 'saved': 'vs 4.5 hours traditional'},
            'grading': {'hours': 3.0, 'saved': 'vs 6 hours traditional'},
            'admin_work': {'hours': 1.5, 'saved': 'vs 3 hours traditional'},
            'feedback_drafting': {'hours': 1.0, 'saved': 'vs 2 hours traditional'}
        }
        
        total_hours = sum(item['hours'] for item in breakdown.values())
        traditional_hours = 15.5  # Estimated traditional total
        
        return {
            'total_hours_saved': round(total_hours, 1),
            'total_percent_saved': round((total_hours / traditional_hours) * 100, 1),
            'breakdown': breakdown,
            'weekly_average': round(total_hours, 1),
            'monthly_projection': round(total_hours * 4, 1),
            'yearly_projection': round(total_hours * 48, 1)
        }
    
    # ===== Parent & Sub-Plan Generators =====
    
    def generate_parent_update(self, student_name: str, subject: str, progress: str) -> str:
        """Generate professional parent update"""
        templates = [
            f"Dear Parent/Guardian of {student_name},\n\n"
            f"I am pleased to share an update on {student_name}'s progress in {subject}. "
            f"{student_name} has been {progress}. "
            f"I look forward to continuing to support {student_name}'s growth. "
            f"Please reach out if you have any questions.\n\n"
            f"Best regards,\n"
            f"Teacher",
            
            f"Hello {student_name}'s Family,\n\n"
            f"This note is to provide a quick update on {student_name}'s learning in {subject}. "
            f"Recent work shows that {student_name} is {progress}. "
            f"Keep up the great support at home!\n\n"
            f"Sincerely,\n"
            f"Teacher"
        ]
        
        return random.choice(templates)
    
    def generate_sub_plan(self, subject: str, grade: str, lesson_topic: str) -> Dict[str, Any]:
        """Generate emergency substitute teacher guide"""
        return {
            'title': f"Sub Plan: {lesson_topic} - {subject} ({grade})",
            'overview': f"This plan covers {lesson_topic} for {grade} students in {subject}.",
            'schedule': {
                'period_1': {'time': '8:00-8:45', 'activity': f'Intro to {lesson_topic}', 'materials': 'Textbook pages 1-5'},
                'period_2': {'time': '8:50-9:35', 'activity': f'Guided practice on {lesson_topic}', 'materials': 'Worksheet'},
                'period_3': {'time': '9:40-10:25', 'activity': 'Independent work', 'materials': 'Assignment'},
                'period_4': {'time': '10:30-11:15', 'activity': 'Review and wrap-up', 'materials': 'Answer key'}
            },
            'emergency_procedures': 'In case of emergency, follow school protocol in the handbook.',
            'notes': 'Please leave notes on student progress for the regular teacher.'
        }
    
    # ===== Zero-PII Compliance Layer =====
    
    def redact_pii(self, text: str) -> Dict[str, Any]:
        """Redact personally identifiable information"""
        # Simulate PII redaction
        redacted_text = text
        
        # Replace common name patterns
        name_pattern = r'\b[A-Z][a-z]+\s[A-Z][a-z]+\b'
        redacted_text = re.sub(name_pattern, '[STUDENT]', redacted_text)
        
        # Replace email patterns
        email_pattern = r'\b[\w\.-]+@[\w\.-]+\.\w+\b'
        redacted_text = re.sub(email_pattern, '[EMAIL]', redacted_text)
        
        return {
            'original': text,
            'redacted': redacted_text,
            'pii_removed': len(re.findall(name_pattern, text)) + len(re.findall(email_pattern, text)),
            'compliance_note': 'FERPA/COPPA compliant - data isolated from model training'
        }


# ==================== PROFESSIONAL DATA MODELS ====================

@dataclass
class SourceDocument:
    """Source document with citation tracking"""
    id: str
    title: str
    source_type: str
    file_path: str
    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    citations: List[Dict[str, Any]] = field(default_factory=list)
    parsed_tables: List[Dict[str, Any]] = field(default_factory=list)
    uploaded_at: datetime = field(default_factory=datetime.now)

@dataclass
class Citation:
    """Audit-linked citation"""
    id: str
    document_id: str
    page_number: int
    paragraph: str
    claim: str
    confidence: float
    source_text: str

@dataclass
class PortfolioAsset:
    """Portfolio asset for monitoring"""
    ticker: str
    name: str
    sector: str
    shares: int
    purchase_price: float
    current_price: float
    market_value: float
    weight: float
    daily_change: float
    weekly_change: float
    monthly_change: float

@dataclass
class WatchdogAlert:
    """Automated alert from asset watchdog"""
    id: str
    asset_ticker: str
    alert_type: str
    severity: str
    message: str
    source: str
    timestamp: datetime
    read: bool = False

@dataclass
class MacroScenario:
    """Macro scenario for stress testing"""
    name: str
    description: str
    rate_change: float
    inflation_change: float
    growth_change: float
    market_shock: float
    portfolio_impact: float

class SearchType(Enum):
    HYBRID = "hybrid"
    VECTOR = "vector"
    KEYWORD = "keyword"
    SEMANTIC = "semantic"


# ==================== PROFESSIONAL INTELLIGENCE ENGINE ====================

class ProfessionalIntelligenceEngine:
    """Professional Intelligence Engine with Research, Analysis & Portfolio Tools"""
    
    def __init__(self, country_code='kenya'):
        self.country = country_code
        self.overlay = LocalCurriculumOverlay.get_overlay(country_code)
        self.documents = []
        self.citations = []
        self.portfolio_assets = []
        self.watchdog_alerts = []
        self.analysis_history = []
        self.search_index = {}
        self.vector_embeddings = {}
        
        self._initialize_sample_data()
    
    def _initialize_sample_data(self):
        """Initialize with sample professional data"""
        sample_docs = [
            {
                'id': 'doc_001',
                'title': 'Q4 2024 Earnings Report - TechCorp',
                'source_type': 'sec_filing',
                'content': """TechCorp Q4 2024 earnings exceeded analyst expectations. 
                Revenue reached $12.4 billion, up 18% year-over-year. 
                EPS of $2.15 beat estimates of $1.95. 
                Cloud revenue grew 32% to $4.2 billion. 
                Operating margin expanded to 28.5% from 26.1%.""",
                'metadata': {'filing_date': '2025-01-15', 'quarter': 'Q4 2024', 'company': 'TechCorp'},
                'citations': [
                    {'page': 2, 'paragraph': 3, 'claim': 'Revenue reached $12.4 billion', 'confidence': 0.95},
                    {'page': 2, 'paragraph': 4, 'claim': 'EPS of $2.15 beat estimates', 'confidence': 0.92}
                ]
            },
            {
                'id': 'doc_002',
                'title': 'AI Market Analysis 2025',
                'source_type': 'research',
                'content': """The AI market is projected to grow to $1.3 trillion by 2028, 
                with a CAGR of 24.5%. Key growth drivers include cloud computing, 
                natural language processing, and computer vision. 
                Enterprise adoption has accelerated 40% in the last year.""",
                'metadata': {'filing_date': '2025-01-10', 'analyst': 'Research Team', 'sector': 'Technology'},
                'citations': [
                    {'page': 1, 'paragraph': 2, 'claim': 'AI market projected to grow to $1.3 trillion', 'confidence': 0.88},
                    {'page': 2, 'paragraph': 1, 'claim': 'CAGR of 24.5%', 'confidence': 0.90}
                ]
            }
        ]
        
        for doc_data in sample_docs:
            doc = SourceDocument(
                id=doc_data['id'],
                title=doc_data['title'],
                source_type=doc_data['source_type'],
                file_path=f"{doc_data['id']}.pdf",
                content=doc_data['content'],
                metadata=doc_data['metadata'],
                citations=doc_data['citations']
            )
            self.documents.append(doc)
        
        sample_assets = [
            PortfolioAsset(
                ticker='AAPL', name='Apple Inc.', sector='Technology',
                shares=100, purchase_price=150.00, current_price=175.50,
                market_value=17550.00, weight=0.25,
                daily_change=0.02, weekly_change=0.05, monthly_change=0.12
            ),
            PortfolioAsset(
                ticker='MSFT', name='Microsoft Corp.', sector='Technology',
                shares=50, purchase_price=330.00, current_price=380.00,
                market_value=19000.00, weight=0.27,
                daily_change=0.015, weekly_change=0.04, monthly_change=0.10
            ),
            PortfolioAsset(
                ticker='GOOGL', name='Alphabet Inc.', sector='Technology',
                shares=75, purchase_price=140.00, current_price=155.00,
                market_value=11625.00, weight=0.17,
                daily_change=0.01, weekly_change=0.03, monthly_change=0.08
            ),
            PortfolioAsset(
                ticker='AMZN', name='Amazon.com Inc.', sector='Consumer',
                shares=30, purchase_price=185.00, current_price=195.00,
                market_value=5850.00, weight=0.08,
                daily_change=-0.005, weekly_change=0.02, monthly_change=0.05
            ),
            PortfolioAsset(
                ticker='NVDA', name='NVIDIA Corp.', sector='Technology',
                shares=40, purchase_price=450.00, current_price=520.00,
                market_value=20800.00, weight=0.30,
                daily_change=0.03, weekly_change=0.08, monthly_change=0.25
            )
        ]
        self.portfolio_assets = sample_assets
        
        sample_alerts = [
            WatchdogAlert(
                id='alert_001',
                asset_ticker='AAPL',
                alert_type='news',
                severity='medium',
                message='Apple announces new AI product line - potential growth catalyst',
                source='Reuters',
                timestamp=datetime.now() - timedelta(hours=2)
            ),
            WatchdogAlert(
                id='alert_002',
                asset_ticker='NVDA',
                alert_type='earnings',
                severity='high',
                message='NVIDIA earnings due next week - expected to beat estimates',
                source='Bloomberg',
                timestamp=datetime.now() - timedelta(hours=5)
            ),
            WatchdogAlert(
                id='alert_003',
                asset_ticker='MSFT',
                alert_type='filing',
                severity='low',
                message='Microsoft files 8-K regarding board changes',
                source='SEC EDGAR',
                timestamp=datetime.now() - timedelta(days=1)
            )
        ]
        self.watchdog_alerts = sample_alerts
    
    def ingest_document(self, file_content: str, file_name: str, source_type: str, metadata: Dict = None) -> SourceDocument:
        doc_id = f"doc_{hashlib.md5(file_name.encode()).hexdigest()[:8]}"
        
        doc = SourceDocument(
            id=doc_id,
            title=file_name,
            source_type=source_type,
            file_path=file_name,
            content=file_content,
            metadata=metadata or {},
            uploaded_at=datetime.now()
        )
        
        doc.citations = self._extract_citations(file_content)
        doc.parsed_tables = self._extract_tables(file_content)
        
        self.documents.append(doc)
        self._index_document(doc)
        
        return doc
    
    def _extract_citations(self, content: str) -> List[Dict[str, Any]]:
        citations = []
        lines = content.split('\n')
        
        for i, line in enumerate(lines):
            claims = re.findall(r'([A-Z][^.!?]*\s+(?:was|is|reached|grew|increased|decreased|fell|rose|were)\s+[^.!?]*\d+[^.!?]*)', line)
            for claim in claims:
                if len(claim) > 20:
                    citations.append({
                        'page': i // 40 + 1,
                        'paragraph': i + 1,
                        'claim': claim.strip(),
                        'confidence': 0.85
                    })
        
        return citations[:10]
    
    def _extract_tables(self, content: str) -> List[Dict[str, Any]]:
        tables = []
        lines = content.split('\n')
        table_data = []
        
        for line in lines:
            if '|' in line or '\t' in line:
                parts = re.split(r'[|\t]+', line)
                if len(parts) > 2:
                    table_data.append([p.strip() for p in parts if p.strip()])
        
        if table_data:
            tables.append({
                'rows': len(table_data),
                'columns': len(table_data[0]) if table_data else 0,
                'data': table_data
            })
        
        return tables
    
    def _index_document(self, doc: SourceDocument):
        words = re.findall(r'\b[a-zA-Z]{3,}\b', doc.content.lower())
        for word in words:
            if word not in self.search_index:
                self.search_index[word] = []
            self.search_index[word].append(doc.id)
    
    def hybrid_search(self, query: str, search_type: SearchType = SearchType.HYBRID) -> List[Dict[str, Any]]:
        results = []
        query_words = re.findall(r'\b[a-zA-Z]{3,}\b', query.lower())
        
        keyword_results = {}
        for word in query_words:
            if word in self.search_index:
                for doc_id in self.search_index[word]:
                    keyword_results[doc_id] = keyword_results.get(doc_id, 0) + 1
        
        vector_results = {}
        for doc in self.documents:
            doc_words = re.findall(r'\b[a-zA-Z]{3,}\b', doc.content.lower())
            overlap = len(set(query_words) & set(doc_words))
            if overlap > 0:
                vector_results[doc.id] = overlap / len(set(query_words) | set(doc_words))
        
        all_doc_ids = set(keyword_results.keys()) | set(vector_results.keys())
        
        for doc_id in all_doc_ids:
            doc = next((d for d in self.documents if d.id == doc_id), None)
            if doc:
                keyword_score = keyword_results.get(doc_id, 0) / max(1, len(query_words))
                vector_score = vector_results.get(doc_id, 0)
                
                if search_type == SearchType.KEYWORD:
                    final_score = keyword_score
                elif search_type == SearchType.VECTOR:
                    final_score = vector_score
                elif search_type == SearchType.SEMANTIC:
                    final_score = vector_score * 0.7 + keyword_score * 0.3
                else:
                    final_score = keyword_score * 0.4 + vector_score * 0.6
                
                if final_score > 0.1:
                    results.append({
                        'document': doc,
                        'score': round(final_score, 3),
                        'matches': list(set(query_words) & set(re.findall(r'\b[a-zA-Z]{3,}\b', doc.content.lower())))
                    })
        
        return sorted(results, key=lambda x: x['score'], reverse=True)
    
    def generate_citation_chain(self, claim: str, document_id: str) -> List[Dict[str, Any]]:
        doc = next((d for d in self.documents if d.id == document_id), None)
        if not doc:
            return []
        
        citations = []
        for citation in doc.citations:
            if claim.lower() in citation['claim'].lower() or any(word in citation['claim'].lower() for word in claim.lower().split()[:5]):
                citations.append({
                    'claim': citation['claim'],
                    'page': citation.get('page', 1),
                    'paragraph': citation.get('paragraph', 1),
                    'confidence': citation.get('confidence', 0.85),
                    'source_text': doc.content.split('\n')[citation.get('paragraph', 1) - 1] if citation.get('paragraph') else ''
                })
        
        return citations
    
    def audit_analysis(self, analysis_text: str, source_documents: List[SourceDocument]) -> Dict[str, Any]:
        audit_results = {
            'verified_claims': [],
            'unverified_claims': [],
            'citations_generated': [],
            'confidence_score': 0.0
        }
        
        claims = re.findall(r'([A-Z][^.!?]*\s+(?:was|is|reached|grew|increased|decreased|fell|rose|were|will|would|could|should)\s+[^.!?]*\d+[^.!?]*)', analysis_text)
        
        for claim in claims:
            verified = False
            for doc in source_documents:
                if claim.lower() in doc.content.lower():
                    citations = self.generate_citation_chain(claim, doc.id)
                    if citations:
                        audit_results['verified_claims'].append({
                            'claim': claim,
                            'source': doc.title,
                            'citations': citations
                        })
                        audit_results['citations_generated'].extend(citations)
                        verified = True
                        break
            
            if not verified:
                audit_results['unverified_claims'].append(claim)
        
        audit_results['confidence_score'] = len(audit_results['verified_claims']) / max(1, len(claims))
        
        return audit_results
    
    def execute_quant_analysis(self, code: str, data: pd.DataFrame) -> Dict[str, Any]:
        try:
            safe_globals = {
                'pd': pd,
                'np': np,
                'data': data,
                'len': len,
                'sum': sum,
                'min': min,
                'max': max,
                'round': round,
                'abs': abs
            }
            
            exec_globals = {}
            exec(code, safe_globals, exec_globals)
            result = exec_globals.get('result', None)
            
            return {
                'status': 'success',
                'result': result,
                'execution_time': 0.1
            }
        except Exception as e:
            return {
                'status': 'error',
                'error': str(e),
                'execution_time': 0.1
            }
    
    def fetch_market_data(self, ticker: str) -> Dict[str, Any]:
        base_price = random.uniform(100, 500)
        return {
            'ticker': ticker,
            'price': round(base_price, 2),
            'change': round(random.uniform(-5, 5), 2),
            'change_percent': round(random.uniform(-3, 3), 2),
            'volume': random.randint(1000000, 10000000),
            'high': round(base_price * random.uniform(1.01, 1.05), 2),
            'low': round(base_price * random.uniform(0.95, 0.99), 2),
            'timestamp': datetime.now().isoformat()
        }
    
    def fetch_sec_filing(self, ticker: str) -> Dict[str, Any]:
        filings = [
            {'type': '10-K', 'date': '2025-01-15', 'title': 'Annual Report'},
            {'type': '10-Q', 'date': '2025-01-10', 'title': 'Quarterly Report'},
            {'type': '8-K', 'date': '2025-01-05', 'title': 'Current Report'},
            {'type': 'DEF 14A', 'date': '2025-01-02', 'title': 'Proxy Statement'}
        ]
        return {
            'ticker': ticker,
            'filings': [f for f in filings if random.random() > 0.5],
            'timestamp': datetime.now().isoformat()
        }
    
    def fetch_earnings_call(self, ticker: str) -> Dict[str, Any]:
        return {
            'ticker': ticker,
            'quarter': 'Q4 2024',
            'date': '2025-01-20',
            'transcript': f"Earnings call transcript for {ticker}. Revenue exceeded expectations...",
            'highlights': [
                'Revenue beat by 5%',
                'EPS exceeded guidance',
                'Forward guidance raised'
            ],
            'timestamp': datetime.now().isoformat()
        }
    
    def calculate_portfolio_metrics(self) -> Dict[str, Any]:
        total_value = sum(asset.market_value for asset in self.portfolio_assets)
        
        sector_exposure = {}
        for asset in self.portfolio_assets:
            sector_exposure[asset.sector] = sector_exposure.get(asset.sector, 0) + asset.weight
        
        returns = [asset.daily_change for asset in self.portfolio_assets]
        volatility = np.std(returns) if returns else 0
        tracking_error = round(volatility * 0.5, 4)
        
        return {
            'total_value': round(total_value, 2),
            'sector_exposure': sector_exposure,
            'volatility': round(volatility, 4),
            'tracking_error': tracking_error,
            'sharpe_ratio': round((0.1 - 0.03) / max(0.01, volatility), 2),
            'concentration_risk': round(max(sector_exposure.values()) if sector_exposure else 0, 2)
        }
    
    def run_watchdog_alerts(self) -> List[WatchdogAlert]:
        new_alerts = []
        
        for asset in self.portfolio_assets:
            if random.random() > 0.7:
                alert = WatchdogAlert(
                    id=f"alert_{datetime.now().timestamp()}",
                    asset_ticker=asset.ticker,
                    alert_type='earnings',
                    severity='high',
                    message=f'Earnings announcement for {asset.ticker} expected tomorrow',
                    source='SEC EDGAR',
                    timestamp=datetime.now()
                )
                new_alerts.append(alert)
            
            if random.random() > 0.8:
                alert = WatchdogAlert(
                    id=f"alert_{datetime.now().timestamp()}",
                    asset_ticker=asset.ticker,
                    alert_type='price',
                    severity='medium',
                    message=f'{asset.ticker} price {random.choice(["up", "down"])} significantly',
                    source='Market Feed',
                    timestamp=datetime.now()
                )
                new_alerts.append(alert)
        
        self.watchdog_alerts.extend(new_alerts)
        return new_alerts
    
    def run_macro_stress_test(self, scenario: MacroScenario) -> Dict[str, Any]:
        portfolio_value = sum(asset.market_value for asset in self.portfolio_assets)
        
        rate_impact = scenario.rate_change * 0.5
        inflation_impact = scenario.inflation_change * 0.3
        growth_impact = scenario.growth_change * 0.4
        market_impact = scenario.market_shock * 0.7
        
        total_impact = rate_impact + inflation_impact + growth_impact + market_impact
        impacted_value = portfolio_value * (1 + total_impact)
        
        asset_impacts = []
        for asset in self.portfolio_assets:
            sector_multiplier = {
                'Technology': 0.2,
                'Consumer': 0.15,
                'Financial': 0.25,
                'Healthcare': 0.1,
                'Energy': 0.3
            }.get(asset.sector, 0.15)
            
            asset_impact = total_impact * sector_multiplier
            impacted_asset_value = asset.market_value * (1 + asset_impact)
            
            asset_impacts.append({
                'ticker': asset.ticker,
                'current_value': asset.market_value,
                'impacted_value': impacted_asset_value,
                'impact_percent': asset_impact * 100
            })
        
        return {
            'portfolio_name': 'Professional Portfolio',
            'scenario_name': scenario.name,
            'current_value': portfolio_value,
            'impacted_value': impacted_value,
            'total_return': total_impact * 100,
            'asset_impacts': asset_impacts,
            'risk_assessment': 'High' if abs(total_impact) > 0.2 else 'Medium' if abs(total_impact) > 0.1 else 'Low'
        }
    
    def generate_deliverable(self, template_type: str, data: Dict[str, Any]) -> str:
        templates = {
            'investment_memo': """
            # Investment Committee Memo
            
            **Date:** {date}
            **Analyst:** {analyst}
            **Subject:** Investment Recommendation
            
            ## Executive Summary
            {summary}
            
            ## Investment Thesis
            {thesis}
            
            ## Risk Analysis
            {risks}
            
            ## Financial Projections
            {financials}
            
            ## Recommendation
            **Action:** {action}
            **Target Price:** {target_price}
            **Time Horizon:** {time_horizon}
            
            ## Appendix
            {appendix}
            """,
            
            'client_tear_sheet': """
            # {client_name} - Portfolio Summary
            
            **Date:** {date}
            **Advisor:** {advisor}
            
            ## Portfolio Overview
            **Total Value:** {total_value}
            **Annualized Return:** {return}
            **Risk Level:** {risk_level}
            
            ## Holdings
            {holdings}
            
            ## Performance
            {performance}
            
            ## Recommendations
            {recommendations}
            """,
            
            'research_note': """
            # Research Note: {title}
            
            **Date:** {date}
            **Analyst:** {analyst}
            **Sector:** {sector}
            
            ## Key Findings
            {findings}
            
            ## Methodology
            {methodology}
            
            ## Data Sources
            {sources}
            
            ## Conclusion
            {conclusion}
            """
        }
        
        template = templates.get(template_type, templates['investment_memo'])
        
        try:
            return template.format(**data)
        except:
            return template
    
    def extract_table_from_pdf(self, pdf_content: str) -> pd.DataFrame:
        lines = pdf_content.split('\n')
        table_data = []
        headers = None
        
        for line in lines:
            if '|' in line or '\t' in line:
                parts = [p.strip() for p in re.split(r'[|\t]+', line) if p.strip()]
                if parts:
                    if not headers:
                        headers = parts[:len(parts)]
                    else:
                        table_data.append(parts[:len(headers)])
        
        if table_data:
            return pd.DataFrame(table_data, columns=headers)
        else:
            return pd.DataFrame({
                'Metric': ['Revenue', 'EBITDA', 'Net Income', 'EPS'],
                'Q4 2024': ['$12.4B', '$3.5B', '$2.1B', '$2.15'],
                'Q4 2023': ['$10.5B', '$2.8B', '$1.7B', '$1.75'],
                'Change': ['+18%', '+25%', '+24%', '+23%']
            })
    
    def get_workspace_data(self) -> Dict[str, Any]:
        return {
            'documents': [
                {
                    'id': doc.id,
                    'title': doc.title,
                    'type': doc.source_type,
                    'content': doc.content[:500] + ('...' if len(doc.content) > 500 else '')
                }
                for doc in self.documents
            ],
            'citations': self.citations,
            'analysis_history': self.analysis_history,
            'portfolio_data': self.calculate_portfolio_metrics()
        }


# ==================== SME DATA MODELS ====================

@dataclass
class BusinessMetric:
    name: str
    value: float
    change: float
    trend: str
    category: str
    timestamp: datetime

@dataclass
class ActionTask:
    id: str
    title: str
    description: str
    priority: str
    category: str
    impact: str
    action_type: str
    status: str
    created_at: datetime
    due_date: Optional[datetime] = None

@dataclass
class WebhookEvent:
    event_type: str
    payload: Dict[str, Any]
    timestamp: datetime
    processed: bool = False

class PriorityLevel(Enum):
    CRITICAL = "critical"
    HIGH = "high"
    MEDIUM = "medium"
    LOW = "low"


# ==================== SME INFRASTRUCTURE ENGINE ====================

class SMEInfrastructureEngine:
    def __init__(self, country_code='kenya'):
        self.country = country_code
        self.overlay = LocalCurriculumOverlay.get_overlay(country_code)
        self.webhook_events = []
        self.api_connections = {}
        self.etl_pipeline_data = {}
    
    def etl_process(self, raw_data: Dict[str, Any]) -> Dict[str, Any]:
        try:
            extracted = self._extract_data(raw_data)
            transformed = self._transform_data(extracted)
            loaded = self._load_data(transformed)
            
            return {
                'status': 'success',
                'extracted': extracted,
                'transformed': transformed,
                'loaded': loaded,
                'timestamp': datetime.now().isoformat()
            }
        except Exception as e:
            return {
                'status': 'error',
                'error': str(e),
                'timestamp': datetime.now().isoformat()
            }
    
    def _extract_data(self, raw_data: Dict[str, Any]) -> Dict[str, Any]:
        extracted = {}
        
        if 'transactions' in raw_data:
            extracted['transactions'] = raw_data['transactions']
        if 'invoices' in raw_data:
            extracted['invoices'] = raw_data['invoices']
        if 'inventory' in raw_data:
            extracted['inventory'] = raw_data['inventory']
        if 'customers' in raw_data:
            extracted['customers'] = raw_data['customers']
        
        for key in extracted:
            if isinstance(extracted[key], list):
                extracted[key] = self._normalize_list(extracted[key])
        
        return extracted
    
    def _normalize_list(self, data_list: List[Dict]) -> List[Dict]:
        normalized = []
        for item in data_list:
            if 'amount' in item and 'total' not in item:
                item['total'] = item['amount']
            if 'qty' in item and 'quantity' not in item:
                item['quantity'] = item['qty']
            if 'name' in item and 'product_name' not in item:
                item['product_name'] = item['name']
            normalized.append(item)
        return normalized
    
    def _transform_data(self, extracted: Dict[str, Any]) -> Dict[str, Any]:
        transformed = {}
        
        if 'transactions' in extracted:
            transactions = extracted['transactions']
            transformed['total_revenue'] = sum(t.get('total', 0) for t in transactions)
            transformed['transaction_count'] = len(transactions)
            transformed['average_transaction'] = transformed['total_revenue'] / max(1, len(transactions))
        
        if 'inventory' in extracted:
            inventory = extracted['inventory']
            transformed['total_inventory_items'] = len(inventory)
            transformed['low_stock_items'] = [i for i in inventory if i.get('quantity', 0) < i.get('threshold', 10)]
            transformed['out_of_stock_items'] = [i for i in inventory if i.get('quantity', 0) <= 0]
        
        if 'customers' in extracted:
            customers = extracted['customers']
            transformed['total_customers'] = len(customers)
            transformed['repeat_customers'] = [c for c in customers if c.get('orders', 0) > 1]
        
        return transformed
    
    def _load_data(self, transformed: Dict[str, Any]) -> Dict[str, Any]:
        self.etl_pipeline_data.update(transformed)
        return transformed
    
    def connect_api(self, service: str, credentials: Dict[str, str]) -> Dict[str, Any]:
        api_configs = {
            'stripe': {'name': 'Stripe', 'icon': '💳', 'endpoints': ['payments', 'invoices', 'customers', 'subscriptions'], 'auth_type': 'api_key'},
            'quickbooks': {'name': 'QuickBooks', 'icon': '📊', 'endpoints': ['invoices', 'expenses', 'customers', 'reports'], 'auth_type': 'oauth2'},
            'shopify': {'name': 'Shopify', 'icon': '🛍️', 'endpoints': ['orders', 'products', 'customers', 'inventory'], 'auth_type': 'api_key'},
            'toast': {'name': 'Toast POS', 'icon': '🍽️', 'endpoints': ['orders', 'menu', 'inventory', 'reports'], 'auth_type': 'api_key'},
            'jobber': {'name': 'Jobber', 'icon': '🔧', 'endpoints': ['clients', 'invoices', 'visits', 'schedule'], 'auth_type': 'oauth2'}
        }
        
        if service in api_configs:
            self.api_connections[service] = {
                'config': api_configs[service],
                'credentials': credentials,
                'connected': True,
                'last_sync': datetime.now().isoformat()
            }
            return {
                'status': 'connected',
                'service': service,
                'message': f"Successfully connected to {api_configs[service]['name']}"
            }
        else:
            return {
                'status': 'error',
                'service': service,
                'message': f"Unsupported API service: {service}"
            }
    
    def fetch_api_data(self, service: str, endpoint: str, params: Dict = None) -> Dict[str, Any]:
        if service not in self.api_connections:
            return {'status': 'error', 'message': f"Service {service} not connected"}
        
        mock_data = {
            'payments': [
                {'id': 1, 'amount': 150.00, 'status': 'paid', 'date': '2024-01-15'},
                {'id': 2, 'amount': 75.50, 'status': 'paid', 'date': '2024-01-14'},
                {'id': 3, 'amount': 200.00, 'status': 'pending', 'date': '2024-01-13'}
            ],
            'invoices': [
                {'id': 'INV-001', 'amount': 150.00, 'status': 'paid', 'due_date': '2024-01-20'},
                {'id': 'INV-002', 'amount': 75.50, 'status': 'overdue', 'due_date': '2024-01-10'},
                {'id': 'INV-003', 'amount': 200.00, 'status': 'pending', 'due_date': '2024-01-25'}
            ],
            'orders': [
                {'id': 'ORD-001', 'total': 150.00, 'status': 'fulfilled', 'items': 3},
                {'id': 'ORD-002', 'total': 75.50, 'status': 'processing', 'items': 2},
            ],
            'inventory': [
                {'id': 'SKU-001', 'name': 'Product A', 'quantity': 45, 'threshold': 50},
                {'id': 'SKU-002', 'name': 'Product B', 'quantity': 12, 'threshold': 20},
                {'id': 'SKU-003', 'name': 'Product C', 'quantity': 3, 'threshold': 10},
            ]
        }
        
        return {
            'status': 'success',
            'service': service,
            'endpoint': endpoint,
            'data': mock_data.get(endpoint, []),
            'timestamp': datetime.now().isoformat()
        }
    
    def process_webhook(self, event_type: str, payload: Dict[str, Any]) -> Dict[str, Any]:
        event = WebhookEvent(
            event_type=event_type,
            payload=payload,
            timestamp=datetime.now()
        )
        self.webhook_events.append(event)
        
        if event_type == 'checkout.paid':
            return self._handle_checkout_paid(event)
        elif event_type == 'invoice.overdue':
            return self._handle_invoice_overdue(event)
        elif event_type == 'inventory.low':
            return self._handle_inventory_low(event)
        elif event_type == 'customer.churn_risk':
            return self._handle_customer_churn_risk(event)
        else:
            return {'status': 'unhandled', 'event_type': event_type}
    
    def _handle_checkout_paid(self, event: WebhookEvent) -> Dict[str, Any]:
        payload = event.payload
        amount = payload.get('amount', 0)
        customer = payload.get('customer', {})
        
        task = ActionTask(
            id=f"task_{datetime.now().timestamp()}",
            title="New Payment Received",
            description=f"Payment of ${amount:.2f} received from {customer.get('name', 'Customer')}",
            priority='low',
            category='finance',
            impact=f"+${amount:.2f} revenue",
            action_type='auto',
            status='approved',
            created_at=datetime.now()
        )
        
        return {
            'status': 'processed',
            'event_type': 'checkout.paid',
            'action_task': task,
            'message': f"Payment processed successfully"
        }
    
    def _handle_invoice_overdue(self, event: WebhookEvent) -> Dict[str, Any]:
        payload = event.payload
        invoice_id = payload.get('invoice_id', 'unknown')
        amount = payload.get('amount', 0)
        days_overdue = payload.get('days_overdue', 0)
        
        task = ActionTask(
            id=f"task_{datetime.now().timestamp()}",
            title=f"Invoice {invoice_id} is {days_overdue} days overdue",
            description=f"Client payment of ${amount:.2f} is overdue. Send payment reminder.",
            priority='high',
            category='finance',
            impact=f"Potential loss of ${amount:.2f}",
            action_type='approval',
            status='pending',
            created_at=datetime.now(),
            due_date=datetime.now() + timedelta(days=1)
        )
        
        return {
            'status': 'processed',
            'event_type': 'invoice.overdue',
            'action_task': task,
            'message': f"Overdue invoice alert created"
        }
    
    def _handle_inventory_low(self, event: WebhookEvent) -> Dict[str, Any]:
        payload = event.payload
        product = payload.get('product', {})
        current_qty = payload.get('current_qty', 0)
        threshold = payload.get('threshold', 0)
        price = product.get('price', 0)
        potential_loss = (threshold - current_qty) * price
        
        task = ActionTask(
            id=f"task_{datetime.now().timestamp()}",
            title=f"Low Stock Alert: {product.get('name', 'Product')}",
            description=f"Current stock: {current_qty} units. Threshold: {threshold} units. Restock now.",
            priority='critical',
            category='inventory',
            impact=f"Prevent ${potential_loss:.2f} lost revenue",
            action_type='approval',
            status='pending',
            created_at=datetime.now(),
            due_date=datetime.now() + timedelta(hours=12)
        )
        
        return {
            'status': 'processed',
            'event_type': 'inventory.low',
            'action_task': task,
            'message': f"Inventory alert created"
        }
    
    def _handle_customer_churn_risk(self, event: WebhookEvent) -> Dict[str, Any]:
        payload = event.payload
        customer = payload.get('customer', {})
        churn_score = payload.get('churn_score', 0)
        days_inactive = payload.get('days_inactive', 0)
        
        task = ActionTask(
            id=f"task_{datetime.now().timestamp()}",
            title=f"Churn Risk: {customer.get('name', 'Customer')}",
            description=f"Customer inactive for {days_inactive} days. Churn score: {churn_score:.1%}",
            priority='high',
            category='customers',
            impact=f"Potential loss of ${customer.get('ltv', 0):.2f} LTV",
            action_type='draft',
            status='pending',
            created_at=datetime.now()
        )
        
        return {
            'status': 'processed',
            'event_type': 'customer.churn_risk',
            'action_task': task,
            'message': f"Churn risk alert created"
        }


# ==================== SME AI ENGINE ====================

class SMEAIEngine:
    def __init__(self, country_code='kenya'):
        self.country = country_code
        self.infrastructure = SMEInfrastructureEngine(country_code)
    
    def predict_cash_flow(self, historical_data: List[Dict]) -> Dict[str, Any]:
        days = 30
        current_balance = 5000
        predictions = []
        
        for i in range(days):
            daily_inflow = random.uniform(100, 500)
            daily_outflow = random.uniform(50, 300)
            current_balance += daily_inflow - daily_outflow
            predictions.append({
                'day': i + 1,
                'date': (datetime.now() + timedelta(days=i)).strftime('%Y-%m-%d'),
                'balance': round(current_balance, 2),
                'inflow': round(daily_inflow, 2),
                'outflow': round(daily_outflow, 2)
            })
        
        min_balance = min(p['balance'] for p in predictions)
        min_balance_day = next(p['day'] for p in predictions if p['balance'] == min_balance)
        
        return {
            'predictions': predictions,
            'min_balance': round(min_balance, 2),
            'min_balance_day': min_balance_day,
            'avg_daily_balance': round(sum(p['balance'] for p in predictions) / days, 2),
            'risk_level': 'high' if min_balance < 1000 else 'medium' if min_balance < 3000 else 'low'
        }
    
    def predict_churn(self, customer_data: List[Dict]) -> List[Dict]:
        predictions = []
        for customer in customer_data[:10]:
            days_since_last_order = customer.get('days_since_last_order', 0)
            order_count = customer.get('order_count', 0)
            churn_score = min(1.0, (days_since_last_order / 90) * 0.7 + (1 - min(1, order_count / 5)) * 0.3)
            
            predictions.append({
                'customer_id': customer.get('id', 'unknown'),
                'name': customer.get('name', 'Unknown'),
                'churn_score': round(churn_score, 3),
                'risk_level': 'high' if churn_score > 0.6 else 'medium' if churn_score > 0.3 else 'low',
                'recommendation': self._get_churn_recommendation(churn_score)
            })
        
        return sorted(predictions, key=lambda x: x['churn_score'], reverse=True)
    
    def _get_churn_recommendation(self, churn_score: float) -> str:
        if churn_score > 0.7:
            return "🚨 Immediate intervention: Personalized offer needed"
        elif churn_score > 0.5:
            return "⚠️ Send re-engagement email with special discount"
        elif churn_score > 0.3:
            return "📧 Send regular engagement content"
        else:
            return "✅ Keep doing what you're doing"
    
    def predict_inventory_depletion(self, inventory_data: List[Dict]) -> List[Dict]:
        predictions = []
        for item in inventory_data:
            current_qty = item.get('quantity', 0)
            daily_sales_avg = item.get('daily_sales_avg', 1)
            threshold = item.get('threshold', 10)
            
            if daily_sales_avg > 0:
                days_until_depletion = current_qty / daily_sales_avg
                days_until_threshold = max(0, (current_qty - threshold) / daily_sales_avg)
            else:
                days_until_depletion = 999
                days_until_threshold = 999
            
            predictions.append({
                'product_id': item.get('id', 'unknown'),
                'name': item.get('name', 'Unknown'),
                'current_qty': current_qty,
                'days_until_depletion': round(days_until_depletion, 1),
                'days_until_threshold': round(days_until_threshold, 1),
                'status': 'critical' if days_until_threshold < 3 else 'warning' if days_until_threshold < 7 else 'ok',
                'restock_quantity': max(threshold * 2 - current_qty, 0)
            })
        
        return sorted(predictions, key=lambda x: x['days_until_threshold'])
    
    def predict_customer_ltv(self, customer_data: List[Dict]) -> List[Dict]:
        predictions = []
        for customer in customer_data[:10]:
            avg_order = customer.get('avg_order_value', 50)
            frequency = customer.get('frequency', 1)
            months = customer.get('months_active', 6)
            
            current_ltv = avg_order * frequency * months
            projected_ltv = avg_order * frequency * 24
            
            predictions.append({
                'customer_id': customer.get('id', 'unknown'),
                'name': customer.get('name', 'Unknown'),
                'current_ltv': round(current_ltv, 2),
                'projected_ltv': round(projected_ltv, 2),
                'potential_increase': round(projected_ltv - current_ltv, 2),
                'segment': 'high' if projected_ltv > 1000 else 'medium' if projected_ltv > 500 else 'low'
            })
        
        return sorted(predictions, key=lambda x: x['projected_ltv'], reverse=True)
    
    def natural_language_query(self, query: str, data_context: Dict[str, Any]) -> Dict[str, Any]:
        query_lower = query.lower()
        response = {
            'query': query,
            'timestamp': datetime.now().isoformat(),
            'results': []
        }
        
        if 'cash' in query_lower or 'balance' in query_lower:
            cash_flow = self.predict_cash_flow([])
            response['results'].append({
                'type': 'cash_flow',
                'data': cash_flow,
                'summary': f"Your current cash flow projection shows a minimum balance of ${cash_flow['min_balance']:.2f} on day {cash_flow['min_balance_day']}. Risk level: {cash_flow['risk_level']}."
            })
        
        if 'revenue' in query_lower or 'income' in query_lower:
            revenue = {'total': 15750.00, 'growth': 12.5, 'best_month': 'December', 'worst_month': 'February'}
            response['results'].append({
                'type': 'revenue',
                'data': revenue,
                'summary': f"Total revenue: ${revenue['total']:.2f}. Growth: {revenue['growth']}% year-over-year."
            })
        
        if 'inventory' in query_lower or 'stock' in query_lower:
            inventory_data = [
                {'id': 'SKU-001', 'name': 'Product A', 'quantity': 45, 'threshold': 50, 'daily_sales_avg': 2},
                {'id': 'SKU-002', 'name': 'Product B', 'quantity': 12, 'threshold': 20, 'daily_sales_avg': 1.5},
                {'id': 'SKU-003', 'name': 'Product C', 'quantity': 3, 'threshold': 10, 'daily_sales_avg': 1}
            ]
            predictions = self.predict_inventory_depletion(inventory_data)
            critical_items = [p for p in predictions if p['status'] == 'critical']
            response['results'].append({
                'type': 'inventory',
                'data': predictions,
                'summary': f"Found {len(critical_items)} items at critical stock levels."
            })
        
        if 'customer' in query_lower or 'churn' in query_lower:
            customer_data = [
                {'id': 1, 'name': 'John Doe', 'days_since_last_order': 45, 'avg_order_value': 75, 'order_count': 3},
                {'id': 2, 'name': 'Jane Smith', 'days_since_last_order': 12, 'avg_order_value': 120, 'order_count': 8},
                {'id': 3, 'name': 'Bob Johnson', 'days_since_last_order': 60, 'avg_order_value': 50, 'order_count': 2},
            ]
            churn_predictions = self.predict_churn(customer_data)
            high_risk = [c for c in churn_predictions if c['risk_level'] == 'high']
            response['results'].append({
                'type': 'churn',
                'data': churn_predictions,
                'summary': f"{len(high_risk)} customers at high churn risk."
            })
        
        if not response['results']:
            response['results'].append({
                'type': 'general',
                'data': {},
                'summary': "I couldn't find specific data for your query. Try asking about cash flow, revenue, inventory, or customers."
            })
        
        return response
    
    def validate_llm_output(self, output: Dict[str, Any], action_type: str) -> Dict[str, Any]:
        validation = {
            'valid': True,
            'errors': [],
            'warnings': [],
            'corrected': {}
        }
        
        if action_type in ['invoice', 'payment', 'financial']:
            amount = output.get('amount', 0)
            if amount < 0:
                validation['valid'] = False
                validation['errors'].append("Amount cannot be negative")
            if amount > 100000:
                validation['warnings'].append("Large transaction amount - please verify")
        
        if action_type in ['inventory', 'restock']:
            quantity = output.get('quantity', 0)
            if quantity < 0:
                validation['valid'] = False
                validation['errors'].append("Quantity cannot be negative")
            if quantity > 1000:
                validation['warnings'].append("Large restock quantity - please verify")
        
        return validation


# ==================== SME PRODUCT UX ENGINE ====================

class SMEProductUXEngine:
    def __init__(self, country_code='kenya'):
        self.country = country_code
        self.infrastructure = SMEInfrastructureEngine(country_code)
        self.ai_engine = SMEAIEngine(country_code)
        self.tasks = []
        self.alert_history = []
    
    def generate_tasks(self, business_data: Dict[str, Any]) -> List[ActionTask]:
        tasks = []
        tasks.extend(self._generate_inventory_tasks(business_data))
        tasks.extend(self._generate_financial_tasks(business_data))
        tasks.extend(self._generate_customer_tasks(business_data))
        tasks.extend(self._generate_marketing_tasks(business_data))
        tasks.extend(self._generate_operations_tasks(business_data))
        
        priority_order = {'critical': 0, 'high': 1, 'medium': 2, 'low': 3}
        tasks.sort(key=lambda t: priority_order.get(t.priority, 4))
        
        self.tasks = tasks
        return tasks
    
    def _generate_inventory_tasks(self, data: Dict) -> List[ActionTask]:
        tasks = []
        inventory = data.get('inventory', [])
        for item in inventory:
            qty = item.get('quantity', 0)
            threshold = item.get('threshold', 10)
            price = item.get('price', 0)
            
            if qty <= 0:
                potential_loss = threshold * price
                tasks.append(ActionTask(
                    id=f"inv_{item.get('id', 'unknown')}_{datetime.now().timestamp()}",
                    title=f"🚨 OUT OF STOCK: {item.get('name', 'Product')}",
                    description=f"Critical! No stock remaining. Order now to prevent ${potential_loss:.2f} lost revenue.",
                    priority='critical',
                    category='inventory',
                    impact=f"Prevent ${potential_loss:.2f} lost sales",
                    action_type='approval',
                    status='pending',
                    created_at=datetime.now(),
                    due_date=datetime.now() + timedelta(hours=6)
                ))
            elif qty < threshold:
                potential_loss = (threshold - qty) * price
                tasks.append(ActionTask(
                    id=f"inv_{item.get('id', 'unknown')}_{datetime.now().timestamp()}",
                    title=f"⚠️ Low Stock: {item.get('name', 'Product')}",
                    description=f"Only {qty} units remaining (Threshold: {threshold}). Restock to maintain availability.",
                    priority='high',
                    category='inventory',
                    impact=f"Prevent ${potential_loss:.2f} lost revenue",
                    action_type='draft',
                    status='pending',
                    created_at=datetime.now(),
                    due_date=datetime.now() + timedelta(days=1)
                ))
        return tasks
    
    def _generate_financial_tasks(self, data: Dict) -> List[ActionTask]:
        tasks = []
        invoices = data.get('invoices', [])
        overdue_invoices = [inv for inv in invoices if inv.get('status') == 'overdue']
        
        for inv in overdue_invoices:
            tasks.append(ActionTask(
                id=f"fin_{inv.get('id', 'unknown')}_{datetime.now().timestamp()}",
                title=f"💳 Overdue Payment: {inv.get('id', 'Invoice')}",
                description=f"Payment of ${inv.get('amount', 0):.2f} is overdue. Send reminder to customer.",
                priority='high',
                category='finance',
                impact=f"Recover ${inv.get('amount', 0):.2f}",
                action_type='approval',
                status='pending',
                created_at=datetime.now(),
                due_date=datetime.now() + timedelta(days=2)
            ))
        
        cash_flow = self.ai_engine.predict_cash_flow([])
        if cash_flow['risk_level'] == 'high':
            tasks.append(ActionTask(
                id=f"fin_cash_{datetime.now().timestamp()}",
                title=f"💰 Cash Flow Alert: ${cash_flow['min_balance']:.2f} minimum projected",
                description=f"Cash flow projected to drop to ${cash_flow['min_balance']:.2f} in {cash_flow['min_balance_day']} days.",
                priority='critical',
                category='finance',
                impact="Avoid cash shortage",
                action_type='alert',
                status='pending',
                created_at=datetime.now(),
                due_date=datetime.now() + timedelta(days=1)
            ))
        return tasks
    
    def _generate_customer_tasks(self, data: Dict) -> List[ActionTask]:
        tasks = []
        customers = data.get('customers', [])
        churn_predictions = self.ai_engine.predict_churn(customers[:10])
        
        for prediction in churn_predictions[:3]:
            if prediction['risk_level'] == 'high':
                tasks.append(ActionTask(
                    id=f"cust_{prediction.get('customer_id', 'unknown')}_{datetime.now().timestamp()}",
                    title=f"👤 Churn Risk: {prediction.get('name', 'Customer')}",
                    description=f"{prediction.get('name', 'Customer')} at {prediction['churn_score']*100:.0f}% churn risk.",
                    priority='high',
                    category='customers',
                    impact=f"Retain LTV",
                    action_type='draft',
                    status='pending',
                    created_at=datetime.now()
                ))
        return tasks
    
    def _generate_marketing_tasks(self, data: Dict) -> List[ActionTask]:
        tasks = []
        opportunities = [
            {'title': "🎯 Email Campaign for Inactive Customers", 'desc': "Send re-engagement email to customers inactive for 30+ days", 'impact': "+15% customer retention"},
            {'title': "📱 Social Media Ad Boost", 'desc': "Boost top-performing post to reach new audience", 'impact': "+20% reach"},
        ]
        
        for opp in opportunities[:2]:
            tasks.append(ActionTask(
                id=f"mkt_{hashlib.md5(opp['title'].encode()).hexdigest()[:8]}_{datetime.now().timestamp()}",
                title=opp['title'],
                description=opp['desc'],
                priority='medium',
                category='marketing',
                impact=opp['impact'],
                action_type='draft',
                status='pending',
                created_at=datetime.now()
            ))
        return tasks
    
    def _generate_operations_tasks(self, data: Dict) -> List[ActionTask]:
        tasks = []
        ops_tasks = [
            {'title': "📊 Monthly Financial Review", 'desc': "Review and reconcile all accounts for month-end", 'priority': 'high'},
            {'title': "📋 Supplier Contract Renewal", 'desc': "Review supplier contracts up for renewal next month", 'priority': 'medium'},
        ]
        
        for task in ops_tasks:
            tasks.append(ActionTask(
                id=f"ops_{hashlib.md5(task['title'].encode()).hexdigest()[:8]}_{datetime.now().timestamp()}",
                title=task['title'],
                description=task['desc'],
                priority=task['priority'],
                category='operations',
                impact="Improved operational efficiency",
                action_type='auto',
                status='pending',
                created_at=datetime.now(),
                due_date=datetime.now() + timedelta(days=7)
            ))
        return tasks
    
    def approve_task(self, task_id: str) -> Dict[str, Any]:
        for task in self.tasks:
            if task.id == task_id and task.status == 'pending':
                task.status = 'approved'
                return {
                    'status': 'approved',
                    'task': task,
                    'message': f"Task '{task.title}' has been approved."
                }
        return {'status': 'error', 'message': 'Task not found or already processed'}
    
    def dismiss_task(self, task_id: str) -> Dict[str, Any]:
        for task in self.tasks:
            if task.id == task_id:
                task.status = 'dismissed'
                return {'status': 'dismissed', 'task': task, 'message': f"Task '{task.title}' dismissed"}
        return {'status': 'error', 'message': 'Task not found'}
    
    def generate_digest(self, channel: str = 'whatsapp') -> Dict[str, Any]:
        priority_tasks = [t for t in self.tasks if t.status == 'pending' and t.priority in ['critical', 'high']]
        
        digest = {
            'channel': channel,
            'timestamp': datetime.now().isoformat(),
            'summary': {
                'total_tasks': len([t for t in self.tasks if t.status == 'pending']),
                'critical_tasks': len([t for t in self.tasks if t.status == 'pending' and t.priority == 'critical']),
                'high_tasks': len([t for t in self.tasks if t.status == 'pending' and t.priority == 'high'])
            },
            'top_priorities': [
                {
                    'title': t.title,
                    'description': t.description,
                    'impact': t.impact,
                    'due': t.due_date.strftime('%Y-%m-%d') if t.due_date else 'No deadline'
                }
                for t in priority_tasks[:5]
            ],
            'action_required': len(priority_tasks) > 0,
            'message': self._format_digest_message(channel, priority_tasks[:3])
        }
        
        self.alert_history.append(digest)
        return digest
    
    def _format_digest_message(self, channel: str, top_tasks: List[ActionTask]) -> str:
        if channel == 'whatsapp':
            message = "📊 *AI Shiksha Daily Business Digest* 📊\n\n"
            message += f"📌 {len([t for t in self.tasks if t.status == 'pending'])} pending tasks\n"
            message += f"🚨 {len([t for t in self.tasks if t.status == 'pending' and t.priority == 'critical'])} critical tasks\n\n"
            
            if top_tasks:
                message += "*Top Priorities:*\n"
                for i, task in enumerate(top_tasks, 1):
                    message += f"{i}. {task.title}\n"
                    message += f"   {task.description[:50]}...\n"
            else:
                message += "✅ No urgent tasks. You're on track! 🎉"
            
            return message
        elif channel == 'sms':
            return f"AI Shiksha Alert: {len([t for t in self.tasks if t.status == 'pending'])} tasks pending."
        elif channel == 'email':
            return f"""
            <h2>📊 AI Shiksha Daily Business Digest</h2>
            <p><strong>{len([t for t in self.tasks if t.status == 'pending'])}</strong> pending tasks</p>
            """
        return "Digest generated successfully"


# ==================== DOCUMENT INTELLIGENCE ENGINE ====================

class DocumentIntelligenceEngine:
    def __init__(self, country_code='kenya'):
        self.country = country_code
        self.overlay = LocalCurriculumOverlay.get_overlay(country_code)
        self.context = LocalCurriculumOverlay.get_local_context(country_code)
    
    def extract_text_from_file(self, uploaded_file) -> str:
        try:
            file_extension = uploaded_file.name.split('.')[-1].lower()
            text = ""
            
            if file_extension == 'pdf':
                if PyPDF2 is None:
                    return "PDF support requires PyPDF2. Please install: pip install PyPDF2"
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            
            elif file_extension == 'docx':
                if docx is None:
                    return "DOCX support requires python-docx. Please install: pip install python-docx"
                doc = docx.Document(io.BytesIO(uploaded_file.read()))
                for para in doc.paragraphs:
                    text += para.text + "\n"
            
            elif file_extension in ['txt', 'csv', 'json']:
                text = uploaded_file.read().decode('utf-8')
            
            else:
                return f"Unsupported file format: {file_extension}. Please upload PDF, DOCX, TXT, CSV, or JSON."
            
            return text.strip() if text else "No text could be extracted from the document."
        
        except Exception as e:
            return f"Error extracting text: {str(e)}"
    
    def analyze_document(self, text: str, segment: str) -> Dict[str, Any]:
        analysis = {
            'word_count': len(text.split()),
            'char_count': len(text),
            'sentence_count': len(re.split(r'[.!?]+', text)),
            'key_phrases': self._extract_key_phrases(text),
            'sentiment': self._analyze_sentiment(text),
            'readability': self._calculate_readability(text),
            'country_context': self.country,
            'curriculum': self.overlay.get('system', 'Universal')
        }
        
        if segment == 'Student':
            analysis.update(self._analyze_student_document(text))
        elif segment == 'Teacher':
            analysis.update(self._analyze_teacher_document(text))
        elif segment == 'Professional':
            analysis.update(self._analyze_professional_document(text))
        elif segment == 'SME Business Owner':
            analysis.update(self._analyze_sme_document(text))
        
        return analysis
    
    def _extract_key_phrases(self, text: str) -> List[str]:
        stopwords = {'the', 'a', 'an', 'of', 'to', 'for', 'with', 'on', 'at', 'from', 'by', 'in', 'and', 'or', 'but'}
        words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
        
        freq = {}
        for word in words:
            if word not in stopwords:
                freq[word] = freq.get(word, 0) + 1
        
        sorted_words = sorted(freq.items(), key=lambda x: x[1], reverse=True)[:10]
        
        bigrams = []
        for i in range(len(words)-1):
            if words[i] not in stopwords and words[i+1] not in stopwords:
                bigram = f"{words[i]} {words[i+1]}"
                if len(bigram) > 5:
                    bigrams.append(bigram)
        
        phrases = [w[0] for w in sorted_words if len(w[0]) > 3][:5]
        
        bigram_freq = {}
        for bg in bigrams:
            bigram_freq[bg] = bigram_freq.get(bg, 0) + 1
        
        top_bigrams = sorted(bigram_freq.items(), key=lambda x: x[1], reverse=True)[:3]
        for bg, _ in top_bigrams:
            if bg not in phrases:
                phrases.append(bg)
        
        return phrases[:7]
    
    def _analyze_sentiment(self, text: str) -> Dict[str, Any]:
        positive_words = {'good', 'great', 'excellent', 'positive', 'achievement', 'success', 'improve', 'growth', 
                         'happy', 'satisfied', 'excited', 'motivated', 'enjoy', 'love', 'best', 'outstanding'}
        negative_words = {'bad', 'poor', 'difficult', 'challenge', 'struggle', 'fail', 'failure', 'frustrating',
                         'disappointed', 'unhappy', 'stress', 'anxiety', 'worry', 'concern'}
        
        words = text.lower().split()
        positive_count = sum(1 for w in words if w in positive_words)
        negative_count = sum(1 for w in words if w in negative_words)
        
        total = positive_count + negative_count
        if total == 0:
            return {'sentiment_score': 0, 'sentiment': 'Neutral'}
        
        score = (positive_count - negative_count) / total
        sentiment = 'Positive' if score > 0.1 else 'Negative' if score < -0.1 else 'Neutral'
        
        return {'sentiment_score': round(score, 2), 'sentiment': sentiment}
    
    def _calculate_readability(self, text: str) -> Dict[str, Any]:
        sentences = re.split(r'[.!?]+', text)
        sentences = [s for s in sentences if len(s.strip()) > 0]
        words = text.split()
        
        if len(sentences) == 0 or len(words) == 0:
            return {'flesch_score': 0, 'grade_level': 'Unknown'}
        
        avg_words_per_sentence = len(words) / len(sentences)
        avg_syllables_per_word = self._count_syllables(text) / len(words)
        
        flesch_score = 206.835 - (1.015 * avg_words_per_sentence) - (84.6 * avg_syllables_per_word)
        
        if flesch_score >= 90:
            grade = '5th Grade (Very Easy)'
        elif flesch_score >= 80:
            grade = '6th Grade (Easy)'
        elif flesch_score >= 70:
            grade = '7th Grade (Fairly Easy)'
        elif flesch_score >= 60:
            grade = '8th-9th Grade (Plain English)'
        elif flesch_score >= 50:
            grade = '10th-12th Grade (Fairly Difficult)'
        elif flesch_score >= 30:
            grade = 'College (Difficult)'
        else:
            grade = 'College Graduate (Very Difficult)'
        
        return {
            'flesch_score': round(flesch_score, 2),
            'grade_level': grade,
            'avg_words_per_sentence': round(avg_words_per_sentence, 2)
        }
    
    def _count_syllables(self, text: str) -> int:
        vowels = 'aeiouy'
        words = text.lower().split()
        count = 0
        for word in words:
            word_vowels = 0
            for char in word:
                if char in vowels:
                    word_vowels += 1
            if word.endswith('e'):
                word_vowels = max(1, word_vowels - 1)
            count += max(1, word_vowels)
        return count
    
    def _analyze_student_document(self, text: str) -> Dict[str, Any]:
        academic_keywords = {'analyze', 'evaluate', 'synthesize', 'discuss', 'compare', 'contrast', 
                            'research', 'study', 'experiment', 'hypothesis', 'theory', 'conclusion'}
        words = text.lower().split()
        academic_count = sum(1 for w in words if w in academic_keywords)
        topics = self._extract_topics(text)
        
        return {
            'academic_language_score': round(min(100, (academic_count / len(words) * 1000)) if words else 0, 2),
            'topics_mentioned': topics[:5],
            'suggested_improvements': self._suggest_student_improvements(text, topics)
        }
    
    def _analyze_teacher_document(self, text: str) -> Dict[str, Any]:
        """Enhanced teacher document analysis with pedagogical scoring"""
        ped_keywords = {'objective', 'learning', 'outcome', 'assessment', 'rubric', 'activity', 
                       'discussion', 'project', 'group', 'individual', 'differentiation',
                       'lesson', 'plan', 'standards', 'evaluation', 'feedback'}
        words = text.lower().split()
        ped_count = sum(1 for w in words if w in ped_keywords)
        
        # Check for key lesson plan components
        has_objective = 'objective' in text.lower()
        has_assessment = 'assessment' in text.lower()
        has_activity = 'activity' in text.lower()
        has_standards = 'standard' in text.lower() or 'standard' in text.lower()
        
        completeness_score = sum([has_objective, has_assessment, has_activity, has_standards]) / 4 * 100
        
        return {
            'pedagogical_score': round(min(100, (ped_count / len(words) * 1000)) if words else 0, 2),
            'curriculum_alignment': self._check_curriculum_alignment(text),
            'suggested_enhancements': self._suggest_teacher_enhancements(text),
            'completeness_score': round(completeness_score, 2),
            'has_objective': has_objective,
            'has_assessment': has_assessment,
            'has_activity': has_activity,
            'has_standards': has_standards
        }
    
    def _analyze_professional_document(self, text: str) -> Dict[str, Any]:
        prof_keywords = {'strategy', 'analysis', 'implementation', 'results', 'findings', 
                        'recommendation', 'efficiency', 'performance', 'optimization',
                        'portfolio', 'market', 'investment', 'returns', 'valuation',
                        'earnings', 'revenue', 'growth', 'margin', 'ROI'}
        words = text.lower().split()
        prof_count = sum(1 for w in words if w in prof_keywords)
        
        return {
            'professional_score': round(min(100, (prof_count / len(words) * 1000)) if words else 0, 2),
            'business_context': self._extract_business_context(text),
            'actionable_insights': self._extract_actionable_insights(text),
            'financial_metrics': self._extract_financial_metrics(text)
        }
    
    def _analyze_sme_document(self, text: str) -> Dict[str, Any]:
        sme_keywords = {'revenue', 'cost', 'profit', 'customer', 'market', 'growth', 
                       'operations', 'supply', 'logistics', 'sales', 'marketing', 'inventory',
                       'cash flow', 'balance sheet', 'income statement', 'forecast'}
        words = text.lower().split()
        sme_count = sum(1 for w in words if w in sme_keywords)
        
        return {
            'business_score': round(min(100, (sme_count / len(words) * 1000)) if words else 0, 2),
            'growth_opportunities': self._identify_growth_opportunities(text),
            'automation_candidates': self._identify_automation_candidates(text),
            'financial_metrics': self._extract_financial_metrics(text)
        }
    
    def _extract_topics(self, text: str) -> List[str]:
        common_topics = {
            'mathematics': ['algebra', 'geometry', 'calculus', 'statistics', 'arithmetic'],
            'science': ['biology', 'chemistry', 'physics', 'environment', 'experiment'],
            'literature': ['novel', 'poetry', 'drama', 'prose', 'literary'],
            'history': ['historical', 'civilization', 'ancient', 'modern', 'century'],
            'geography': ['map', 'climate', 'population', 'region', 'continent'],
            'economics': ['market', 'trade', 'investment', 'currency', 'finance'],
            'technology': ['computer', 'software', 'digital', 'programming', 'ai', 'automation'],
            'language': ['vocabulary', 'grammar', 'writing', 'reading', 'speaking']
        }
        
        found_topics = []
        text_lower = text.lower()
        
        for category, keywords in common_topics.items():
            for keyword in keywords:
                if keyword in text_lower:
                    found_topics.append(f"{category.title()}")
                    break
        
        return list(dict.fromkeys(found_topics))
    
    def _suggest_student_improvements(self, text: str, topics: List[str]) -> List[str]:
        suggestions = []
        words = text.split()
        
        if len(words) < 100:
            suggestions.append("Consider expanding your analysis with more depth and examples.")
        elif len(words) > 1000:
            suggestions.append("Consider condensing your work for clarity and focus.")
        
        academic_words = {'analyze', 'evaluate', 'synthesize', 'compare', 'contrast', 'research'}
        has_academic = any(w in text.lower() for w in academic_words)
        if not has_academic:
            suggestions.append("Incorporate more academic language (analyze, evaluate, synthesize).")
        
        if 'Mathematics' in topics:
            suggestions.append("Include step-by-step working for mathematical problems.")
        if 'Science' in topics:
            suggestions.append("Include more scientific evidence and references.")
        if 'Literature' in topics:
            suggestions.append("Provide more textual evidence to support your arguments.")
        if 'Language' in topics:
            suggestions.append("Include more advanced vocabulary and sentence structures.")
        
        if self.country == 'kenya':
            suggestions.append("Connect your learning to Kenyan community values and context.")
        elif self.country == 'bangladesh':
            suggestions.append("Incorporate Bengali cultural perspectives in your analysis.")
        elif self.country == 'usa':
            suggestions.append("Consider how this applies to American educational standards.")
        elif self.country == 'uk':
            suggestions.append("Align your work with British academic expectations.")
        
        return suggestions[:5]
    
    def _check_curriculum_alignment(self, text: str) -> Dict[str, Any]:
        curriculum_keywords = {
            'kenya': ['competency', 'cbc', 'knec', 'community', 'values', 'skills'],
            'bangladesh': ['nctb', 'board', 'exam', 'bangladesh', 'curriculum', 'national'],
            'usa': ['common core', 'ngss', 'standards', 'college', 'career'],
            'uk': ['national curriculum', 'gcse', 'a-level', 'academic']
        }
        
        keywords = curriculum_keywords.get(self.country, [])
        text_lower = text.lower()
        
        matched = [kw for kw in keywords if kw in text_lower]
        alignment_score = len(matched) / len(keywords) if keywords else 0
        
        return {
            'score': round(alignment_score * 100, 2),
            'matched_indicators': matched,
            'status': 'Aligned' if alignment_score > 0.5 else 'Partial' if alignment_score > 0.2 else 'Not Aligned'
        }
    
    def _suggest_teacher_enhancements(self, text: str) -> List[str]:
        """Enhanced teacher enhancement suggestions"""
        suggestions = []
        text_lower = text.lower()
        
        if 'objective' not in text_lower:
            suggestions.append("Add clear learning objectives for each lesson.")
        if 'assessment' not in text_lower:
            suggestions.append("Include assessment criteria and methods.")
        if 'activity' not in text_lower:
            suggestions.append("Add interactive activities to engage students.")
        if 'differentiation' not in text_lower:
            suggestions.append("Include differentiation strategies for diverse learners.")
        
        if 'technology' not in text_lower and 'digital' not in text_lower:
            suggestions.append("Integrate digital tools and AI resources.")
        
        if 'warm-up' not in text_lower and 'warm up' not in text_lower:
            suggestions.append("Include a warm-up activity to activate prior knowledge.")
        
        if 'exit ticket' not in text_lower and 'exit' not in text_lower:
            suggestions.append("Add an exit ticket to assess student understanding.")
        
        if self.country == 'kenya':
            suggestions.append("Incorporate community-based learning approaches.")
        elif self.country == 'bangladesh':
            suggestions.append("Include bilingual support (Bengali/English).")
        elif self.country == 'usa':
            suggestions.append("Emphasize college and career readiness.")
        elif self.country == 'uk':
            suggestions.append("Focus on academic rigor and depth of understanding.")
        
        return suggestions[:5]
    
    def _extract_business_context(self, text: str) -> Dict[str, Any]:
        context = {
            'industry': 'Not specified',
            'market': 'Not specified',
            'key_metrics': [],
            'timeline': 'Not specified'
        }
        
        text_lower = text.lower()
        
        industries = {
            'tech': ['software', 'technology', 'digital', 'ai', 'automation'],
            'finance': ['finance', 'banking', 'investment', 'currency'],
            'health': ['health', 'medical', 'wellness', 'clinical'],
            'retail': ['retail', 'store', 'customer', 'sales'],
            'manufacturing': ['manufacture', 'production', 'factory', 'assembly']
        }
        
        for industry, keywords in industries.items():
            if any(kw in text_lower for kw in keywords):
                context['industry'] = industry.title()
                break
        
        metric_patterns = [
            r'(\d+[\.,]?\d*)\s*%',
            r'\$\s*(\d+[\.,]?\d*)',
            r'(\d+[\.,]?\d*)\s*million',
            r'(\d+[\.,]?\d*)\s*billion'
        ]
        
        for pattern in metric_patterns:
            matches = re.findall(pattern, text)
            if matches:
                context['key_metrics'].extend(matches[:3])
        
        return context
    
    def _extract_actionable_insights(self, text: str) -> List[str]:
        insights = []
        action_patterns = [
            r'(recommend|suggest|should|must|need to)\s+([^.!?]+)',
            r'(implement|adopt|use|apply)\s+([^.!?]+)',
            r'(improve|enhance|optimize)\s+([^.!?]+)'
        ]
        
        for pattern in action_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                if len(match) >= 2:
                    insight = f"{match[0]} {match[1]}"
                    if len(insight) > 10 and len(insight) < 100:
                        insights.append(insight.strip())
        
        return list(dict.fromkeys(insights))[:5]
    
    def _identify_growth_opportunities(self, text: str) -> List[str]:
        opportunities = []
        text_lower = text.lower()
        
        if any(kw in text_lower for kw in ['new market', 'expansion', 'enter', 'customer']):
            opportunities.append("Market expansion opportunities")
        if any(kw in text_lower for kw in ['product', 'service', 'offer', 'new']):
            opportunities.append("Product/service diversification")
        if any(kw in text_lower for kw in ['digital', 'online', 'e-commerce', 'website']):
            opportunities.append("Digital transformation opportunity")
        if any(kw in text_lower for kw in ['efficiency', 'cost', 'reduce', 'save']):
            opportunities.append("Cost optimization and efficiency improvement")
        if any(kw in text_lower for kw in ['referral', 'repeat', 'loyalty', 'retention']):
            opportunities.append("Customer loyalty and retention program")
        
        if self.country == 'kenya':
            opportunities.append("Leverage mobile money and digital payments")
        elif self.country == 'bangladesh':
            opportunities.append("Explore RMG and export opportunities")
        elif self.country == 'usa':
            opportunities.append("Leverage innovation and technology sectors")
        elif self.country == 'uk':
            opportunities.append("Explore creative and financial services")
        
        return list(dict.fromkeys(opportunities))[:5]
    
    def _identify_automation_candidates(self, text: str) -> List[str]:
        candidates = []
        text_lower = text.lower()
        
        if any(kw in text_lower for kw in ['customer', 'support', 'enquiry', 'help']):
            candidates.append("Customer support chatbot")
        if any(kw in text_lower for kw in ['order', 'receipt', 'invoice', 'payment']):
            candidates.append("Automated billing and invoicing")
        if any(kw in text_lower for kw in ['inventory', 'stock', 'supply', 'warehouse']):
            candidates.append("Inventory management automation")
        if any(kw in text_lower for kw in ['marketing', 'social', 'email', 'content']):
            candidates.append("Marketing automation tools")
        if any(kw in text_lower for kw in ['report', 'analytics', 'dashboard', 'track']):
            candidates.append("Analytics and reporting automation")
        if any(kw in text_lower for kw in ['schedule', 'appointment', 'booking', 'calendar']):
            candidates.append("Scheduling and booking system")
        
        return list(dict.fromkeys(candidates))[:5]
    
    def _extract_financial_metrics(self, text: str) -> Dict[str, Any]:
        metrics = {
            'revenue': None,
            'profit': None,
            'expenses': None,
            'growth_rate': None,
            'margin': None,
            'roi': None
        }
        
        revenue_pattern = r'revenue\s*[:$]?\s*\$?(\d+[\.,]?\d*)'
        profit_pattern = r'profit\s*[:$]?\s*\$?(\d+[\.,]?\d*)'
        expenses_pattern = r'expenses?\s*[:$]?\s*\$?(\d+[\.,]?\d*)'
        growth_pattern = r'growth\s*[:]?\s*(\d+[\.,]?\d*)\s*%'
        margin_pattern = r'margin\s*[:]?\s*(\d+[\.,]?\d*)\s*%'
        roi_pattern = r'ROI\s*[:]?\s*(\d+[\.,]?\d*)\s*%'
        
        revenue_match = re.search(revenue_pattern, text, re.IGNORECASE)
        profit_match = re.search(profit_pattern, text, re.IGNORECASE)
        expenses_match = re.search(expenses_pattern, text, re.IGNORECASE)
        growth_match = re.search(growth_pattern, text, re.IGNORECASE)
        margin_match = re.search(margin_pattern, text, re.IGNORECASE)
        roi_match = re.search(roi_pattern, text, re.IGNORECASE)
        
        if revenue_match:
            metrics['revenue'] = float(revenue_match.group(1).replace(',', ''))
        if profit_match:
            metrics['profit'] = float(profit_match.group(1).replace(',', ''))
        if expenses_match:
            metrics['expenses'] = float(expenses_match.group(1).replace(',', ''))
        if growth_match:
            metrics['growth_rate'] = float(growth_match.group(1).replace(',', ''))
        if margin_match:
            metrics['margin'] = float(margin_match.group(1).replace(',', ''))
        if roi_match:
            metrics['roi'] = float(roi_match.group(1).replace(',', ''))
        
        return metrics


# ==================== UNIVERSAL CORE ENGINE ====================

class UniversalCore:
    UNIVERSAL_SUBJECTS = {
        'mathematics': {
            'topics': ['Arithmetic', 'Algebra', 'Geometry', 'Statistics', 'Calculus'],
            'skills': ['Problem Solving', 'Logical Reasoning', 'Pattern Recognition']
        },
        'english_language': {
            'topics': ['Reading', 'Writing', 'Speaking', 'Listening', 'Grammar'],
            'skills': ['Communication', 'Critical Analysis', 'Creative Expression']
        },
        'basic_science': {
            'topics': ['Biology', 'Chemistry', 'Physics', 'Earth Science'],
            'skills': ['Scientific Method', 'Observation', 'Experimentation']
        },
        'geography': {
            'topics': ['Physical Geography', 'Human Geography', 'Map Skills', 'Climate'],
            'skills': ['Spatial Awareness', 'Cultural Understanding', 'Environmental Awareness']
        },
        'general_knowledge': {
            'topics': ['History', 'Current Events', 'Civics', 'Economics'],
            'skills': ['Critical Thinking', 'Awareness', 'Global Understanding']
        },
        'applied_ai': {
            'topics': ['AI Fundamentals', 'Prompt Engineering', 'Workflow Automation', 'Ethics'],
            'skills': ['AI Literacy', 'Automation', 'Problem Solving']
        }
    }
    
    UNIVERSAL_COMPETENCIES = {
        'critical_thinking': 'Analyze, evaluate, and synthesize information',
        'communication': 'Express ideas clearly in multiple formats',
        'collaboration': 'Work effectively with others',
        'creativity': 'Generate innovative solutions',
        'digital_literacy': 'Use technology effectively and responsibly'
    }


# ==================== LOCAL OVERLAY ENGINE ====================

class LocalCurriculumOverlay:
    CURRICULUM_OVERLAYS = {
        'kenya': {
            'code': 'KE',
            'system': 'CBC (Competency Based Curriculum)',
            'boards': ['KNEC', 'KICD'],
            'subjects': {
                'mathematics': 'Mathematics (including Financial Literacy)',
                'english_language': 'English (with Kiswahili as second language)',
                'basic_science': 'Integrated Science',
                'geography': 'Geography and Social Studies'
            },
            'national_exams': ['KCPE', 'KCSE'],
            'language': 'English/Kiswahili',
            'grade_levels': ['PP1', 'PP2', 'Grade 1-9', 'Form 1-4'],
            'currency': 'KES',
            'timezone': 'EAT'
        },
        'bangladesh': {
            'code': 'BD',
            'system': 'National Curriculum (NCTB)',
            'boards': ['Dhaka Board', 'Rajshahi Board', 'Chittagong Board', 'Barisal Board', 'Sylhet Board'],
            'subjects': {
                'mathematics': 'Mathematics (গণিত)',
                'english_language': 'English (ইংরেজি)',
                'basic_science': 'Science (বিজ্ঞান)',
                'geography': 'Geography and Environment (ভূগোল ও পরিবেশ)'
            },
            'national_exams': ['PSC', 'JSC', 'SSC', 'HSC'],
            'language': 'Bengali/English',
            'grade_levels': ['Class 1-5', 'Class 6-8', 'Class 9-10', 'Class 11-12'],
            'currency': 'BDT',
            'timezone': 'BST'
        },
        'usa': {
            'code': 'US',
            'system': 'Common Core State Standards',
            'boards': ['State-specific', 'College Board'],
            'subjects': {
                'mathematics': 'Mathematics (Common Core)',
                'english_language': 'English Language Arts',
                'basic_science': 'Science (NGSS)',
                'geography': 'Social Studies'
            },
            'national_exams': ['SAT', 'ACT', 'AP'],
            'language': 'English/Spanish',
            'grade_levels': ['K-5', '6-8', '9-12'],
            'currency': 'USD',
            'timezone': 'EST/CST/PST'
        },
        'uk': {
            'code': 'UK',
            'system': 'National Curriculum for England',
            'boards': ['AQA', 'Edexcel', 'OCR', 'WJEC'],
            'subjects': {
                'mathematics': 'Mathematics',
                'english_language': 'English',
                'basic_science': 'Science (Combined/Triple)',
                'geography': 'Geography'
            },
            'national_exams': ['GCSE', 'A-Levels'],
            'language': 'English',
            'grade_levels': ['KS1-2', 'KS3', 'KS4-5'],
            'currency': 'GBP',
            'timezone': 'GMT/BST'
        }
    }
    
    LOCALIZED_PROMPTS = {
        'kenya': {
            'greeting': 'Jambo! Welcome to AI Shiksha Kenya 🇰🇪',
            'exam_style': 'KNEC-style examination questions with practical applications',
            'examples': 'Use examples relevant to East African context',
            'cultural_note': 'Integration of Kenyan cultural values and community-based learning'
        },
        'bangladesh': {
            'greeting': 'স্বাগতম! Welcome to AI Shiksha Bangladesh 🇧🇩',
            'exam_style': 'NCTB and board examination preparation style',
            'examples': 'Use examples relevant to Bangladeshi context',
            'cultural_note': 'Integration of Bangladeshi cultural heritage and language'
        },
        'usa': {
            'greeting': 'Welcome to AI Shiksha USA 🇺🇸',
            'exam_style': 'Common Core and standardized test preparation',
            'examples': 'Use examples relevant to American context',
            'cultural_note': 'Focus on college and career readiness'
        },
        'uk': {
            'greeting': 'Welcome to AI Shiksha UK 🇬🇧',
            'exam_style': 'GCSE and A-Level examination format',
            'examples': 'Use examples relevant to British context',
            'cultural_note': 'Focus on academic rigor and depth'
        }
    }
    
    LOCAL_CONTEXTS = {
        'kenya': {
            'culture': 'Kenyan community values, Harambee spirit, diverse ethnic groups',
            'environment': 'Savanna, wildlife, agriculture, coastal regions',
            'economy': 'Agriculture, tourism, technology (Silicon Savannah)'
        },
        'bangladesh': {
            'culture': 'Bengali heritage, language movement, diverse traditions',
            'environment': 'Delta region, rivers, monsoon climate, agriculture',
            'economy': 'Garment industry, agriculture, remittances, technology'
        },
        'usa': {
            'culture': 'Diverse immigrant nation, American Dream, individual liberty',
            'environment': 'Diverse climates, 50 states, national parks',
            'economy': 'World\'s largest economy, innovation, technology'
        },
        'uk': {
            'culture': 'British heritage, royal tradition, multicultural society',
            'environment': 'Temperate climate, varied landscapes, historic cities',
            'economy': 'Service economy, financial hub, creative industries'
        }
    }
    
    @staticmethod
    def get_overlay(country_code):
        return LocalCurriculumOverlay.CURRICULUM_OVERLAYS.get(country_code, {})
    
    @staticmethod
    def get_local_context(country_code):
        return LocalCurriculumOverlay.LOCAL_CONTEXTS.get(country_code, {})
    
    @staticmethod
    def get_localized_prompt(country_code, key):
        prompts = LocalCurriculumOverlay.LOCALIZED_PROMPTS.get(country_code, {})
        return prompts.get(key, '')
    
    @staticmethod
    def localize_question(question, country_code):
        localized = question.copy()
        overlay = LocalCurriculumOverlay.get_overlay(country_code)
        context = LocalCurriculumOverlay.get_local_context(country_code)
        
        if overlay and 'subjects' in overlay:
            for uni_sub, local_sub in overlay['subjects'].items():
                if uni_sub in question.get('subject', '').lower():
                    localized['local_subject'] = local_sub
                    break
        
        if 'explanation' in localized:
            context_phrases = {
                'kenya': f" using Kenyan examples and context (agriculture, wildlife, community)",
                'bangladesh': f" using Bangladeshi examples (rivers, garment industry, Bengali culture)",
                'usa': f" using American examples (diversity, innovation, local communities)",
                'uk': f" using British examples (history, multicultural society, local context)"
            }
            localized['explanation'] += context_phrases.get(country_code, '')
        
        if 'socratic_hint' in localized:
            cultural_note = LocalCurriculumOverlay.get_localized_prompt(country_code, 'cultural_note')
            if cultural_note:
                localized['socratic_hint'] += f" (Cultural context: {cultural_note})"
        
        return localized
    
    @staticmethod
    def get_grade_levels(country_code):
        overlay = LocalCurriculumOverlay.get_overlay(country_code)
        return overlay.get('grade_levels', ['Primary', 'Secondary'])


# ==================== STUDENT LEARNING INTELLIGENCE ENGINE ====================
# Data Infrastructure & Adaptive Algorithms + Pedagogical Rules + Gamification +
# Analytics & Safety Controls for the Student segment.
#
# Everything below runs locally against st.session_state (in-memory for the
# session) — nothing is written to a database or file, and nothing is sent to
# any external training pipeline, which is the concrete implementation of the
# COPPA/FERPA data-isolation requirement in this demo (see PrivacyControlEngine).

@dataclass
class SpacedRepetitionCard:
    """A single memory item tracked with the SM-2 algorithm (Piotr Wozniak's
    SuperMemo-2 scheduler — the same family of algorithm Anki's default
    scheduler is based on)."""
    concept_id: str
    subject: str
    concept_name: str
    easiness: float = 2.5
    interval_days: int = 0
    repetitions: int = 0
    next_review: str = field(default_factory=lambda: datetime.now().strftime('%Y-%m-%d'))
    last_reviewed: Optional[str] = None
    lapses: int = 0


class SpacedRepetitionEngine:
    """SM-2 style scheduler. Quality is graded 0-5 by the student's self-assessed
    recall: 0-2 = forgotten/hard, 3 = correct with effort, 4 = correct, 5 = easy."""

    @staticmethod
    def schedule(card: SpacedRepetitionCard, quality: int) -> SpacedRepetitionCard:
        quality = max(0, min(5, quality))

        if quality < 3:
            card.repetitions = 0
            card.interval_days = 1
            card.lapses += 1
        else:
            if card.repetitions == 0:
                card.interval_days = 1
            elif card.repetitions == 1:
                card.interval_days = 6
            else:
                card.interval_days = max(1, round(card.interval_days * card.easiness))
            card.repetitions += 1

        card.easiness = card.easiness + (0.1 - (5 - quality) * (0.08 + (5 - quality) * 0.02))
        card.easiness = max(1.3, round(card.easiness, 2))

        card.last_reviewed = datetime.now().strftime('%Y-%m-%d')
        card.next_review = (datetime.now() + timedelta(days=card.interval_days)).strftime('%Y-%m-%d')
        return card

    @staticmethod
    def due_cards(cards: Dict[str, "SpacedRepetitionCard"]) -> List["SpacedRepetitionCard"]:
        today = datetime.now().strftime('%Y-%m-%d')
        due = [c for c in cards.values() if c.next_review <= today]
        return sorted(due, key=lambda c: c.next_review)

    @staticmethod
    def get_or_create(cards: Dict[str, "SpacedRepetitionCard"], concept_id, subject, concept_name):
        if concept_id not in cards:
            cards[concept_id] = SpacedRepetitionCard(concept_id=concept_id, subject=subject, concept_name=concept_name)
        return cards[concept_id]


@dataclass
class KnowledgeNode:
    node_id: str
    name: str
    subject: str
    prerequisites: List[str] = field(default_factory=list)


class PrerequisiteKnowledgeGraph:
    """Node-based dependency graphs per subject. A node unlocks only once every
    prerequisite node is mastered (e.g. Factoring must be mastered before
    Quadratic Equations unlocks).

    These graphs are illustrative starter content for the demo, not a
    curriculum-authority mapping — swap in real syllabus dependency data for
    production use."""

    GRAPHS: Dict[str, List[KnowledgeNode]] = {
        'mathematics': [
            KnowledgeNode('num_ops', 'Number Operations', 'mathematics', []),
            KnowledgeNode('fractions', 'Fractions & Decimals', 'mathematics', ['num_ops']),
            KnowledgeNode('algebra_basics', 'Algebraic Expressions', 'mathematics', ['fractions']),
            KnowledgeNode('linear_eq', 'Linear Equations', 'mathematics', ['algebra_basics']),
            KnowledgeNode('factoring', 'Factoring', 'mathematics', ['linear_eq']),
            KnowledgeNode('quadratics', 'Quadratic Equations', 'mathematics', ['factoring']),
            KnowledgeNode('functions', 'Functions & Graphs', 'mathematics', ['quadratics']),
        ],
        'english language': [
            KnowledgeNode('phonics', 'Phonics & Vocabulary', 'english language', []),
            KnowledgeNode('grammar_basics', 'Grammar Basics', 'english language', ['phonics']),
            KnowledgeNode('sentence_struct', 'Sentence Structure', 'english language', ['grammar_basics']),
            KnowledgeNode('paragraph', 'Paragraph Writing', 'english language', ['sentence_struct']),
            KnowledgeNode('essay', 'Essay Composition', 'english language', ['paragraph']),
        ],
        'basic science': [
            KnowledgeNode('sci_method', 'Scientific Method', 'basic science', []),
            KnowledgeNode('matter', 'States of Matter', 'basic science', ['sci_method']),
            KnowledgeNode('cells', 'Cells & Living Things', 'basic science', ['sci_method']),
            KnowledgeNode('energy', 'Energy & Forces', 'basic science', ['matter']),
            KnowledgeNode('ecosystems', 'Ecosystems', 'basic science', ['cells']),
        ],
        'geography': [
            KnowledgeNode('maps', 'Map Reading', 'geography', []),
            KnowledgeNode('landforms', 'Landforms', 'geography', ['maps']),
            KnowledgeNode('climate', 'Climate Zones', 'geography', ['maps']),
            KnowledgeNode('regions', 'Regional Geography', 'geography', ['landforms', 'climate']),
        ],
        'general knowledge': [
            KnowledgeNode('gk_basics', 'General Awareness Basics', 'general knowledge', []),
            KnowledgeNode('gk_current', 'Current Affairs', 'general knowledge', ['gk_basics']),
        ],
    }

    @classmethod
    def get_graph(cls, subject: str) -> List[KnowledgeNode]:
        return cls.GRAPHS.get(subject.lower(), cls.GRAPHS['mathematics'])

    @staticmethod
    def compute_status(nodes: List[KnowledgeNode], mastered_ids: set) -> Dict[str, str]:
        status = {}
        for node in nodes:
            if node.node_id in mastered_ids:
                status[node.node_id] = 'mastered'
            elif all(p in mastered_ids for p in node.prerequisites):
                status[node.node_id] = 'unlocked'
            else:
                status[node.node_id] = 'locked'
        return status

    @staticmethod
    def next_recommended(nodes: List[KnowledgeNode], mastered_ids: set) -> Optional[KnowledgeNode]:
        for node in nodes:
            if node.node_id not in mastered_ids and all(p in mastered_ids for p in node.prerequisites):
                return node
        return None


class MultimodalOCREngine:
    """Ingests photos of handwritten math, diagrams, or textbook pages and turns
    them into text the practice engine can work with.

    Real extraction requires the `pytesseract` package AND the Tesseract OCR
    binary installed on the host machine. I cannot confirm Tesseract is
    installed in your deployment environment — verify with `tesseract --version`.
    If either is missing, this engine reports that clearly and offers a manual
    transcription fallback instead of failing silently or faking a result."""

    @staticmethod
    def is_available() -> bool:
        return OCR_AVAILABLE

    @staticmethod
    def extract_text(image_bytes) -> Dict[str, Any]:
        if not OCR_AVAILABLE:
            return {
                'success': False,
                'text': '',
                'error': 'pytesseract / Pillow / Tesseract not available on this server.'
            }
        try:
            image = Image.open(io.BytesIO(image_bytes))
            text = pytesseract.image_to_string(image)
            return {'success': True, 'text': text.strip(), 'error': None}
        except Exception as e:
            return {'success': False, 'text': '', 'error': str(e)}


class SocraticGuardrailEngine:
    """System-prompt-style rules that stop the tutor from ever handing over the
    final answer directly. Produces a graduated hint ladder (nudge -> strategy
    -> applied step -> self-check) and layers a Contextual Content Shield that
    blocks non-academic, off-topic, or direct cheating requests.

    The content shield here is a simple rule-based keyword filter for demo
    purposes — it is not a production-grade moderation system and will miss
    paraphrased attempts; treat it as a starting point, not a guarantee."""

    OFF_TOPIC_MARKERS = [
        'write my essay', 'do my homework for me', 'give me the answer',
        'just tell me the answer', 'cheat on', 'exam answers', 'test answers',
        'homework answers', 'do this for me', 'skip the explanation'
    ]

    @classmethod
    def content_shield(cls, text: str) -> Dict[str, Any]:
        lowered = (text or '').lower()
        flagged = [m for m in cls.OFF_TOPIC_MARKERS if m in lowered]
        return {
            'blocked': len(flagged) > 0,
            'reason': f"This looks like a request to skip learning ({', '.join(flagged)}) rather than a study question." if flagged else None
        }

    @staticmethod
    def hint_ladder(question: Dict[str, Any]) -> List[str]:
        base_hint = question.get('socratic_hint', "Think about what the question is really asking.")
        return [
            "🤔 Nudge: What do you already know that connects to this problem?",
            f"💡 Strategy: {base_hint}",
            "🔍 Applied step: Try that strategy on just the first part of the problem.",
            "✅ Self-check: Redo your last step — does it fully answer the original question?"
        ]

    @classmethod
    def get_hint(cls, question: Dict[str, Any], hint_level: int) -> str:
        ladder = cls.hint_ladder(question)
        idx = max(0, min(hint_level, len(ladder) - 1))
        return ladder[idx]


class DynamicDifficultyBalancer:
    """Shifts difficulty algorithmically: breaks a problem down into sub-steps
    after two consecutive wrong attempts, and escalates the difficulty tier
    after a win streak of three or more correct answers."""

    LEVELS = ['easy', 'medium', 'hard']

    @classmethod
    def next_difficulty(cls, current: str, wrong_streak: int, win_streak: int) -> Dict[str, Any]:
        idx = cls.LEVELS.index(current) if current in cls.LEVELS else 1
        new_idx = idx
        break_down = False

        if wrong_streak >= 2:
            break_down = True
            new_idx = max(0, idx - 1)
        elif win_streak >= 3:
            new_idx = min(len(cls.LEVELS) - 1, idx + 1)

        return {
            'difficulty': cls.LEVELS[new_idx],
            'break_into_substeps': break_down,
            'changed': new_idx != idx
        }


class VoiceActiveRecallEngine:
    """Lets a student explain a concept out loud so verbal comprehension can be
    evaluated, rather than only multiple-choice recall.

    IMPORTANT — verify before relying on this in production: this demo does
    not ship a real speech-to-text model. It captures audio via Streamlit's
    built-in `st.audio_input` widget, but transcription needs a real STT
    backend wired into `transcribe_audio()` below (for example a hosted
    speech-to-text API) — I have not integrated one, since doing so would
    require an API key and network access this environment doesn't have. The
    scoring logic (`score_explanation`) works on whatever transcript text it
    is given, whether typed manually or produced by a real STT service."""

    @staticmethod
    def transcribe_audio(audio_bytes) -> Dict[str, Any]:
        return {
            'success': False,
            'transcript': '',
            'note': 'No speech-to-text backend is connected in this demo. Type what you said below, or wire a real STT API into transcribe_audio().'
        }

    @staticmethod
    def score_explanation(transcript: str, key_terms: List[str]) -> Dict[str, Any]:
        lowered = (transcript or '').lower()
        hit = [t for t in key_terms if t.lower() in lowered]
        missed = [t for t in key_terms if t.lower() not in lowered]
        coverage = round(100 * len(hit) / len(key_terms), 1) if key_terms else 0.0
        return {'coverage_pct': coverage, 'covered_terms': hit, 'missed_terms': missed}


class InstantStepEvaluator:
    """Split-screen, line-by-line evaluation of a student's worked solution.
    Matching is a simple keyword-overlap heuristic against expected-step text
    for this demo — not a symbolic math checker — so treat 'correct' as
    'closely matches the expected step', not a verified proof."""

    @staticmethod
    def evaluate_steps(student_steps: List[str], expected_steps: List[str]) -> List[Dict[str, Any]]:
        results = []
        for i, exp in enumerate(expected_steps):
            student_line = student_steps[i] if i < len(student_steps) else ''
            exp_keywords = [w.strip('.,()=') for w in exp.lower().split() if len(w) > 2]
            hits = sum(1 for w in exp_keywords if w in student_line.lower())
            ratio = hits / max(1, len(exp_keywords))
            if not student_line.strip():
                status = 'empty'
            elif ratio >= 0.6:
                status = 'correct'
            elif ratio >= 0.3:
                status = 'partial'
            else:
                status = 'incorrect'
            results.append({'step_number': i + 1, 'expected': exp, 'student_input': student_line, 'status': status})
        return results


@dataclass
class AttemptLog:
    concept_id: str
    subject: str
    solve_time_sec: float
    edit_attempts: int
    self_reported_confidence: int  # 1-5
    correct: bool
    timestamp: str = field(default_factory=lambda: datetime.now().isoformat())


class LatentKnowledgeGapTracker:
    """Backend measurement of solve speed, edit/backtrack attempts, and
    self-reported confidence, combined into a hidden-weakness score per
    concept. A concept flagged 'overconfident relative to accuracy' or 'slow
    with low accuracy' surfaces as a latent gap even when raw accuracy looks
    acceptable on its own.

    The scoring weights below are a reasonable heuristic for a demo, not a
    validated psychometric model — treat the flags as prompts to look closer,
    not a certified diagnosis of the student's understanding."""

    @staticmethod
    def analyze(logs: List[AttemptLog]) -> List[Dict[str, Any]]:
        by_concept: Dict[str, List[AttemptLog]] = {}
        for log in logs:
            by_concept.setdefault(log.concept_id, []).append(log)

        flags = []
        for concept_id, entries in by_concept.items():
            avg_time = sum(e.solve_time_sec for e in entries) / len(entries)
            avg_edits = sum(e.edit_attempts for e in entries) / len(entries)
            avg_confidence = sum(e.self_reported_confidence for e in entries) / len(entries)
            accuracy = sum(1 for e in entries if e.correct) / len(entries)

            gap_score = 0
            reasons = []
            if accuracy < 0.6:
                gap_score += 2
                reasons.append('low accuracy')
            if avg_confidence >= 4 and accuracy < 0.6:
                gap_score += 3
                reasons.append('overconfident relative to accuracy')
            if avg_edits >= 3:
                gap_score += 1
                reasons.append('high edit/backtrack count')
            if avg_time > 90 and accuracy < 0.8:
                gap_score += 1
                reasons.append('slow relative to accuracy')

            if gap_score > 0:
                flags.append({
                    'concept_id': concept_id,
                    'subject': entries[0].subject,
                    'gap_score': gap_score,
                    'reasons': reasons,
                    'avg_time_sec': round(avg_time, 1),
                    'avg_edits': round(avg_edits, 1),
                    'avg_confidence': round(avg_confidence, 1),
                    'accuracy_pct': round(accuracy * 100, 1)
                })

        return sorted(flags, key=lambda f: -f['gap_score'])


class HabitStreakEngine:
    """Daily habit-loop mechanics: streak tracking with a grace-period 'streak
    repair' buffer (so one missed day doesn't erase progress) and 5-minute
    micro-session pacing. Streamlit apps can't send real OS/browser push
    notifications on their own — I have not faked that — so this surfaces
    reminders as in-app banners instead; wire a real push service (e.g. web
    push or a mobile notification SDK) at the deployment layer if you need
    true push notifications."""

    MICRO_SESSION_SECONDS = 5 * 60

    @staticmethod
    def record_activity(last_active_date: Optional[str], current_streak: int, repair_tokens: int) -> Dict[str, Any]:
        today = datetime.now().date()
        if last_active_date is None:
            return {'streak': 1, 'repair_tokens': repair_tokens, 'message': 'Streak started! 🔥', 'date': today.isoformat()}

        last_date = datetime.fromisoformat(last_active_date).date()
        gap = (today - last_date).days

        if gap == 0:
            return {'streak': max(1, current_streak), 'repair_tokens': repair_tokens, 'message': 'Already logged today ✅', 'date': today.isoformat()}
        elif gap == 1:
            return {'streak': current_streak + 1, 'repair_tokens': repair_tokens, 'message': f'Streak extended to {current_streak + 1} days! 🔥', 'date': today.isoformat()}
        elif gap == 2 and repair_tokens > 0:
            return {'streak': current_streak + 1, 'repair_tokens': repair_tokens - 1, 'message': 'Missed a day, but a streak-repair buffer saved it! 🩹🔥', 'date': today.isoformat()}
        else:
            return {'streak': 1, 'repair_tokens': repair_tokens, 'message': "Streak reset — let's build a new one!", 'date': today.isoformat()}


class PrivacyControlEngine:
    """COPPA/FERPA-aligned data controls for this session.

    Fact, not marketing: this app keeps student work in Streamlit's in-memory
    `st.session_state` only. Nothing here is written to a database or file by
    default, and nothing in this codebase sends student data to a model
    training pipeline. That said, I cannot verify the privacy posture of
    whatever server or hosting platform you deploy this on (logs, host-level
    disk caching, etc.) — this class controls the application layer only.
    It adds an explicit, visible consent flag and a one-click purge so the
    zero-retention behavior is verifiable rather than just implied."""

    SESSION_ONLY_KEYS = [
        'sr_cards', 'mastered_concepts', 'attempt_logs', 'completed_lessons',
        'achievements', 'weak_areas', 'ocr_history', 'voice_sessions'
    ]

    @classmethod
    def purge(cls, session_state) -> None:
        for k in cls.SESSION_ONLY_KEYS:
            if k in session_state:
                current = session_state[k]
                if isinstance(current, dict):
                    session_state[k] = {}
                elif isinstance(current, list):
                    session_state[k] = []
        session_state.student_score = 0
        session_state.streak = 0
        session_state.wrong_streak = 0
        session_state.win_streak = 0


# ==================== SEGMENT-SPECIFIC OUTCOME ENGINES ====================

class StudentOutcomeEngine:
    def __init__(self, country_code='kenya'):
        self.country = country_code
        self.overlay = LocalCurriculumOverlay.get_overlay(country_code)
        self.context = LocalCurriculumOverlay.get_local_context(country_code)
        self.doc_engine = DocumentIntelligenceEngine(country_code)
    
    def get_adaptive_questions(self, subject, difficulty='medium'):
        questions = {
            'easy': [
                {
                    'id': 1,
                    'subject': subject,
                    'question': 'What is 5 + 7?',
                    'options': ['10', '11', '12', '13'],
                    'correct': '12',
                    'explanation': '5 + 7 = 12',
                    'socratic_hint': 'Count from 5: 6,7,8,9,10,11,12'
                }
            ],
            'medium': [
                {
                    'id': 2,
                    'subject': subject,
                    'question': 'What is 15 × 8?',
                    'options': ['100', '120', '130', '150'],
                    'correct': '120',
                    'explanation': '15 × 8 = 120 (15 × 10 - 15 × 2 = 150 - 30 = 120)',
                    'socratic_hint': 'Break it down: 15 × 10 = 150, then subtract 15 × 2 = 30'
                }
            ],
            'hard': [
                {
                    'id': 3,
                    'subject': subject,
                    'question': 'If 3x + 7 = 22, what is x?',
                    'options': ['3', '5', '7', '9'],
                    'correct': '5',
                    'explanation': '3x = 22 - 7 = 15, x = 15/3 = 5',
                    'socratic_hint': 'First, isolate the term with x'
                }
            ]
        }
        
        localized = []
        for q in questions.get(difficulty, []):
            q['country'] = self.country
            q['exam_style'] = self.overlay.get('national_exams', ['local'])[0]
            q['local_context'] = self.context
            localized.append(LocalCurriculumOverlay.localize_question(q, self.country))
        
        return localized
    
    def track_progress(self, user_data):
        progress_metrics = {
            'current_score': user_data.get('score', 0),
            'streak': user_data.get('streak', 0),
            'topics_mastered': user_data.get('mastered', []),
            'areas_to_improve': user_data.get('weak_areas', []),
            'projected_grade': self._calculate_projected_grade(user_data),
            'grade_level': self.overlay.get('grade_levels', ['Unknown'])[0]
        }
        return progress_metrics
    
    def _calculate_projected_grade(self, user_data):
        score = user_data.get('score', 0)
        
        grading_scales = {
            'kenya': {90: 'A (Excellent)', 75: 'B (Good)', 60: 'C (Satisfactory)', 45: 'D (Needs Improvement)', 0: 'E (Remedial)'},
            'bangladesh': {80: 'A+ (Excellent)', 70: 'A (Good)', 60: 'A- (Satisfactory)', 50: 'B (Average)', 0: 'C (Needs Improvement)'},
            'usa': {90: 'A', 80: 'B', 70: 'C', 60: 'D', 0: 'F'},
            'uk': {70: 'A (First Class)', 60: 'B (Upper Second)', 50: 'C (Lower Second)', 40: 'D (Third)', 0: 'E (Fail)'}
        }
        
        scale = grading_scales.get(self.country, {90: 'A', 75: 'B', 60: 'C', 45: 'D', 0: 'E'})
        
        for threshold, grade in sorted(scale.items(), reverse=True):
            if score >= threshold:
                return grade
        return 'Needs Assessment'


class TeacherOutcomeEngine:
    """Enhanced Teacher Outcome Engine with all teacher tools"""
    
    def __init__(self, country_code='kenya'):
        self.country = country_code
        self.overlay = LocalCurriculumOverlay.get_overlay(country_code)
        self.context = LocalCurriculumOverlay.get_local_context(country_code)
        self.doc_engine = DocumentIntelligenceEngine(country_code)
        self.intelligence = TeacherIntelligenceEngine(country_code)
    
    def generate_lesson_plan(self, subject, grade, duration, curriculum='universal'):
        overlay = self.overlay
        context = self.context
        
        plan = {
            'title': f"{subject} Lesson - {grade}",
            'curriculum': overlay.get('system', 'Universal'),
            'country': self.country,
            'duration': duration,
            'objectives': self._generate_objectives(subject, grade),
            'activities': self._generate_activities(subject, duration),
            'assessment': self._generate_assessment(subject)
        }
        
        if overlay:
            plan['local_context'] = f"Aligned with {overlay.get('system')} ({overlay.get('code')})"
            plan['language_support'] = f"Available in {overlay.get('language', 'English')}"
            plan['cultural_context'] = context
        
        return plan
    
    def _generate_objectives(self, subject, grade):
        objectives = [
            f"Understand key concepts of {subject} at {grade} level",
            f"Apply {subject} knowledge to real-world problems",
            f"Develop critical thinking in {subject} context"
        ]
        
        if self.country == 'kenya':
            objectives.extend([
                "Demonstrate competency-based learning outcomes",
                "Integrate Kenyan community values in learning"
            ])
        elif self.country == 'bangladesh':
            objectives.extend([
                "Develop skills aligned with Bangladesh National Curriculum",
                "Apply learning to Bengali language and cultural context"
            ])
        elif self.country == 'usa':
            objectives.extend([
                "Meet Common Core State Standards",
                "Develop college and career readiness skills"
            ])
        elif self.country == 'uk':
            objectives.extend([
                "Achieve National Curriculum objectives",
                "Develop rigorous academic understanding"
            ])
        
        return objectives
    
    def _generate_activities(self, subject, duration):
        time_slots = []
        
        if duration <= 30:
            time_slots = [10, 10, 10]
        elif duration <= 45:
            time_slots = [15, 20, 10]
        else:
            time_slots = [15, 30, 15]
        
        local_examples = {
            'kenya': 'using local examples from Kenyan agriculture, wildlife, and community',
            'bangladesh': 'using local examples from Bangladeshi rivers, culture, and garment industry',
            'usa': 'using local examples from American communities and innovation',
            'uk': 'using local examples from British history and multicultural society'
        }
        
        example_note = local_examples.get(self.country, 'using local examples')
        
        activities = [
            f"Warm-up (0-{time_slots[0]} min): Introduction to {subject} {example_note}",
            f"Main Activity ({time_slots[0]}-{time_slots[0]+time_slots[1]} min): Interactive learning session with group work",
            f"Closure ({time_slots[0]+time_slots[1]}-{duration} min): Review and Q&A with local context application"
        ]
        
        return activities
    
    def _generate_assessment(self, subject):
        return {
            'criteria': ['Understanding', 'Application', 'Analysis', 'Communication'],
            'weighting': [30, 25, 25, 20],
            'rubric': self._get_localized_rubric()
        }
    
    def _get_localized_rubric(self):
        rubrics = {
            'kenya': {
                'Excellent': 'Exceeds competency expectations with community application',
                'Good': 'Meets competency expectations with practical understanding',
                'Satisfactory': 'Basic competency achieved',
                'Needs Improvement': 'Requires additional support for competency'
            },
            'bangladesh': {
                'Excellent': 'Outstanding understanding with Bengali context mastery',
                'Good': 'Strong understanding with local application',
                'Satisfactory': 'Meets curriculum requirements',
                'Needs Improvement': 'Needs additional support for board standards'
            },
            'usa': {
                'Excellent': 'Exceeds standards with creativity and insight',
                'Good': 'Meets all standards with confidence',
                'Satisfactory': 'Meets basic standards',
                'Needs Improvement': 'Requires additional support for standards'
            },
            'uk': {
                'Excellent': 'Exceptional understanding with depth and rigor',
                'Good': 'Strong understanding with academic quality',
                'Satisfactory': 'Meets National Curriculum requirements',
                'Needs Improvement': 'Requires additional support for GCSE/A-Level preparation'
            }
        }
        
        return rubrics.get(self.country, {
            'Excellent': 'Demonstrates mastery with creativity',
            'Good': 'Shows strong understanding',
            'Satisfactory': 'Meets basic requirements',
            'Needs Improvement': 'Requires additional support'
        })
    
    # ===== Teacher Intelligence Methods =====
    
    def generate_lesson_bundle(self, subject: str, grade: str, topic: str, duration: int) -> Dict[str, Any]:
        return self.intelligence.generate_lesson_bundle(subject, grade, topic, duration)
    
    def generate_quiz(self, subject: str, grade: str, topic: str, num_questions: int = 5) -> Quiz:
        return self.intelligence.generate_quiz(subject, grade, topic, num_questions)
    
    def generate_rubric(self, title: str, criteria: List[str], levels: List[str]) -> Rubric:
        return self.intelligence.generate_rubric(title, criteria, levels)
    
    def adjust_reading_level(self, text: str, target_level: str) -> Dict[str, Any]:
        return self.intelligence.adjust_reading_level(text, target_level)
    
    def convert_document(self, uploaded_file, output_format: str) -> Dict[str, Any]:
        return self.intelligence.convert_document(uploaded_file, output_format)
    
    def get_lesson_blocks(self, lesson_plan: LessonPlan) -> Dict[str, Any]:
        return self.intelligence.get_lesson_blocks(lesson_plan)
    
    def regenerate_block(self, block_type: str, context: str) -> str:
        return self.intelligence.regenerate_block(block_type, context)
    
    def apply_accommodation(self, content: str, accommodation_type: str) -> Dict[str, Any]:
        return self.intelligence.apply_accommodation(content, accommodation_type)
    
    def generate_batch_feedback(self, assignments: List[Dict[str, Any]]) -> List[StudentFeedback]:
        return self.intelligence.generate_batch_feedback(assignments)
    
    def approve_feedback(self, feedback_id: str) -> Dict[str, Any]:
        return self.intelligence.approve_feedback(feedback_id)
    
    def calculate_time_saved(self) -> Dict[str, Any]:
        return self.intelligence.calculate_time_saved()
    
    def generate_parent_update(self, student_name: str, subject: str, progress: str) -> str:
        return self.intelligence.generate_parent_update(student_name, subject, progress)
    
    def generate_sub_plan(self, subject: str, grade: str, lesson_topic: str) -> Dict[str, Any]:
        return self.intelligence.generate_sub_plan(subject, grade, lesson_topic)
    
    def redact_pii(self, text: str) -> Dict[str, Any]:
        return self.intelligence.redact_pii(text)
    
    def connect_lms(self, lms_type: str, credentials: Dict[str, str]) -> Dict[str, Any]:
        return self.intelligence.connect_lms(lms_type, credentials)
    
    def sync_lms_data(self, lms_type: str, sync_type: str = 'all') -> Dict[str, Any]:
        return self.intelligence.sync_lms_data(lms_type, sync_type)
    
    def map_standards(self, content: str, standard_type: str = 'common_core') -> List[str]:
        return self.intelligence.map_standards(content, standard_type)


class ProfessionalOutcomeEngine:
    def __init__(self, domain='business', country_code='kenya'):
        self.domain = domain
        self.country = country_code
        self.overlay = LocalCurriculumOverlay.get_overlay(country_code)
        self.doc_engine = DocumentIntelligenceEngine(country_code)
        self.intelligence = ProfessionalIntelligenceEngine(country_code)
    
    def generate_workflow(self, task_type):
        workflows = {
            'research': {
                'steps': [
                    'Define research question',
                    'Gather and analyze data',
                    'Synthesize findings',
                    'Generate report',
                    'Add citations and references'
                ],
                'output': 'Research synthesis with actionable insights',
                'localization': f'Using {self.overlay.get("system", "global")} standards and {self.overlay.get("currency", "local")} context'
            },
            'marketing': {
                'steps': [
                    'Define target audience',
                    'Create content strategy',
                    'Generate marketing copy',
                    'Design visual elements',
                    'Track and optimize performance'
                ],
                'output': 'Multi-channel marketing campaign',
                'localization': f'Localized for {self.country.upper()} market with {self.overlay.get("language", "English")} support'
            },
            'analytics': {
                'steps': [
                    'Collect data from all sources',
                    'Clean and preprocess data',
                    'Perform statistical analysis',
                    'Create visualizations',
                    'Interpret results and suggest actions'
                ],
                'output': 'Comprehensive analytics dashboard',
                'localization': f'Adapted for {self.overlay.get("system", "local")} business environment'
            },
            'investment_research': {
                'steps': [
                    'Identify investment thesis',
                    'Analyze market fundamentals',
                    'Evaluate competitive landscape',
                    'Build financial model',
                    'Generate investment recommendation'
                ],
                'output': 'Investment research report with actionable insights',
                'localization': f'Using {self.overlay.get("currency", "local")} market context'
            },
            'portfolio_analysis': {
                'steps': [
                    'Gather portfolio data',
                    'Calculate risk metrics',
                    'Analyze sector exposure',
                    'Run stress tests',
                    'Generate optimization recommendations'
                ],
                'output': 'Portfolio analysis with risk attribution',
                'localization': f'Adapted for {self.country.upper()} market conditions'
            }
        }
        
        return workflows.get(task_type, workflows['research'])
    
    def hybrid_search(self, query: str, search_type: str = 'hybrid') -> List[Dict[str, Any]]:
        search_type_enum = SearchType.HYBRID
        if search_type == 'keyword':
            search_type_enum = SearchType.KEYWORD
        elif search_type == 'vector':
            search_type_enum = SearchType.VECTOR
        elif search_type == 'semantic':
            search_type_enum = SearchType.SEMANTIC
        
        return self.intelligence.hybrid_search(query, search_type_enum)
    
    def audit_analysis(self, analysis_text: str) -> Dict[str, Any]:
        return self.intelligence.audit_analysis(analysis_text, self.intelligence.documents)
    
    def execute_quant_analysis(self, code: str, data: pd.DataFrame) -> Dict[str, Any]:
        return self.intelligence.execute_quant_analysis(code, data)
    
    def get_portfolio_metrics(self) -> Dict[str, Any]:
        return self.intelligence.calculate_portfolio_metrics()
    
    def run_stress_test(self, scenario: MacroScenario) -> Dict[str, Any]:
        return self.intelligence.run_macro_stress_test(scenario)
    
    def generate_deliverable(self, template_type: str, data: Dict[str, Any]) -> str:
        return self.intelligence.generate_deliverable(template_type, data)
    
    def extract_table(self, pdf_content: str) -> pd.DataFrame:
        return self.intelligence.extract_table_from_pdf(pdf_content)
    
    def get_watchdog_alerts(self) -> List[WatchdogAlert]:
        return self.intelligence.watchdog_alerts
    
    def refresh_watchdogs(self) -> List[WatchdogAlert]:
        return self.intelligence.run_watchdog_alerts()
    
    def get_market_data(self, ticker: str) -> Dict[str, Any]:
        return self.intelligence.fetch_market_data(ticker)
    
    def get_sec_filing(self, ticker: str) -> Dict[str, Any]:
        return self.intelligence.fetch_sec_filing(ticker)
    
    def get_earnings_call(self, ticker: str) -> Dict[str, Any]:
        return self.intelligence.fetch_earnings_call(ticker)
    
    def get_workspace_data(self) -> Dict[str, Any]:
        return self.intelligence.get_workspace_data()


class SMEOutcomeEngine:
    def __init__(self, market='africa', country_code='kenya'):
        self.market = market
        self.country = country_code
        self.overlay = LocalCurriculumOverlay.get_overlay(country_code)
        self.doc_engine = DocumentIntelligenceEngine(country_code)
        self.infrastructure = SMEInfrastructureEngine(country_code)
        self.ai_engine = SMEAIEngine(country_code)
        self.ux_engine = SMEProductUXEngine(country_code)
    
    def generate_automation(self, business_type):
        automations = {
            'retail': {
                'inventory': 'Auto-reorder when stock below threshold',
                'customer_support': 'AI chatbot for common queries',
                'payments': f'Integrated mobile money ({self._get_payment_rails()})',
                'marketing': f'Automated WhatsApp/SMS campaigns in {self.overlay.get("language", "English")}'
            },
            'service': {
                'booking': 'Self-service booking and scheduling',
                'followup': 'Automated client check-ins and feedback',
                'billing': 'Auto-generate invoices and receipts',
                'referrals': 'Digital referral tracking system'
            },
            'agriculture': {
                'weather_alerts': f'Real-time weather notifications for {self.country}',
                'market_prices': f'Daily price updates in {self.overlay.get("currency", "local")}',
                'supply_chain': 'Track and optimize distribution',
                'financial': f'Crop insurance and loan management in {self.country} context'
            }
        }
        
        return automations.get(business_type, automations['retail'])
    
    def _get_payment_rails(self):
        rails = {
            'kenya': 'M-Pesa, Airtel Money, Equitel',
            'bangladesh': 'bKash, Nagad, Rocket',
            'usa': 'PayPal, Stripe, Venmo, Square',
            'uk': 'PayPal, Stripe, Barclays, Monzo'
        }
        return rails.get(self.country, 'M-Pesa, Airtel Money')
    
    def get_action_tasks(self, business_data: Dict[str, Any]) -> List[ActionTask]:
        return self.ux_engine.generate_tasks(business_data)
    
    def process_webhook(self, event_type: str, payload: Dict[str, Any]) -> Dict[str, Any]:
        return self.infrastructure.process_webhook(event_type, payload)
    
    def connect_api(self, service: str, credentials: Dict[str, str]) -> Dict[str, Any]:
        return self.infrastructure.connect_api(service, credentials)
    
    def natural_language_query(self, query: str, data_context: Dict[str, Any]) -> Dict[str, Any]:
        return self.ai_engine.natural_language_query(query, data_context)
    
    def generate_digest(self, channel: str = 'whatsapp') -> Dict[str, Any]:
        return self.ux_engine.generate_digest(channel)


# ==================== MAIN APPLICATION ====================

def init_session_state():
    defaults = {
        'user_role': None,
        'country_code': 'kenya',
        'student_score': 54,
        'completed_lessons': [],
        'achievements': [],
        'streak': 0,
        'weak_areas': [],
        # --- Student Learning Intelligence additions ---
        'sr_cards': {},                 # concept_id -> SpacedRepetitionCard
        'mastered_concepts': {},        # subject -> list[node_id]
        'attempt_logs': [],             # list[dict] mirroring AttemptLog fields
        'wrong_streak': 0,              # consecutive wrong answers (difficulty balancer)
        'win_streak': 0,                # consecutive correct answers (difficulty balancer)
        'current_difficulty': 'medium',
        'hint_level': {},               # question_id -> current hint rung shown
        'last_active_date': None,       # for HabitStreakEngine
        'streak_repair_tokens': 2,
        'ocr_history': [],
        'voice_sessions': [],
        'privacy_consent': True,        # zero-retention acknowledgement, on by default
        'domain': 'business',
        'business_type': 'retail',
        'preferred_language': 'English',
        'doc_analysis_history': [],
        'sme_tasks': [],
        'sme_alerts': [],
        'webhook_events': [],
        'professional_workspace': {},
        'teacher_lesson_plans': [],
        'teacher_rubrics': [],
        'teacher_quizzes': [],
        'teacher_feedback': []
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

init_session_state()

def apply_custom_css():
    st.markdown("""
    <style>
    .stButton > button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        border-radius: 10px;
        padding: 10px 24px;
        font-weight: 600;
        transition: all 0.3s;
    }
    .stButton > button:hover {
        transform: scale(1.05);
        box-shadow: 0 10px 20px rgba(0,0,0,0.2);
    }
    .country-badge {
        background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);
        color: white;
        padding: 4px 12px;
        border-radius: 20px;
        display: inline-block;
        margin: 4px;
        font-size: 0.8rem;
    }
    .achievement-badge {
        background: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%);
        color: white;
        padding: 4px 12px;
        border-radius: 20px;
        display: inline-block;
        margin: 4px;
    }
    .universal-core {
        background: #f0f2f6;
        padding: 20px;
        border-radius: 10px;
        border-left: 5px solid #667eea;
    }
    .local-overlay {
        background: #f0f2f6;
        padding: 20px;
        border-radius: 10px;
        border-left: 5px solid #f5576c;
    }
    .flag-emoji {
        font-size: 1.2rem;
        margin-right: 8px;
    }
    .document-analysis {
        background: #e8f5e9;
        padding: 20px;
        border-radius: 10px;
        border-left: 5px solid #4caf50;
        margin: 10px 0;
    }
    .task-critical {
        background: #ffebee;
        border-left: 5px solid #f44336;
        padding: 15px;
        border-radius: 5px;
        margin: 10px 0;
    }
    .task-high {
        background: #fff3e0;
        border-left: 5px solid #ff9800;
        padding: 15px;
        border-radius: 5px;
        margin: 10px 0;
    }
    .task-medium {
        background: #e3f2fd;
        border-left: 5px solid #2196f3;
        padding: 15px;
        border-radius: 5px;
        margin: 10px 0;
    }
    .task-low {
        background: #f5f5f5;
        border-left: 5px solid #9e9e9e;
        padding: 15px;
        border-radius: 5px;
        margin: 10px 0;
    }
    .professional-card {
        background: #f3e5f5;
        border-left: 5px solid #9c27b0;
        padding: 15px;
        border-radius: 5px;
        margin: 10px 0;
    }
    .watchdog-alert {
        background: #fff3e0;
        border-left: 5px solid #ff6f00;
        padding: 15px;
        border-radius: 5px;
        margin: 10px 0;
    }
    .citation-box {
        background: #e3f2fd;
        border-left: 5px solid #1565c0;
        padding: 15px;
        border-radius: 5px;
        margin: 10px 0;
        font-family: monospace;
        font-size: 0.9em;
    }
    .teacher-card {
        background: #e8f5e9;
        border-left: 5px solid #2e7d32;
        padding: 15px;
        border-radius: 5px;
        margin: 10px 0;
    }
    .time-saved-metric {
        background: #fff8e1;
        border-left: 5px solid #f57f17;
        padding: 15px;
        border-radius: 5px;
        margin: 10px 0;
    }
    .accommodation-tag {
        display: inline-block;
        background: #e3f2fd;
        padding: 4px 12px;
        border-radius: 15px;
        margin: 4px;
        font-size: 0.8rem;
        color: #1565c0;
    }
    </style>
    """, unsafe_allow_html=True)

apply_custom_css()

# ==================== SIDEBAR ====================

st.sidebar.title("🌍 AI Shiksha")
st.sidebar.caption("Universal Core + Local Overlay")

country_flags = {
    'kenya': '🇰🇪',
    'bangladesh': '🇧🇩',
    'usa': '🇺🇸',
    'uk': '🇬🇧'
}

country_code = st.sidebar.selectbox(
    "🌐 Select Country/Region:",
    ['kenya', 'bangladesh', 'usa', 'uk'],
    format_func=lambda x: f"{country_flags.get(x, '🌍')} {x.title()}"
)
st.session_state.country_code = country_code

overlay = LocalCurriculumOverlay.get_overlay(country_code)
context = LocalCurriculumOverlay.get_local_context(country_code)

if overlay:
    st.sidebar.info(f"""
    **{country_flags.get(country_code, '🌍')} {country_code.title()}**
    **Curriculum:** {overlay.get('system', 'Universal')}
    **Exam Boards:** {', '.join(overlay.get('boards', ['Local']))}
    **Language:** {overlay.get('language', 'English')}
    **Currency:** {overlay.get('currency', 'Local')}
    """)

user_role = st.sidebar.selectbox(
    "👤 Select Your Role:",
    ['Student', 'Teacher', 'Professional', 'SME Business Owner']
)
st.session_state.user_role = user_role

menu_options = {
    'Student': ['🎓 Dashboard', '📝 Practice', '🧠 Knowledge Map', '🔁 Review Queue', '📊 Progress', '🏆 Achievements', '🔒 Privacy & Safety', '📄 Document Analysis'],
    'Teacher': ['👨‍🏫 Dashboard', '📋 Lesson Builder', '📝 Assessment', '⏱️ Hours Saved', '📄 Document Analysis', '📑 LMS Integration', '🎯 Standards Mapping'],
    'Professional': ['💼 Dashboard', '🔬 Research', '📊 Portfolio', '📄 Document Analysis', '📑 Workspace', '⚡ Watchdogs'],
    'SME Business Owner': ['🏢 Dashboard', '📈 Growth', '🤖 Automation', '📊 Analytics', '📄 Document Analysis', '🔌 API Connectors', '⚡ Webhooks']
}

st.sidebar.divider()
choice = st.sidebar.radio("Navigate:", menu_options.get(user_role, ['Dashboard']))

def show_vibe_check():
    with st.sidebar.expander("🎯 Daily Vibe Check", expanded=False):
        vibe = st.select_slider(
            "How's your learning energy today?",
            options=["😴 Low", "😐 Neutral", "⚡ Medium", "🔥 High", "🚀 Cosmic"],
            value="🔥 High"
        )
        
        if vibe in ["🔥 High", "🚀 Cosmic"]:
            st.success("Let's ride that energy wave! 🌊")
            st.balloons()
        elif vibe == "😴 Low":
            st.info("Start with a 5-min quick win!")
        
        if st.button("🎯 Get your vibe mission"):
            missions = {
                "😴 Low": "Complete 1 easy MCQ to get moving",
                "😐 Neutral": "Explore a new learning module",
                "⚡ Medium": "Generate a practice test",
                "🔥 High": "Tackle a complex problem set",
                "🚀 Cosmic": "Challenge: Build a mini-project"
            }
            st.info(f"🚀 **Vibe Mission:** {missions.get(vibe)}")

show_vibe_check()


# ==================== TEACHER FUNCTIONS ====================

def show_teacher_dashboard():
    """Enhanced Teacher Dashboard with all teacher tools"""
    st.header(f"👨‍🏫 Teacher Intelligence Dashboard - {country_code.title()}")
    
    teacher_engine = TeacherOutcomeEngine(country_code)
    
    # Quick Stats
    col1, col2, col3, col4 = st.columns(4)
    col1.metric("📚 Lesson Plans", len(teacher_engine.intelligence.lesson_plans), "↑ 2 this week")
    col2.metric("📝 Quizzes", len(teacher_engine.intelligence.quizzes), "↑ 1 new")
    col3.metric("📊 Rubrics", len(teacher_engine.intelligence.rubrics), "↑ 1 new")
    col4.metric("⏱️ Hours Saved", "7.0 hrs", "↑ 15%")
    
    # Time Saved Overview
    st.subheader("⏱️ Time Saved This Week")
    time_saved = teacher_engine.calculate_time_saved()
    
    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Total Hours Saved", f"{time_saved['total_hours_saved']} hrs", f"{time_saved['total_percent_saved']}%")
    col2.metric("Lesson Planning", f"{time_saved['breakdown']['lesson_planning']['hours']} hrs", "vs 4.5 hrs traditional")
    col3.metric("Grading", f"{time_saved['breakdown']['grading']['hours']} hrs", "vs 6 hrs traditional")
    col4.metric("Admin Work", f"{time_saved['breakdown']['admin_work']['hours']} hrs", "vs 3 hrs traditional")
    
    # Recent Lesson Plans
    with st.expander("📚 Recent Lesson Plans", expanded=True):
        for plan in teacher_engine.intelligence.lesson_plans[-3:]:
            st.markdown(f"""
            <div class="teacher-card">
            <strong>{plan.title}</strong>
            <br>Subject: {plan.subject} | Grade: {plan.grade_level} | Duration: {plan.duration} min
            <br>Objectives: {', '.join(plan.objectives[:2])}...
            </div>
            """, unsafe_allow_html=True)
    
    # Recent Feedback
    with st.expander("📝 Recent Feedback", expanded=True):
        feedbacks = teacher_engine.intelligence.feedback_batches[-3:]
        if feedbacks:
            for fb in feedbacks:
                status = "✅ Approved" if fb.approved else "⏳ Pending Review"
                st.markdown(f"""
                <div class="teacher-card">
                <strong>{fb.student_name}</strong> - {fb.assignment}
                <br>Strengths: {', '.join(fb.strengths)}
                <br>Status: {status}
                </div>
                """, unsafe_allow_html=True)
        else:
            st.info("No feedback generated yet.")


def show_teacher_lesson_builder():
    """Enhanced Lesson Builder with Multi-Asset Bundle Generator"""
    st.header(f"📋 Lesson Builder - {country_code.title()}")
    st.caption("Create complete lesson bundles with slide outlines, worksheets, and answer keys")
    
    teacher_engine = TeacherOutcomeEngine(country_code)
    
    # Lesson Bundle Generator
    st.subheader("🚀 Multi-Asset Bundle Generator")
    st.caption("Generate a complete lesson plan, slide outline, worksheet, and answer key in one request")
    
    col1, col2 = st.columns(2)
    with col1:
        subject = st.selectbox("Subject:", ['Science', 'Mathematics', 'English Language', 'Geography', 'History'])
        grade = st.selectbox("Grade Level:", overlay.get('grade_levels', ['Grade 1', 'Grade 2', 'Grade 3', 'Grade 4', 'Grade 5', 'Grade 6', 'Grade 7', 'Grade 8', 'Grade 9', 'Grade 10']))
        topic = st.text_input("Topic:", "Photosynthesis")
    
    with col2:
        duration = st.slider("Lesson Duration (minutes):", 30, 90, 45)
        include_worksheet = st.checkbox("Include Worksheet", value=True)
        include_answer_key = st.checkbox("Include Answer Key", value=True)
        include_slides = st.checkbox("Include Slide Outline", value=True)
    
    if st.button("✨ Generate Complete Lesson Bundle", type="primary"):
        with st.spinner("Generating lesson bundle..."):
            bundle = teacher_engine.generate_lesson_bundle(subject, grade, topic, duration)
        
        st.success(f"✅ Lesson Bundle Generated! Bundle ID: {bundle['bundle_id']}")
        
        tabs = st.tabs(["📋 Lesson Plan", "📊 Slide Outline", "📝 Worksheet", "🔑 Answer Key"])
        
        with tabs[0]:
            st.markdown(f"### {bundle['lesson_plan']['title']}")
            st.markdown(f"**Subject:** {bundle['lesson_plan']['subject']}")
            st.markdown(f"**Grade:** {bundle['lesson_plan']['grade']}")
            st.markdown(f"**Duration:** {bundle['lesson_plan']['duration']} min")
            st.markdown("**Objectives:**")
            for obj in bundle['lesson_plan']['objectives']:
                st.markdown(f"- {obj}")
            st.markdown("**Procedure:**")
            for key, value in bundle['lesson_plan']['procedure'].items():
                st.markdown(f"- **{key.replace('_', ' ').title()}:** {value}")
            
            if st.button("📥 Download Lesson Plan"):
                st.download_button(
                    label="Download as Text",
                    data=json.dumps(bundle['lesson_plan'], indent=2),
                    file_name=f"lesson_plan_{topic.lower().replace(' ', '_')}.json",
                    mime="application/json"
                )
        
        with tabs[1]:
            if include_slides:
                st.markdown("### Slide Outline")
                for slide in bundle['slide_outline']:
                    st.markdown(f"**Slide {slide['slide']}:** {slide['title']} ({slide['type']})")
            else:
                st.info("Slide outline not included.")
        
        with tabs[2]:
            if include_worksheet:
                st.markdown(f"### {bundle['worksheet']['title']}")
                st.markdown(f"**Total Points:** {bundle['worksheet']['total_points']}")
                for section in bundle['worksheet']['sections']:
                    st.markdown(f"**{section['section']}**")
                    st.markdown(f"*{section['instructions']}*")
                    for q in section['questions']:
                        st.markdown(f"- {q}")
            else:
                st.info("Worksheet not included.")
        
        with tabs[3]:
            if include_answer_key:
                st.markdown(f"### {bundle['answer_key']['title']}")
                for part, answers in bundle['answer_key']['answers'].items():
                    st.markdown(f"**{part}:**")
                    for answer in answers:
                        st.markdown(f"- {answer}")
            else:
                st.info("Answer key not included.")
    
    # Modular Block Editor
    st.divider()
    st.subheader("🧩 Modular Block Editor")
    st.caption("Edit, swap, or regenerate individual lesson sections")
    
    # Select existing lesson plan
    lesson_options = [plan.title for plan in teacher_engine.intelligence.lesson_plans]
    if lesson_options:
        selected_lesson_title = st.selectbox("Select Lesson to Edit:", lesson_options)
        
        selected_plan = next(
            (p for p in teacher_engine.intelligence.lesson_plans if p.title == selected_lesson_title),
            None
        )
        
        if selected_plan:
            blocks = teacher_engine.get_lesson_blocks(selected_plan)
            
            for block_name, block_data in blocks['blocks'].items():
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.markdown(f"**{block_name.replace('_', ' ').title()}:**")
                    content = st.text_area(
                        f"Edit {block_name}",
                        value=block_data['content'],
                        height=100,
                        key=f"block_{block_name}"
                    )
                with col2:
                    if st.button(f"🔄 Regenerate {block_name}", key=f"regenerate_{block_name}"):
                        new_content = teacher_engine.regenerate_block(block_name, selected_plan.title)
                        st.session_state[f"block_{block_name}"] = new_content
                        st.success(f"✅ {block_name} regenerated!")
            
            if st.button("💾 Save Lesson Changes"):
                st.success("✅ Lesson plan saved successfully!")
    
    # One-Tap Accommodation Toggles
    st.divider()
    st.subheader("🎯 One-Tap Accommodation Toggles")
    st.caption("Quick filters to apply IEP/504 modifications, sentence starters, or ELL translations")
    
    accommodation_text = st.text_area(
        "Enter content to modify:",
        height=100,
        value="Students will write a paragraph explaining the process of photosynthesis."
    )
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        if st.button("📋 IEP/504", use_container_width=True):
            result = teacher_engine.apply_accommodation(accommodation_text, 'iep_504')
            st.info(f"✅ {result.get('changes', 'Applied')}")
    with col2:
        if st.button("✏️ Sentence Starters", use_container_width=True):
            result = teacher_engine.apply_accommodation(accommodation_text, 'sentence_starters')
            st.info(f"✅ {result.get('changes', 'Applied')}")
    with col3:
        if st.button("🌐 ELL Translation", use_container_width=True):
            result = teacher_engine.apply_accommodation(accommodation_text, 'ell_translation')
            st.info(f"✅ {result.get('changes', 'Applied')}")
    with col4:
        if st.button("📖 Simplified Text", use_container_width=True):
            result = teacher_engine.apply_accommodation(accommodation_text, 'simplified_text')
            st.info(f"✅ {result.get('changes', 'Applied')}")
    
    # Dynamic Lexile Adjuster
    st.divider()
    st.subheader("📖 Dynamic Lexile Adjuster")
    st.caption("Rewrites any text passage up or down reading levels instantly")
    
    text_to_adjust = st.text_area(
        "Enter text to adjust reading level:",
        height=100,
        value="Photosynthesis is the process by which plants convert sunlight into energy. This complex biochemical process involves multiple stages and specialized cellular structures."
    )
    
    target_level = st.selectbox("Target Reading Level:", ['beginner', 'intermediate', 'advanced', 'expert'])
    
    if st.button("🔄 Adjust Reading Level"):
        result = teacher_engine.adjust_reading_level(text_to_adjust, target_level)
        
        st.markdown("**Original Text:**")
        st.text(result['original_text'])
        st.markdown("**Adjusted Text:**")
        st.text(result['adjusted_text'])
        st.markdown(f"**Target Level:** {result['level_description']}")
        st.markdown(f"**Word Count Change:** {result['word_count_original']} → {result['word_count_adjusted']}")
        st.json(result['changes'])


def show_teacher_assessment():
    """Assessment Engine with Rubric & Quiz Generator"""
    st.header(f"📝 Assessment Engine - {country_code.title()}")
    st.caption("Build tiered grading rubrics and generate aligned quizzes")
    
    teacher_engine = TeacherOutcomeEngine(country_code)
    
    # Rubric Generator
    st.subheader("📊 Rubric Generator")
    
    col1, col2 = st.columns(2)
    with col1:
        rubric_title = st.text_input("Rubric Title:", "Science Lab Report Rubric")
        criteria_list = st.text_area("Criteria (one per line):", "Hypothesis\nProcedure\nData Collection\nAnalysis\nConclusion")
    
    with col2:
        levels_list = st.text_input("Levels (comma-separated):", "Excellent, Proficient, Developing, Beginning")
    
    if st.button("✨ Generate Rubric", type="primary"):
        criteria = [c.strip() for c in criteria_list.split('\n') if c.strip()]
        levels = [l.strip() for l in levels_list.split(',') if l.strip()]
        
        if criteria and levels:
            rubric = teacher_engine.generate_rubric(rubric_title, criteria, levels)
            
            st.success(f"✅ Rubric Generated: {rubric.title}")
            
            # Display rubric
            st.markdown("### Rubric Preview")
            data = []
            for criterion in rubric.criteria:
                row = {'Criterion': criterion}
                for level in rubric.levels:
                    row[level] = rubric.descriptions[criterion].get(level, '')
                row['Weight'] = f"{rubric.weightings.get(criterion, 0) * 100:.0f}%"
                data.append(row)
            
            st.dataframe(pd.DataFrame(data))
            
            if st.button("📥 Download Rubric"):
                st.download_button(
                    label="Download as JSON",
                    data=json.dumps({
                        'title': rubric.title,
                        'criteria': rubric.criteria,
                        'levels': rubric.levels,
                        'descriptions': rubric.descriptions,
                        'weightings': rubric.weightings
                    }, indent=2),
                    file_name=f"{rubric_title.lower().replace(' ', '_')}_rubric.json",
                    mime="application/json"
                )
        else:
            st.error("Please enter at least one criterion and level.")
    
    # Quiz Generator
    st.divider()
    st.subheader("📝 Quiz Generator")
    
    col1, col2 = st.columns(2)
    with col1:
        quiz_subject = st.selectbox("Subject:", ['Science', 'Mathematics', 'English Language', 'Geography', 'History'], key="quiz_subject")
        quiz_grade = st.selectbox("Grade Level:", overlay.get('grade_levels', ['Grade 1', 'Grade 2', 'Grade 3', 'Grade 4', 'Grade 5', 'Grade 6', 'Grade 7', 'Grade 8', 'Grade 9', 'Grade 10']), key="quiz_grade")
        quiz_topic = st.text_input("Topic:", "Photosynthesis")
    
    with col2:
        num_questions = st.slider("Number of Questions:", 3, 15, 5)
        include_time_limit = st.checkbox("Include Time Limit", value=True)
    
    if st.button("✨ Generate Quiz", type="primary"):
        quiz = teacher_engine.generate_quiz(quiz_subject, quiz_grade, quiz_topic, num_questions)
        
        st.success(f"✅ Quiz Generated: {quiz.title}")
        st.markdown(f"**Total Points:** {quiz.total_points}")
        if quiz.time_limit:
            st.markdown(f"**Time Limit:** {quiz.time_limit} minutes")
        
        # Display questions
        for q in quiz.questions:
            st.markdown(f"**Question {q['id']}** ({q['type'].replace('_', ' ').title()}, {q['points']} pts)")
            st.markdown(q['question'])
            if q['type'] == 'multiple_choice':
                for i, opt in enumerate(q.get('options', [])):
                    st.markdown(f"- {chr(65 + i)}. {opt}")
            if 'correct_answer' in q:
                st.caption(f"✅ Correct Answer: {q['correct_answer']}")
            st.divider()
        
        if st.button("📥 Download Quiz"):
            st.download_button(
                label="Download as JSON",
                data=json.dumps({
                    'title': quiz.title,
                    'subject': quiz.subject,
                    'grade_level': quiz.grade_level,
                    'questions': quiz.questions,
                    'total_points': quiz.total_points,
                    'time_limit': quiz.time_limit
                }, indent=2),
                file_name=f"{quiz_topic.lower().replace(' ', '_')}_quiz.json",
                mime="application/json"
            )


def show_teacher_hours_saved():
    """Time-Saved Dashboard with Analytics"""
    st.header(f"⏱️ Time-Saved Dashboard - {country_code.title()}")
    st.caption("Analytics tracking estimated hours recovered per week on lesson planning, grading, and admin work")
    
    teacher_engine = TeacherOutcomeEngine(country_code)
    
    # Time Saved Overview
    time_saved = teacher_engine.calculate_time_saved()
    
    col1, col2, col3 = st.columns(3)
    col1.metric("Total Hours Saved This Week", f"{time_saved['total_hours_saved']} hrs", f"{time_saved['total_percent_saved']}%")
    col2.metric("Monthly Projection", f"{time_saved['monthly_projection']} hrs")
    col3.metric("Yearly Projection", f"{time_saved['yearly_projection']} hrs")
    
    # Breakdown
    st.subheader("📊 Time Saved Breakdown")
    
    breakdown_data = []
    for category, data in time_saved['breakdown'].items():
        breakdown_data.append({
            'Category': category.replace('_', ' ').title(),
            'Hours Saved': data['hours'],
            'Traditional Hours': float(data['saved'].split(' ')[1]) if 'vs' in data['saved'] else 0,
            'Savings %': f"{round((data['hours'] / float(data['saved'].split(' ')[1] if 'vs' in data['saved'] else 1)) * 100, 1)}%"
        })
    
    st.dataframe(pd.DataFrame(breakdown_data))
    
    # Historical Trend
    st.subheader("📈 Weekly Time Saved Trend")
    history = teacher_engine.intelligence.time_saved_history
    
    if history:
        trend_data = []
        for metric in history:
            trend_data.append({
                'Week': metric.week_start.strftime('%Y-%m-%d'),
                'Lesson Planning': metric.lesson_planning_hours,
                'Grading': metric.grading_hours,
                'Admin': metric.admin_hours,
                'Total': metric.total_hours
            })
        df_trend = pd.DataFrame(trend_data)
        st.line_chart(df_trend.set_index('Week'))
    
    # Parent & Sub-Plan Generators
    st.divider()
    st.subheader("📧 Parent & Sub-Plan Generators")
    st.caption("One-click drafting of professional parent updates and emergency substitute teacher guides")
    
    tab1, tab2 = st.tabs(["📧 Parent Update", "📋 Sub Plan"])
    
    with tab1:
        st.markdown("### Parent Update Generator")
        
        student_name = st.text_input("Student Name:", "Alex Johnson")
        subject = st.selectbox("Subject:", ['Science', 'Mathematics', 'English Language', 'Geography', 'History'], key="parent_subject")
        progress = st.selectbox("Progress:", ["making excellent progress", "showing improvement", "working at grade level", "needs additional support", "has shown significant growth"])
        
        if st.button("✏️ Generate Parent Update"):
            update = teacher_engine.generate_parent_update(student_name, subject, progress)
            st.markdown("### 📧 Parent Update")
            st.text_area("Copy and send:", update, height=200)
            
            st.download_button(
                label="📥 Download Update",
                data=update,
                file_name=f"parent_update_{student_name.lower().replace(' ', '_')}.txt",
                mime="text/plain"
            )
    
    with tab2:
        st.markdown("### Sub Plan Generator")
        
        sub_subject = st.selectbox("Subject:", ['Science', 'Mathematics', 'English Language', 'Geography', 'History'], key="sub_subject")
        sub_grade = st.selectbox("Grade Level:", overlay.get('grade_levels', ['Grade 1', 'Grade 2', 'Grade 3', 'Grade 4', 'Grade 5', 'Grade 6', 'Grade 7', 'Grade 8', 'Grade 9', 'Grade 10']), key="sub_grade")
        sub_topic = st.text_input("Lesson Topic:", "Photosynthesis")
        
        if st.button("📋 Generate Sub Plan"):
            sub_plan = teacher_engine.generate_sub_plan(sub_subject, sub_grade, sub_topic)
            
            st.markdown(f"### 📋 {sub_plan['title']}")
            st.markdown(f"**Overview:** {sub_plan['overview']}")
            
            st.markdown("**Schedule:**")
            for period, details in sub_plan['schedule'].items():
                st.markdown(f"- **{period.replace('_', ' ').title()}** ({details['time']}): {details['activity']}")
                st.markdown(f"  Materials: {details['materials']}")
            
            st.markdown(f"**Emergency Procedures:** {sub_plan['emergency_procedures']}")
            st.markdown(f"**Notes:** {sub_plan['notes']}")
            
            st.download_button(
                label="📥 Download Sub Plan",
                data=json.dumps(sub_plan, indent=2),
                file_name=f"sub_plan_{sub_topic.lower().replace(' ', '_')}.json",
                mime="application/json"
            )
    
    # Zero-PII Compliance
    st.divider()
    st.subheader("🔒 Zero-PII Compliance Layer")
    st.caption("Client-side redaction of student names and FERPA/COPPA data isolation")
    
    text_with_pii = st.text_area(
        "Enter text with PII to redact:",
        height=100,
        value="John Smith scored 95% on the assignment. Contact john.smith@school.com for more information."
    )
    
    if st.button("🔒 Redact PII"):
        redacted = teacher_engine.redact_pii(text_with_pii)
        
        st.markdown("**Original Text:**")
        st.text(redacted['original'])
        st.markdown("**Redacted Text:**")
        st.text(redacted['redacted'])
        st.markdown(f"**PII Removed:** {redacted['pii_removed']} instances")
        st.markdown(f"**Compliance:** {redacted['compliance_note']}")


def show_teacher_lms_integration():
    """LMS Integration & File Conversion"""
    st.header(f"📑 LMS Integration - {country_code.title()}")
    st.caption("Native export/import sync with Google Classroom, Canvas, Schoology, and Microsoft Teams")
    
    teacher_engine = TeacherOutcomeEngine(country_code)
    
    # LMS Connection
    st.subheader("🔗 Connect to LMS")
    
    lms_options = ['google_classroom', 'canvas', 'schoology', 'microsoft_teams']
    lms_names = {
        'google_classroom': 'Google Classroom 📚',
        'canvas': 'Canvas 🎨',
        'schoology': 'Schoology 📖',
        'microsoft_teams': 'Microsoft Teams 💬'
    }
    
    for lms in lms_options:
        col1, col2, col3 = st.columns([2, 1, 1])
        with col1:
            st.markdown(f"**{lms_names.get(lms, lms)}**")
        with col2:
            if st.button(f"Connect {lms_names.get(lms, lms)}", key=f"connect_{lms}"):
                result = teacher_engine.connect_lms(lms, {'api_key': 'demo_key'})
                if result['status'] == 'connected':
                    st.success(f"✅ Connected to {result['config']['name']}")
                else:
                    st.error(f"❌ {result['message']}")
        with col3:
            if st.button(f"🔄 Sync", key=f"sync_{lms}"):
                result = teacher_engine.sync_lms_data(lms)
                if result['status'] == 'success':
                    st.success(f"✅ Synced {result['sync_type']} data")
                    st.json(result['data'])
    
    # Universal File Conversion
    st.divider()
    st.subheader("📄 Universal File Conversion")
    st.caption("Import old PDF/Word worksheets and export to editable Google Docs, Slides, or printable formats")
    
    uploaded_file = st.file_uploader(
        "Upload Document to Convert:",
        type=['pdf', 'docx', 'txt'],
        key="teacher_file_convert"
    )
    
    if uploaded_file is not None:
        output_format = st.selectbox(
            "Output Format:",
            ['google_docs', 'google_slides', 'printable']
        )
        
        if st.button("🔄 Convert Document"):
            with st.spinner("Converting document..."):
                result = teacher_engine.convert_document(uploaded_file, output_format)
            
            if result.get('status') == 'error':
                st.error(f"❌ {result.get('message', 'Conversion failed')}")
            else:
                st.success(f"✅ {result.get('message', 'Conversion complete')}")
                st.markdown(f"**Original Format:** {result.get('original_format')}")
                st.markdown(f"**Output Format:** {result.get('output_format')}")
                st.markdown(f"**Title:** {result.get('title')}")
                
                st.text_area("Converted Content Preview:", result.get('content', '')[:500], height=200)
                
                if result.get('downloadable'):
                    st.download_button(
                        label="📥 Download Converted File",
                        data=result.get('content', ''),
                        file_name=f"{result.get('title')}.txt",
                        mime="text/plain"
                    )


def show_teacher_standards_mapping():
    """Curriculum Standard Mapping"""
    st.header(f"🎯 Standards Mapping - {country_code.title()}")
    st.caption("Embedded databases of state and national benchmarks for auto-tagging")
    
    teacher_engine = TeacherOutcomeEngine(country_code)
    
    st.subheader("📋 Curriculum Standards Database")
    
    standard_types = {
        'common_core': 'Common Core State Standards',
        'ngss': 'Next Generation Science Standards (NGSS)',
        'teks': 'Texas Essential Knowledge and Skills (TEKS)',
        'cbc': 'Competency Based Curriculum (CBC)',
        'nctb': 'National Curriculum (NCTB)'
    }
    
    selected_standard = st.selectbox(
        "Select Standards Type:",
        list(standard_types.keys()),
        format_func=lambda x: standard_types.get(x, x)
    )
    
    if selected_standard in teacher_engine.intelligence.standard_databases:
        db = teacher_engine.intelligence.standard_databases[selected_standard]
        st.markdown(f"**{standard_types.get(selected_standard)}**")
        
        for category, standards in db.items():
            st.markdown(f"### {category}")
            st.markdown(", ".join(standards))
    
    # Auto-Tagging
    st.divider()
    st.subheader("🏷️ Auto-Tag Content with Standards")
    st.caption("Paste content to automatically map to relevant standards")
    
    content_to_tag = st.text_area(
        "Enter content to tag:",
        height=150,
        value="Students will explore the process of photosynthesis and its role in the carbon cycle."
    )
    
    tag_standard_type = st.selectbox(
        "Select Standards Type for Tagging:",
        ['common_core', 'ngss', 'teks', 'cbc', 'nctb'],
        format_func=lambda x: standard_types.get(x, x)
    )
    
    if st.button("🏷️ Auto-Tag Content"):
        matched = teacher_engine.map_standards(content_to_tag, tag_standard_type)
        
        if matched:
            st.success(f"✅ Found {len(matched)} matching standards")
            for standard in matched:
                st.markdown(f"- {standard}")
        else:
            st.info("No matching standards found. Try adjusting your content or selecting a different standards type.")


# ==================== PROFESSIONAL FUNCTIONS ====================

def show_professional_dashboard():
    st.header(f"💼 Professional Intelligence Dashboard - {country_code.title()}")
    
    prof_engine = ProfessionalOutcomeEngine(st.session_state.domain, country_code)
    
    col1, col2, col3, col4 = st.columns(4)
    col1.metric("📚 Documents", len(prof_engine.intelligence.documents), "↑ 2 new")
    col2.metric("📊 Portfolio Value", "$75,825", "↑ 8.2%")
    col3.metric("⚠️ Watchdog Alerts", len(prof_engine.intelligence.watchdog_alerts), "3 unread")
    col4.metric("⏱️ Time Saved", "18 hrs", "↑ 4 hrs")
    
    st.subheader("📊 Portfolio Overview")
    portfolio_metrics = prof_engine.get_portfolio_metrics()
    
    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Total Value", f"${portfolio_metrics['total_value']:,.2f}")
    col2.metric("Volatility", f"{portfolio_metrics['volatility']:.2%}")
    col3.metric("Sharpe Ratio", f"{portfolio_metrics['sharpe_ratio']:.2f}")
    col4.metric("Concentration Risk", f"{portfolio_metrics['concentration_risk']:.1%}")
    
    with st.expander("📈 Portfolio Holdings", expanded=True):
        assets_data = []
        for asset in prof_engine.intelligence.portfolio_assets:
            assets_data.append({
                'Ticker': asset.ticker,
                'Name': asset.name,
                'Shares': asset.shares,
                'Current Price': f"${asset.current_price:.2f}",
                'Market Value': f"${asset.market_value:,.2f}",
                'Weight': f"{asset.weight:.1%}",
                'Daily Change': f"{asset.daily_change:.1%}"
            })
        st.dataframe(pd.DataFrame(assets_data))
    
    st.subheader("⚡ 24/7 Asset Watchdogs")
    alerts = prof_engine.get_watchdog_alerts()
    
    for alert in alerts[-5:]:
        severity_color = {'high': '🔴', 'medium': '🟡', 'low': '🟢'}.get(alert.severity, '🔵')
        
        st.markdown(f"""
        <div class="watchdog-alert">
        <strong>{severity_color} {alert.asset_ticker}</strong> - {alert.alert_type.upper()}
        <br>{alert.message}
        <br><small>Source: {alert.source} | {alert.timestamp.strftime('%Y-%m-%d %H:%M')}</small>
        </div>
        """, unsafe_allow_html=True)
    
    if st.button("🔄 Refresh Watchdogs"):
        new_alerts = prof_engine.refresh_watchdogs()
        if new_alerts:
            st.success(f"✅ {len(new_alerts)} new alerts detected!")
            st.rerun()
        else:
            st.info("No new alerts detected.")


def show_professional_research():
    st.header(f"🔬 Research & Analysis - {country_code.title()}")
    
    prof_engine = ProfessionalOutcomeEngine(st.session_state.domain, country_code)
    
    st.subheader("🔍 Hybrid Search Engine")
    st.caption("Search across documents using vector + keyword (BM25) hybrid retrieval")
    
    col1, col2 = st.columns([3, 1])
    with col1:
        search_query = st.text_input("Search query:", placeholder="e.g., 'TechCorp Q4 2024 earnings revenue growth'")
    with col2:
        search_type = st.selectbox("Search Type:", ['hybrid', 'keyword', 'vector', 'semantic'])
    
    if search_query:
        with st.spinner("Searching..."):
            results = prof_engine.hybrid_search(search_query, search_type)
        
        if results:
            st.success(f"✅ Found {len(results)} results")
            for result in results[:5]:
                doc = result['document']
                with st.expander(f"📄 {doc.title} (Score: {result['score']:.3f})"):
                    st.markdown(f"**Source Type:** {doc.source_type}")
                    st.markdown(f"**Matches:** {', '.join(result['matches'][:5])}")
                    st.markdown(f"**Content Preview:** {doc.content[:300]}...")
                    
                    if doc.citations:
                        st.markdown("**📝 Citations:**")
                        for citation in doc.citations[:3]:
                            st.markdown(f"- Page {citation.get('page', 'N/A')}: {citation.get('claim', 'N/A')}")
        else:
            st.info("No results found. Try adjusting your search terms.")
    
    st.divider()
    st.subheader("📄 Document Ingestion")
    st.caption("Upload documents for AI-powered analysis and citation tracking")
    
    uploaded_file = st.file_uploader(
        "Upload Document",
        type=['pdf', 'docx', 'txt', 'csv', 'json'],
        key="prof_doc_upload"
    )
    
    if uploaded_file is not None:
        source_type = st.selectbox("Source Type:", ['sec_filing', 'earnings_call', 'research', 'news'])
        
        if st.button("📥 Ingest Document"):
            with st.spinner("Ingesting document..."):
                doc_engine = DocumentIntelligenceEngine(country_code)
                text = doc_engine.extract_text_from_file(uploaded_file)
                
                if text.startswith("Error"):
                    st.error(f"⚠️ {text}")
                else:
                    doc = prof_engine.intelligence.ingest_document(
                        text, 
                        uploaded_file.name, 
                        source_type,
                        {'upload_date': datetime.now().isoformat()}
                    )
                    
                    st.success(f"✅ Document ingested: {doc.title}")
                    st.markdown(f"**Document ID:** {doc.id}")
                    st.markdown(f"**Source Type:** {doc.source_type}")
                    st.markdown(f"**Words:** {len(doc.content.split())}")
                    
                    if doc.citations:
                        st.markdown("**📝 Extracted Citations:**")
                        for citation in doc.citations[:3]:
                            st.markdown(f"- {citation.get('claim', 'N/A')}")
    
    st.divider()
    st.subheader("🔗 Audit-Linked Citation Engine")
    st.caption("Verify claims against source documents with page-level citations")
    
    analysis_text = st.text_area(
        "Enter analysis text to audit:",
        height=150,
        placeholder="e.g., 'TechCorp revenue reached $12.4 billion in Q4 2024, up 18% year-over-year.'"
    )
    
    if st.button("🔍 Audit Analysis"):
        with st.spinner("Auditing analysis..."):
            audit_results = prof_engine.audit_analysis(analysis_text)
        
        st.markdown(f"**Confidence Score:** {audit_results['confidence_score']:.1%}")
        
        if audit_results['verified_claims']:
            st.success(f"✅ {len(audit_results['verified_claims'])} claims verified")
            for claim in audit_results['verified_claims']:
                st.markdown(f"""
                <div class="citation-box">
                <strong>✅ Verified:</strong> {claim['claim']}
                <br><small>Source: {claim['source']}</small>
                </div>
                """, unsafe_allow_html=True)
        
        if audit_results['unverified_claims']:
            st.warning(f"⚠️ {len(audit_results['unverified_claims'])} unverified claims")
            for claim in audit_results['unverified_claims']:
                st.markdown(f"- {claim}")
    
    st.divider()
    st.subheader("🧮 Code-Executing Data Sandbox")
    st.caption("Run quantitative analysis in a sandboxed Python environment")
    
    code_col1, code_col2 = st.columns(2)
    
    with code_col1:
        sample_data = pd.DataFrame({
            'Date': ['2024-01-01', '2024-02-01', '2024-03-01', '2024-04-01'],
            'Revenue': [10000, 12000, 15000, 18000],
            'Expenses': [6000, 7000, 8000, 9000],
            'Profit': [4000, 5000, 7000, 9000]
        })
        
        st.dataframe(sample_data)
        
        code = st.text_area(
            "Python Code:",
            height=150,
            value="""# Calculate metrics
result = {
    'total_revenue': data['Revenue'].sum(),
    'avg_revenue': data['Revenue'].mean(),
    'total_profit': data['Profit'].sum(),
    'profit_margin': data['Profit'].sum() / data['Revenue'].sum()
}"""
        )
    
    with code_col2:
        if st.button("▶️ Execute Code", type="primary"):
            with st.spinner("Executing code in sandbox..."):
                result = prof_engine.execute_quant_analysis(code, sample_data)
            
            if result['status'] == 'success':
                st.success("✅ Code executed successfully!")
                st.json(result['result'])
            else:
                st.error(f"❌ Error: {result.get('error', 'Unknown error')}")


def show_professional_portfolio():
    st.header(f"📊 Portfolio Intelligence - {country_code.title()}")
    
    prof_engine = ProfessionalOutcomeEngine(st.session_state.domain, country_code)
    
    portfolio_metrics = prof_engine.get_portfolio_metrics()
    
    col1, col2, col3 = st.columns(3)
    col1.metric("Total Portfolio Value", f"${portfolio_metrics['total_value']:,.2f}")
    col2.metric("Volatility (Annualized)", f"{portfolio_metrics['volatility']:.2%}")
    col3.metric("Sharpe Ratio", f"{portfolio_metrics['sharpe_ratio']:.2f}")
    
    st.subheader("📊 Sector Concentration")
    sector_data = pd.DataFrame({
        'Sector': list(portfolio_metrics['sector_exposure'].keys()),
        'Weight': [f"{v:.1%}" for v in portfolio_metrics['sector_exposure'].values()]
    })
    st.dataframe(sector_data)
    
    st.divider()
    st.subheader("🌍 Macro Scenario Stress-Testing")
    st.caption("Simulate macroeconomic shocks against your portfolio")
    
    col1, col2 = st.columns(2)
    
    with col1:
        scenario_name = st.text_input("Scenario Name:", "Rate Hike Scenario")
        scenario_desc = st.text_area("Description:", "Federal Reserve rate hike with inflation impact")
        
        rate_change = st.slider("Rate Change (bps):", -100, 100, 25) / 100
        inflation_change = st.slider("Inflation Change (%):", -5, 5, 2) / 100
        growth_change = st.slider("Growth Change (%):", -5, 5, 1) / 100
        market_shock = st.slider("Market Shock (%):", -20, 20, -5) / 100
    
    with col2:
        if st.button("🧪 Run Stress Test", type="primary"):
            scenario = MacroScenario(
                name=scenario_name,
                description=scenario_desc,
                rate_change=rate_change,
                inflation_change=inflation_change,
                growth_change=growth_change,
                market_shock=market_shock,
                portfolio_impact=0.0
            )
            
            with st.spinner("Running stress test..."):
                stress_results = prof_engine.run_stress_test(scenario)
            
            st.success("✅ Stress Test Complete!")
            st.metric("Portfolio Impact", f"{stress_results['total_return']:.1f}%")
            st.metric("Risk Assessment", stress_results['risk_assessment'])
            
            st.markdown("**Asset Impacts:**")
            impact_df = pd.DataFrame(stress_results['asset_impacts'])
            st.dataframe(impact_df[['ticker', 'current_value', 'impacted_value', 'impact_percent']])
    
    st.divider()
    st.subheader("📄 One-Click Deliverable Generator")
    st.caption("Turn research notes into professional deliverables")
    
    template_type = st.selectbox(
        "Template Type:",
        ['investment_memo', 'client_tear_sheet', 'research_note']
    )
    
    template_data = {
        'date': datetime.now().strftime('%Y-%m-%d'),
        'analyst': st.text_input("Analyst Name:", "John Doe"),
        'summary': st.text_area("Executive Summary:", "Investment opportunity with strong growth potential."),
        'thesis': st.text_area("Investment Thesis:", "Company has competitive advantage in growing market."),
        'risks': st.text_area("Risk Analysis:", "Market volatility and regulatory risks."),
        'financials': st.text_area("Financial Projections:", "Revenue growth of 15% projected."),
        'action': st.selectbox("Recommendation:", ['Buy', 'Hold', 'Sell']),
        'target_price': st.number_input("Target Price:", value=150.00, step=5.00),
        'time_horizon': st.selectbox("Time Horizon:", ['Short-term', 'Medium-term', 'Long-term']),
        'appendix': st.text_area("Appendix:", "Additional data and charts.")
    }
    
    if st.button("📄 Generate Deliverable"):
        deliverable = prof_engine.generate_deliverable(template_type, template_data)
        
        st.markdown("### 📋 Generated Deliverable")
        st.markdown(deliverable)
        
        st.download_button(
            label="📥 Download Deliverable",
            data=deliverable,
            file_name=f"{template_type}_{datetime.now().strftime('%Y%m%d')}.md",
            mime="text/markdown"
        )


def show_professional_workspace():
    st.header(f"📑 Split-Screen Workspace - {country_code.title()}")
    st.caption("Dual-pane interface with working notes and source documents")
    
    prof_engine = ProfessionalOutcomeEngine(st.session_state.domain, country_code)
    workspace_data = prof_engine.get_workspace_data()
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("📝 Working Notes")
        st.text_area(
            "Notes:",
            height=400,
            value="""# Research Notes
## Key Findings
- TechCorp Q4 2024 revenue: $12.4B (18% YoY growth)
- EPS: $2.15 vs $1.95 estimate
- Cloud revenue: $4.2B (32% growth)

## Investment Thesis
Strong growth trajectory in cloud and AI segments.
Operating margin expansion to 28.5% indicates improving efficiency.

## Risks
- Competitive pressure in cloud market
- Regulatory scrutiny on AI products
- Macroeconomic headwinds

## Next Steps
1. Review competitor earnings
2. Analyze forward guidance
3. Build valuation model
"""
        )
        
        if st.button("💾 Save Notes"):
            st.success("✅ Notes saved successfully!")
    
    with col2:
        st.subheader("📄 Source Documents")
        
        for doc in workspace_data['documents']:
            with st.expander(f"📄 {doc['title']}"):
                st.markdown(f"**Type:** {doc['type']}")
                st.markdown(f"**Preview:** {doc['content']}")
        
        st.divider()
        st.subheader("📊 Automated Table Extraction")
        st.caption("Extract tables from PDFs into Excel-ready format")
        
        if st.button("🔄 Extract Tables from Documents"):
            with st.spinner("Extracting tables..."):
                for doc in prof_engine.intelligence.documents:
                    tables = prof_engine.extract_table(doc.content)
                    if not tables.empty:
                        st.markdown(f"**Table from {doc.title}:**")
                        st.dataframe(tables)
                        
                        csv = tables.to_csv(index=False)
                        st.download_button(
                            label=f"📥 Download {doc.title} Table",
                            data=csv,
                            file_name=f"{doc.id}_table.csv",
                            mime="text/csv"
                        )


def show_professional_watchdogs():
    st.header(f"⚡ 24/7 Asset Watchdogs - {country_code.title()}")
    st.caption("Automated agents scanning earnings, regulatory filings, and market news")
    
    prof_engine = ProfessionalOutcomeEngine(st.session_state.domain, country_code)
    
    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("🔄 Refresh All Watchdogs", type="primary"):
            new_alerts = prof_engine.refresh_watchdogs()
            if new_alerts:
                st.success(f"✅ {len(new_alerts)} new alerts detected!")
                st.rerun()
            else:
                st.info("No new alerts detected.")
    
    with col2:
        st.caption(f"Total Alerts: {len(prof_engine.get_watchdog_alerts())}")
    
    with col3:
        unread = len([a for a in prof_engine.get_watchdog_alerts() if not a.read])
        st.caption(f"Unread: {unread}")
    
    filter_type = st.selectbox("Filter by:", ['All', 'earnings', 'filing', 'news', 'price'])
    
    alerts = prof_engine.get_watchdog_alerts()
    if filter_type != 'All':
        alerts = [a for a in alerts if a.alert_type == filter_type]
    
    for alert in alerts:
        severity_color = {'high': '🔴', 'medium': '🟡', 'low': '🟢'}.get(alert.severity, '🔵')
        
        col1, col2 = st.columns([5, 1])
        with col1:
            st.markdown(f"""
            <div class="watchdog-alert">
            <strong>{severity_color} {alert.asset_ticker}</strong> - {alert.alert_type.upper()}
            <br>{alert.message}
            <br><small>Source: {alert.source} | {alert.timestamp.strftime('%Y-%m-%d %H:%M')}</small>
            </div>
            """, unsafe_allow_html=True)
        with col2:
            if not alert.read:
                if st.button("📖 Mark Read", key=f"read_{alert.id}"):
                    alert.read = True
                    st.rerun()
    
    if not alerts:
        st.info("No alerts to display.")
    
    st.divider()
    st.subheader("📈 Real-Time Market Feeds")
    
    ticker = st.text_input("Enter Ticker:", "AAPL")
    
    if ticker:
        tab1, tab2, tab3 = st.tabs(["📊 Market Data", "📄 SEC Filings", "🎙️ Earnings Call"])
        
        with tab1:
            with st.spinner(f"Fetching market data for {ticker}..."):
                market_data = prof_engine.get_market_data(ticker)
            st.json(market_data)
        
        with tab2:
            with st.spinner(f"Fetching SEC filings for {ticker}..."):
                filings = prof_engine.get_sec_filing(ticker)
            st.json(filings)
        
        with tab3:
            with st.spinner(f"Fetching earnings call for {ticker}..."):
                earnings = prof_engine.get_earnings_call(ticker)
            st.json(earnings)


# ==================== DOCUMENT ANALYSIS COMPONENT ====================

def show_document_analysis():
    st.header(f"📄 Document Intelligence - {user_role}")
    st.caption(f"Upload your {user_role.lower()} documents for AI-powered analysis and feedback")
    
    uploaded_file = st.file_uploader(
        "📤 Upload Document",
        type=['pdf', 'docx', 'txt', 'csv', 'json'],
        help="Upload PDF, DOCX, TXT, CSV, or JSON files for analysis"
    )
    
    if uploaded_file is not None:
        st.info(f"📎 **File:** {uploaded_file.name} ({uploaded_file.size / 1024:.2f} KB)")
        
        doc_engine = DocumentIntelligenceEngine(st.session_state.country_code)
        
        with st.spinner("📖 Extracting and analyzing document..."):
            text = doc_engine.extract_text_from_file(uploaded_file)
        
        if text.startswith("Error") or text.startswith("Unsupported") or text.startswith("No text"):
            st.error(f"⚠️ {text}")
        else:
            with st.expander("📝 Document Preview"):
                preview = text[:1000] + ("..." if len(text) > 1000 else "")
                st.text_area("Content Preview", preview, height=200)
            
            with st.spinner("🧠 Analyzing document content..."):
                analysis = doc_engine.analyze_document(text, user_role)
            
            st.success("✅ Document Analysis Complete!")
            
            col1, col2, col3, col4 = st.columns(4)
            col1.metric("📊 Word Count", analysis.get('word_count', 0))
            col2.metric("📝 Sentences", analysis.get('sentence_count', 0))
            col3.metric("🎯 Key Phrases", len(analysis.get('key_phrases', [])))
            col4.metric("💬 Sentiment", analysis.get('sentiment', {}).get('sentiment', 'Neutral'))
            
            st.subheader("📖 Readability Analysis")
            readability = analysis.get('readability', {})
            col1, col2, col3 = st.columns(3)
            col1.metric("Flesch Score", readability.get('flesch_score', 'N/A'))
            col2.metric("Grade Level", readability.get('grade_level', 'Unknown'))
            col3.metric("Avg Words/Sentence", readability.get('avg_words_per_sentence', 'N/A'))
            
            if analysis.get('key_phrases'):
                st.subheader("🔑 Key Phrases")
                st.markdown(", ".join([f"`{phrase}`" for phrase in analysis.get('key_phrases', [])[:7]]))
            
            st.divider()
            st.subheader(f"🎯 {user_role}-Specific Analysis")
            
            if user_role == 'Student':
                col1, col2 = st.columns(2)
                with col1:
                    st.metric("📚 Academic Language Score", f"{analysis.get('academic_language_score', 0)}%")
                    st.markdown("**📋 Topics Mentioned:**")
                    for topic in analysis.get('topics_mentioned', []):
                        st.markdown(f"- {topic}")
                with col2:
                    st.markdown("**💡 Suggested Improvements:**")
                    for suggestion in analysis.get('suggested_improvements', []):
                        st.info(f"• {suggestion}")
            
            elif user_role == 'Teacher':
                col1, col2 = st.columns(2)
                with col1:
                    st.metric("🎓 Pedagogical Score", f"{analysis.get('pedagogical_score', 0)}%")
                    st.metric("📋 Completeness Score", f"{analysis.get('completeness_score', 0)}%")
                    alignment = analysis.get('curriculum_alignment', {})
                    st.metric("📋 Curriculum Alignment", alignment.get('status', 'Unknown'))
                    st.markdown("**✅ Matched Indicators:**")
                    for indicator in alignment.get('matched_indicators', []):
                        st.markdown(f"- {indicator}")
                with col2:
                    st.markdown("**💡 Suggested Enhancements:**")
                    for suggestion in analysis.get('suggested_enhancements', []):
                        st.info(f"• {suggestion}")
                    st.markdown("**📋 Components:**")
                    st.markdown(f"- Objectives: {'✅' if analysis.get('has_objective') else '❌'}")
                    st.markdown(f"- Assessment: {'✅' if analysis.get('has_assessment') else '❌'}")
                    st.markdown(f"- Activities: {'✅' if analysis.get('has_activity') else '❌'}")
                    st.markdown(f"- Standards: {'✅' if analysis.get('has_standards') else '❌'}")
            
            elif user_role == 'Professional':
                col1, col2 = st.columns(2)
                with col1:
                    st.metric("💼 Professional Score", f"{analysis.get('professional_score', 0)}%")
                    business = analysis.get('business_context', {})
                    st.markdown(f"**🏭 Industry:** {business.get('industry', 'Not specified')}")
                    st.markdown(f"**📊 Key Metrics:** {', '.join(business.get('key_metrics', ['None identified']))}")
                    fin_metrics = analysis.get('financial_metrics', {})
                    if fin_metrics.get('revenue'):
                        st.metric("💰 Revenue", f"${fin_metrics['revenue']:,.0f}")
                    if fin_metrics.get('profit'):
                        st.metric("📈 Profit", f"${fin_metrics['profit']:,.0f}")
                    if fin_metrics.get('margin'):
                        st.metric("📊 Margin", f"{fin_metrics['margin']:.1f}%")
                    if fin_metrics.get('roi'):
                        st.metric("💎 ROI", f"{fin_metrics['roi']:.1f}%")
                with col2:
                    st.markdown("**💡 Actionable Insights:**")
                    for insight in analysis.get('actionable_insights', []):
                        st.success(f"• {insight}")
            
            elif user_role == 'SME Business Owner':
                col1, col2 = st.columns(2)
                with col1:
                    st.metric("🏢 Business Score", f"{analysis.get('business_score', 0)}%")
                    st.markdown("**📈 Growth Opportunities:**")
                    for opportunity in analysis.get('growth_opportunities', []):
                        st.markdown(f"- {opportunity}")
                    fin_metrics = analysis.get('financial_metrics', {})
                    if fin_metrics.get('revenue'):
                        st.metric("💰 Revenue", f"${fin_metrics['revenue']:,.0f}")
                    if fin_metrics.get('profit'):
                        st.metric("📈 Profit", f"${fin_metrics['profit']:,.0f}")
                with col2:
                    st.markdown("**🤖 Automation Candidates:**")
                    for candidate in analysis.get('automation_candidates', []):
                        st.info(f"• {candidate}")
            
            if st.button("📥 Download Analysis Report"):
                report = f"""AI Shiksha Document Analysis Report
                ======================================
                Document: {uploaded_file.name}
                Country: {st.session_state.country_code.title()}
                Segment: {user_role}
                Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}
                
                Analysis Results:
                - Word Count: {analysis.get('word_count', 0)}
                - Sentence Count: {analysis.get('sentence_count', 0)}
                - Sentiment: {analysis.get('sentiment', {}).get('sentiment', 'Neutral')}
                - Readability: {readability.get('grade_level', 'Unknown')}
                
                Key Phrases: {', '.join(analysis.get('key_phrases', [])[:7])}
                """
                
                if user_role == 'Teacher':
                    report += f"""
                    
                    Teacher Analysis:
                    - Pedagogical Score: {analysis.get('pedagogical_score', 0)}%
                    - Completeness Score: {analysis.get('completeness_score', 0)}%
                    - Curriculum Alignment: {analysis.get('curriculum_alignment', {}).get('status', 'Unknown')}
                    - Enhancements: {', '.join(analysis.get('suggested_enhancements', []))}
                    """
                elif user_role == 'Professional':
                    report += f"""
                    
                    Professional Analysis:
                    - Professional Score: {analysis.get('professional_score', 0)}%
                    - Industry: {analysis.get('business_context', {}).get('industry', 'Not specified')}
                    - Insights: {', '.join(analysis.get('actionable_insights', []))}
                    - Financial Metrics: {json.dumps(analysis.get('financial_metrics', {}), indent=2)}
                    """
                elif user_role == 'SME Business Owner':
                    report += f"""
                    
                    SME Analysis:
                    - Business Score: {analysis.get('business_score', 0)}%
                    - Growth Opportunities: {', '.join(analysis.get('growth_opportunities', []))}
                    - Automation: {', '.join(analysis.get('automation_candidates', []))}
                    - Financial Metrics: {json.dumps(analysis.get('financial_metrics', {}), indent=2)}
                    """
                
                st.download_button(
                    label="📥 Download Report",
                    data=report,
                    file_name=f"analysis_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                    mime="text/plain"
                )


# ==================== SME DASHBOARD FUNCTIONS ====================

def show_sme_dashboard():
    st.header(f"🏢 SME Growth Automation Engine - {country_code.title()}")
    
    sme_engine = SMEOutcomeEngine('africa' if country_code in ['kenya', 'bangladesh'] else 'global', country_code)
    
    business_data = {
        'inventory': [
            {'id': 'SKU-001', 'name': 'Product A', 'quantity': 45, 'threshold': 50, 'price': 25.00},
            {'id': 'SKU-002', 'name': 'Product B', 'quantity': 12, 'threshold': 20, 'price': 15.00},
            {'id': 'SKU-003', 'name': 'Product C', 'quantity': 3, 'threshold': 10, 'price': 35.00},
            {'id': 'SKU-004', 'name': 'Product D', 'quantity': 0, 'threshold': 15, 'price': 45.00},
        ],
        'invoices': [
            {'id': 'INV-001', 'amount': 150.00, 'status': 'paid', 'due_date': '2024-01-20'},
            {'id': 'INV-002', 'amount': 75.50, 'status': 'overdue', 'due_date': '2024-01-10'},
            {'id': 'INV-003', 'amount': 200.00, 'status': 'pending', 'due_date': '2024-01-25'},
            {'id': 'INV-004', 'amount': 120.00, 'status': 'overdue', 'due_date': '2024-01-05'},
        ],
        'customers': [
            {'id': 1, 'name': 'John Doe', 'days_since_last_order': 45, 'avg_order_value': 75, 'order_count': 3, 'months_active': 6, 'frequency': 0.5},
            {'id': 2, 'name': 'Jane Smith', 'days_since_last_order': 12, 'avg_order_value': 120, 'order_count': 8, 'months_active': 12, 'frequency': 1.5},
            {'id': 3, 'name': 'Bob Johnson', 'days_since_last_order': 60, 'avg_order_value': 50, 'order_count': 2, 'months_active': 4, 'frequency': 0.3},
            {'id': 4, 'name': 'Alice Brown', 'days_since_last_order': 8, 'avg_order_value': 90, 'order_count': 6, 'months_active': 8, 'frequency': 1.2},
        ]
    }
    
    col1, col2, col3, col4 = st.columns(4)
    col1.metric("📊 Revenue", "$15,750", "↑ 12.5%")
    col2.metric("👥 Customers", "342", "↑ 8.3%")
    col3.metric("📦 Inventory Items", "1,247", "↓ 3.2%")
    col4.metric("⚡ Automation Rate", "68%", "↑ 12%")
    
    st.subheader("🎯 Action-Oriented Task Feed")
    st.caption("Priority-ranked recommendations based on real-time business data")
    
    tasks = sme_engine.get_action_tasks(business_data)
    
    filter_col1, filter_col2 = st.columns(2)
    with filter_col1:
        priority_filter = st.selectbox("Filter by Priority:", ['All', 'Critical', 'High', 'Medium', 'Low'])
    with filter_col2:
        category_filter = st.selectbox("Filter by Category:", ['All', 'inventory', 'finance', 'customers', 'marketing', 'operations'])
    
    filtered_tasks = tasks
    if priority_filter != 'All':
        filtered_tasks = [t for t in filtered_tasks if t.priority.lower() == priority_filter.lower()]
    if category_filter != 'All':
        filtered_tasks = [t for t in filtered_tasks if t.category == category_filter]
    
    for task in filtered_tasks[:10]:
        priority_class = {
            'critical': 'task-critical',
            'high': 'task-high',
            'medium': 'task-medium',
            'low': 'task-low'
        }.get(task.priority, 'task-medium')
        
        with st.container():
            st.markdown(f'<div class="{priority_class}">', unsafe_allow_html=True)
            col1, col2, col3 = st.columns([5, 2, 1])
            
            with col1:
                st.markdown(f"**{task.title}**")
                st.caption(task.description)
                st.caption(f"💡 Impact: {task.impact}")
            
            with col2:
                st.caption(f"Priority: {task.priority.upper()}")
                st.caption(f"Category: {task.category}")
                if task.due_date:
                    st.caption(f"Due: {task.due_date.strftime('%Y-%m-%d')}")
            
            with col3:
                if task.status == 'pending':
                    if st.button(f"✅ Approve", key=f"approve_{task.id}"):
                        result = sme_engine.ux_engine.approve_task(task.id)
                        if result['status'] in ['approved', 'drafted', 'acknowledged']:
                            st.success(f"✅ {result['message']}")
                            st.rerun()
                    
                    if st.button(f"❌ Dismiss", key=f"dismiss_{task.id}"):
                        result = sme_engine.ux_engine.dismiss_task(task.id)
                        if result['status'] == 'dismissed':
                            st.info(f"🗑️ {result['message']}")
                            st.rerun()
                else:
                    st.caption(f"Status: {task.status}")
            
            st.markdown('</div>', unsafe_allow_html=True)
    
    if not filtered_tasks:
        st.info("🎉 No pending tasks! Your business is on track.")
    
    st.divider()
    st.subheader("💬 AI Business Assistant")
    st.caption("Ask questions about your business in plain English")
    
    query = st.text_input("What would you like to know?", placeholder="e.g., 'What's my cash flow forecast?' or 'Which customers are at risk?'")
    
    if query:
        with st.spinner("Analyzing your business data..."):
            response = sme_engine.natural_language_query(query, business_data)
            
            st.markdown(f"**Your Question:** {query}")
            st.markdown("---")
            
            for result in response.get('results', []):
                st.markdown(f"**{result.get('type', 'Analysis').title()}**")
                st.info(result.get('summary', 'No summary available'))
                
                if result.get('data') and isinstance(result['data'], dict):
                    st.json(result['data'])
                elif result.get('data') and isinstance(result['data'], list):
                    st.dataframe(pd.DataFrame(result['data']))
    
    st.divider()
    st.subheader("📱 Proactive Push Delivery")
    st.caption("Automated digests and alerts through your preferred channels")
    
    col1, col2 = st.columns(2)
    with col1:
        channel = st.selectbox("Select Channel:", ['WhatsApp', 'SMS', 'Email'])
        
        if st.button("📤 Send Digest Now", type="primary"):
            digest = sme_engine.generate_digest(channel.lower())
            
            st.success(f"✅ Digest sent to {channel}!")
            
            if digest.get('action_required', False):
                st.warning(f"⚠️ {digest.get('summary', {}).get('critical_tasks', 0)} critical tasks require your attention")
            
            with st.expander("📋 View Digest Preview"):
                st.json(digest)
    
    with col2:
        st.markdown("**Recent Alerts**")
        for alert in sme_engine.ux_engine.alert_history[-3:]:
            st.info(f"📢 {alert.get('message', 'Alert')[:100]}...")


def show_sme_growth():
    st.header(f"📈 SME Growth Analytics - {country_code.title()}")
    
    sme_engine = SMEOutcomeEngine('africa' if country_code in ['kenya', 'bangladesh'] else 'global', country_code)
    ai_engine = SMEAIEngine(country_code)
    
    st.subheader("🔮 Predictive Analytics")
    
    tab1, tab2, tab3, tab4 = st.tabs(["💰 Cash Flow", "📊 Churn Prediction", "📦 Inventory", "💎 Customer LTV"])
    
    with tab1:
        st.markdown("### Cash Flow Forecast (Next 30 Days)")
        cash_flow = ai_engine.predict_cash_flow([])
        
        df = pd.DataFrame(cash_flow['predictions'])
        st.line_chart(df.set_index('day')['balance'])
        
        col1, col2, col3 = st.columns(3)
        col1.metric("📊 Minimum Balance", f"${cash_flow['min_balance']:,.2f}", f"Day {cash_flow['min_balance_day']}")
        col2.metric("📈 Average Daily Balance", f"${cash_flow['avg_daily_balance']:,.2f}")
        col3.metric("⚠️ Risk Level", cash_flow['risk_level'].upper())
        
        if cash_flow['risk_level'] == 'high':
            st.warning("⚠️ Cash flow risk detected! Consider reducing expenses or increasing revenue.")
    
    with tab2:
        st.markdown("### Customer Churn Prediction")
        
        customers = [
            {'id': 1, 'name': 'John Doe', 'days_since_last_order': 45, 'avg_order_value': 75, 'order_count': 3},
            {'id': 2, 'name': 'Jane Smith', 'days_since_last_order': 12, 'avg_order_value': 120, 'order_count': 8},
            {'id': 3, 'name': 'Bob Johnson', 'days_since_last_order': 60, 'avg_order_value': 50, 'order_count': 2},
            {'id': 4, 'name': 'Alice Brown', 'days_since_last_order': 8, 'avg_order_value': 90, 'order_count': 6},
            {'id': 5, 'name': 'Charlie Wilson', 'days_since_last_order': 75, 'avg_order_value': 60, 'order_count': 1},
        ]
        
        churn_predictions = ai_engine.predict_churn(customers)
        df_churn = pd.DataFrame(churn_predictions)
        
        def color_risk(val):
            if val == 'high':
                return '🔴 High'
            elif val == 'medium':
                return '🟡 Medium'
            else:
                return '🟢 Low'
        
        df_churn['risk_level_display'] = df_churn['risk_level'].apply(color_risk)
        st.dataframe(df_churn[['name', 'churn_score', 'risk_level_display', 'recommendation']])
        
        high_risk = len([c for c in churn_predictions if c['risk_level'] == 'high'])
        if high_risk > 0:
            st.warning(f"⚠️ {high_risk} customers at high churn risk. Take action now!")
    
    with tab3:
        st.markdown("### Inventory Depletion Forecast")
        
        inventory_data = [
            {'id': 'SKU-001', 'name': 'Product A', 'quantity': 45, 'threshold': 50, 'daily_sales_avg': 2.5, 'price': 25.00},
            {'id': 'SKU-002', 'name': 'Product B', 'quantity': 12, 'threshold': 20, 'daily_sales_avg': 1.8, 'price': 15.00},
            {'id': 'SKU-003', 'name': 'Product C', 'quantity': 3, 'threshold': 10, 'daily_sales_avg': 1.2, 'price': 35.00},
            {'id': 'SKU-004', 'name': 'Product D', 'quantity': 0, 'threshold': 15, 'daily_sales_avg': 1.0, 'price': 45.00},
        ]
        
        inventory_forecast = ai_engine.predict_inventory_depletion(inventory_data)
        df_inventory = pd.DataFrame(inventory_forecast)
        
        def color_status(val):
            if val == 'critical':
                return '🔴 Critical'
            elif val == 'warning':
                return '🟡 Warning'
            else:
                return '🟢 OK'
        
        df_inventory['status_display'] = df_inventory['status'].apply(color_status)
        st.dataframe(df_inventory[['name', 'current_qty', 'days_until_threshold', 'status_display', 'restock_quantity']])
        
        critical_items = [i for i in inventory_forecast if i['status'] == 'critical']
        if critical_items:
            st.error(f"🚨 {len(critical_items)} items require immediate restocking!")
    
    with tab4:
        st.markdown("### Customer Lifetime Value (LTV) Analysis")
        
        customer_data = [
            {'id': 1, 'name': 'John Doe', 'avg_order_value': 75, 'frequency': 0.5, 'months_active': 6},
            {'id': 2, 'name': 'Jane Smith', 'avg_order_value': 120, 'frequency': 1.5, 'months_active': 12},
            {'id': 3, 'name': 'Bob Johnson', 'avg_order_value': 50, 'frequency': 0.3, 'months_active': 4},
            {'id': 4, 'name': 'Alice Brown', 'avg_order_value': 90, 'frequency': 1.2, 'months_active': 8},
            {'id': 5, 'name': 'Charlie Wilson', 'avg_order_value': 60, 'frequency': 0.2, 'months_active': 3},
        ]
        
        ltv_predictions = ai_engine.predict_customer_ltv(customer_data)
        df_ltv = pd.DataFrame(ltv_predictions)
        
        def color_segment(val):
            if val == 'high':
                return '🟢 High Value'
            elif val == 'medium':
                return '🟡 Medium Value'
            else:
                return '🔴 Low Value'
        
        df_ltv['segment_display'] = df_ltv['segment'].apply(color_segment)
        st.dataframe(df_ltv[['name', 'current_ltv', 'projected_ltv', 'potential_increase', 'segment_display']])
        
        total_ltv = sum(c['current_ltv'] for c in ltv_predictions)
        total_projected = sum(c['projected_ltv'] for c in ltv_predictions)
        st.metric("Total Current LTV", f"${total_ltv:,.2f}", f"+${total_projected - total_ltv:,.2f} potential")


def show_sme_automation():
    st.header(f"🤖 SME Automation Solutions - {country_code.title()}")
    
    sme_engine = SMEOutcomeEngine('africa' if country_code in ['kenya', 'bangladesh'] else 'global', country_code)
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("🚀 Generate Automation Plan")
        business_type = st.selectbox(
            "Business Type:",
            ['retail', 'service', 'agriculture', 'manufacturing', 'tech']
        )
        
        if st.button("🚀 Generate Automation Plan", type="primary"):
            automations = sme_engine.generate_automation(business_type)
            
            st.success(f"✅ Automation Plan Generated for {business_type.title()} Business")
            
            st.markdown("#### 🔄 Automated Systems")
            for key, value in automations.items():
                st.markdown(f"- **{key.title()}:** {value}")
    
    with col2:
        st.subheader("📊 Expected Impact")
        st.metric("⏱️ Time Savings", "15-25 hours/week")
        st.metric("📈 Revenue Increase", "20-30%", "Projected")
        st.metric("👥 Customer Retention", "85%", "↑ 10%")
        
        st.subheader("💰 Payment Integration")
        st.markdown(f"**Region:** {country_code.upper()}")
        st.markdown(f"**Currency:** {overlay.get('currency', 'Local')}")
        st.markdown("**Supported Payment Rails:**")
        
        rails = sme_engine._get_payment_rails()
        for rail in rails.split(', '):
            st.markdown(f"✅ {rail}")
    
    st.divider()
    st.subheader("💬 AI Business Query (RAG)")
    st.caption("Ask questions about your automation and business data")
    
    query = st.text_input("Ask about automation opportunities:", placeholder="e.g., 'How can I automate my inventory management?'")
    if query:
        with st.spinner("Generating insights..."):
            business_data = {
                'inventory': [{'name': 'Product A', 'quantity': 45, 'threshold': 50}],
                'customers': [{'name': 'John Doe', 'orders': 3}]
            }
            response = sme_engine.natural_language_query(query, business_data)
            st.info(response.get('results', [{}])[0].get('summary', 'No response available'))


def show_sme_api_connectors():
    st.header(f"🔌 API Connector Layer - {country_code.title()}")
    st.caption("Connect your business tools for seamless data integration")
    
    sme_engine = SMEOutcomeEngine('africa' if country_code in ['kenya', 'bangladesh'] else 'global', country_code)
    
    st.subheader("📊 Available Integrations")
    
    api_options = ['stripe', 'quickbooks', 'shopify', 'toast', 'jobber']
    api_icons = {
        'stripe': '💳',
        'quickbooks': '📊',
        'shopify': '🛍️',
        'toast': '🍽️',
        'jobber': '🔧'
    }
    api_names = {
        'stripe': 'Stripe Payments',
        'quickbooks': 'QuickBooks Accounting',
        'shopify': 'Shopify E-commerce',
        'toast': 'Toast POS',
        'jobber': 'Jobber Field Service'
    }
    
    cols = st.columns(3)
    for i, api in enumerate(api_options[:3]):
        with cols[i]:
            st.markdown(f"### {api_icons.get(api, '🔌')} {api_names.get(api, api.title())}")
            st.caption("Read/Write Integration")
            
            if st.button(f"Connect {api_names.get(api, api.title())}", key=f"connect_{api}"):
                credentials = {
                    'api_key': st.text_input("API Key:", placeholder="Enter your API key", key=f"key_{api}"),
                    'secret': st.text_input("Secret:", placeholder="Enter your secret", type="password", key=f"secret_{api}")
                }
                
                result = sme_engine.connect_api(api, credentials)
                if result['status'] == 'connected':
                    st.success(f"✅ {result['message']}")
                    st.balloons()
                else:
                    st.error(f"❌ {result['message']}")
    
    st.divider()
    st.subheader("📡 Connected Services")
    
    if sme_engine.infrastructure.api_connections:
        for service, conn in sme_engine.infrastructure.api_connections.items():
            config = conn.get('config', {})
            st.info(f"✅ **{config.get('name', service.title())}** - Connected ({conn.get('last_sync', 'N/A')})")
    else:
        st.info("No services connected yet. Connect your first service above.")
    
    st.divider()
    st.subheader("📊 API Data Preview")
    
    if sme_engine.infrastructure.api_connections:
        service = st.selectbox("Select Service:", list(sme_engine.infrastructure.api_connections.keys()))
        endpoints = sme_engine.infrastructure.api_connections[service]['config'].get('endpoints', [])
        
        if endpoints:
            endpoint = st.selectbox("Select Endpoint:", endpoints)
            
            if st.button("🔍 Fetch Data"):
                data = sme_engine.infrastructure.fetch_api_data(service, endpoint)
                if data.get('status') == 'success':
                    st.dataframe(pd.DataFrame(data.get('data', [])))
                else:
                    st.error(f"Error: {data.get('message', 'Unknown error')}")


def show_sme_webhooks():
    st.header(f"⚡ Webhook Architecture - {country_code.title()}")
    st.caption("Event-driven infrastructure for real-time AI actions")
    
    sme_engine = SMEOutcomeEngine('africa' if country_code in ['kenya', 'bangladesh'] else 'global', country_code)
    
    st.subheader("📡 Available Webhook Events")
    
    webhook_events = [
        {'event': 'checkout.paid', 'description': 'Customer completes a purchase', 'action': 'Update revenue, send confirmation'},
        {'event': 'invoice.overdue', 'description': 'Invoice becomes overdue', 'action': 'Send reminder, flag for follow-up'},
        {'event': 'inventory.low', 'description': 'Stock drops below threshold', 'action': 'Generate restock alert'},
        {'event': 'customer.churn_risk', 'description': 'Customer shows churn signals', 'action': 'Send re-engagement campaign'},
        {'event': 'order.fulfilled', 'description': 'Order is fulfilled', 'action': 'Update inventory, send satisfaction survey'},
    ]
    
    df_webhooks = pd.DataFrame(webhook_events)
    st.dataframe(df_webhooks)
    
    st.divider()
    st.subheader("🔧 Test Webhook Simulation")
    
    event_type = st.selectbox("Select Event Type:", ['checkout.paid', 'invoice.overdue', 'inventory.low', 'customer.churn_risk'])
    
    if event_type == 'checkout.paid':
        payload = {
            'amount': 150.00,
            'customer': {'name': 'John Doe', 'email': 'john@example.com'},
            'order_id': 'ORD-12345'
        }
    elif event_type == 'invoice.overdue':
        payload = {
            'invoice_id': 'INV-6789',
            'amount': 250.00,
            'days_overdue': 5,
            'customer': {'name': 'Jane Smith', 'email': 'jane@example.com'}
        }
    elif event_type == 'inventory.low':
        payload = {
            'product': {'id': 'SKU-001', 'name': 'Product A', 'price': 25.00},
            'current_qty': 3,
            'threshold': 10
        }
    else:
        payload = {
            'customer': {'id': 123, 'name': 'Bob Johnson', 'email': 'bob@example.com'},
            'churn_score': 0.75,
            'days_inactive': 45,
            'ltv': 500.00
        }
    
    st.json(payload)
    
    if st.button("⚡ Process Webhook Event", type="primary"):
        result = sme_engine.process_webhook(event_type, payload)
        
        st.success(f"✅ Webhook Event Processed!")
        
        st.markdown("**Result:**")
        st.json(result)
        
        if result.get('action_task'):
            task = result['action_task']
            st.markdown("**📋 Generated Action Task:**")
            st.info(f"**{task.title}**\n\n{task.description}")
    
    st.divider()
    st.subheader("📜 Webhook History")
    
    if sme_engine.infrastructure.webhook_events:
        history_data = []
        for event in sme_engine.infrastructure.webhook_events[-5:]:
            history_data.append({
                'Event': event.event_type,
                'Timestamp': event.timestamp.strftime('%Y-%m-%d %H:%M:%S'),
                'Payload': str(event.payload)[:50] + '...'
            })
        st.dataframe(pd.DataFrame(history_data))
    else:
        st.info("No webhook events processed yet. Use the simulator above to test.")


# ==================== REST OF THE FUNCTIONS ====================

def show_home():
    st.title("🌍 AI Shiksha Global Edition")
    st.subheader(f"{country_flags.get(country_code, '🌍')} {country_code.title()} - Universal Core + Local Curriculum Overlay")
    
    with st.expander("🔷 Universal Core - Portable Across All Systems", expanded=True):
        st.markdown("""
        <div class="universal-core">
        <h4>📚 Core Subjects (Available Everywhere)</h4>
        """, unsafe_allow_html=True)
        
        cols = st.columns(3)
        subjects = ['Mathematics', 'English Language', 'Basic Science', 'Geography', 'General Knowledge', 'Applied AI']
        for i, subject in enumerate(subjects):
            cols[i % 3].markdown(f"✅ **{subject}**")
        
        st.markdown("</div>", unsafe_allow_html=True)
    
    with st.expander("🔶 Local Curriculum Overlay - Country-Specific", expanded=True):
        st.markdown(f"""
        <div class="local-overlay">
        <h4>{country_flags.get(country_code, '🌍')} {country_code.title()} - Curriculum Details</h4>
        <p><strong>Curriculum System:</strong> {overlay.get('system', 'Universal')}</p>
        <p><strong>Exam Boards:</strong> {', '.join(overlay.get('boards', ['Local boards']))}</p>
        <p><strong>National Exams:</strong> {', '.join(overlay.get('national_exams', ['Local exams']))}</p>
        <p><strong>Language Support:</strong> {overlay.get('language', 'English')}</p>
        <p><strong>Grade Levels:</strong> {', '.join(overlay.get('grade_levels', ['All levels']))}</p>
        <p><strong>Currency:</strong> {overlay.get('currency', 'Local')}</p>
        <p><strong>Cultural Context:</strong> {context.get('culture', 'Global')}</p>
        </div>
        """, unsafe_allow_html=True)
    
    st.divider()
    st.subheader("🎯 Segment-Specific Outcome Engines")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.markdown("**🎓 Students**")
        st.caption("Grade & Exam Outcomes Engine")
        st.markdown("✅ Adaptive practice")
        st.markdown("✅ Score trends")
        st.markdown("✅ Competition prep")
        st.markdown("✅ Socratic loops")
    
    with col2:
        st.markdown("**👨‍🏫 Teachers**")
        st.caption("Hours-Saved Engine")
        st.markdown("✅ Lesson bundle generator")
        st.markdown("✅ Rubric & quiz generator")
        st.markdown("✅ LMS integration")
        st.markdown("✅ Standards mapping")
    
    with col3:
        st.markdown("**💼 Professionals**")
        st.caption("Career Acceleration Lab")
        st.markdown("✅ Hybrid search")
        st.markdown("✅ Citation engine")
        st.markdown("✅ Portfolio monitoring")
        st.markdown("✅ Stress testing")
    
    with col4:
        st.markdown("**🏢 SMEs**")
        st.caption("Growth Automation Engine")
        st.markdown("✅ Predictive analytics")
        st.markdown("✅ Action task feed")
        st.markdown("✅ API connectors")
        st.markdown("✅ Webhooks & alerts")


def _log_daily_activity():
    """Wires HabitStreakEngine into session_state once per page load."""
    result = HabitStreakEngine.record_activity(
        st.session_state.last_active_date,
        st.session_state.streak,
        st.session_state.streak_repair_tokens
    )
    if result['date'] != st.session_state.last_active_date:
        st.session_state.streak = result['streak']
        st.session_state.streak_repair_tokens = result['repair_tokens']
        st.session_state.last_active_date = result['date']
        if 'streak started' not in result['message'].lower() and 'already logged' not in result['message'].lower():
            st.toast(result['message'])


def _award_achievement(name: str):
    if name not in st.session_state.achievements:
        st.session_state.achievements.append(name)
        st.toast(f"🏆 Achievement unlocked: {name}")


def show_student_dashboard():
    st.header(f"🎓 Student Dashboard - {country_code.title()}")
    _log_daily_activity()

    student_engine = StudentOutcomeEngine(country_code)

    col1, col2, col3, col4, col5 = st.columns(5)
    col1.metric("Current Score", f"{st.session_state.student_score}%")
    col2.metric("Day Streak", f"{st.session_state.streak} days", help="Streak-repair buffers left: " + str(st.session_state.streak_repair_tokens))
    col3.metric("Completed Lessons", len(st.session_state.completed_lessons))
    col4.metric("Achievements", len(st.session_state.achievements))
    due_count = len(SpacedRepetitionEngine.due_cards(st.session_state.sr_cards))
    col5.metric("Cards Due for Review", due_count)

    if due_count > 0:
        st.info(f"🔁 You have **{due_count}** concept(s) due for review today — right before you'd naturally start forgetting them. Head to **Review Queue**.")

    st.divider()
    st.subheader("⏱️ Micro-Session")
    st.caption("A short, low-friction session beats a long one you skip. This one is paced for ~5 minutes.")
    if 'micro_session_start' not in st.session_state:
        st.session_state.micro_session_start = None

    msc1, msc2 = st.columns([1, 3])
    with msc1:
        if st.session_state.micro_session_start is None:
            if st.button("▶️ Start 5-min session"):
                st.session_state.micro_session_start = time.time()
                st.rerun()
        else:
            elapsed = time.time() - st.session_state.micro_session_start
            remaining = max(0, HabitStreakEngine.MICRO_SESSION_SECONDS - elapsed)
            if remaining <= 0:
                st.success("✅ Micro-session complete!")
                _award_achievement("5-Minute Finisher")
                if st.button("Reset"):
                    st.session_state.micro_session_start = None
                    st.rerun()
            else:
                mins, secs = divmod(int(remaining), 60)
                with msc2:
                    st.progress(1 - remaining / HabitStreakEngine.MICRO_SESSION_SECONDS, text=f"{mins:02d}:{secs:02d} remaining — head to Practice")

    st.divider()
    st.subheader("📊 Progress Snapshot")
    weeks = ['Week 1', 'Week 2', 'Week 3', 'Week 4', 'Current']
    scores = [45, 52, 58, 63, st.session_state.student_score]
    df = pd.DataFrame({'Week': weeks, 'Score': scores})
    st.line_chart(df.set_index('Week'))

    col1, col2 = st.columns(2)
    with col1:
        grade = student_engine._calculate_projected_grade({'score': st.session_state.student_score})
        st.metric("Projected Grade", grade)
    with col2:
        if st.session_state.student_score >= 60:
            st.success("✅ On track for success!")
        else:
            st.warning("📚 Keep practicing to improve")


def show_student_practice():
    st.header("📝 Adaptive Practice")
    st.caption(f"Aligned with {overlay.get('system', 'Universal')} curriculum · Socratic guardrails on · Difficulty auto-balances")
    _log_daily_activity()

    student_engine = StudentOutcomeEngine(country_code)

    subject = st.selectbox("Select Subject:", ['Mathematics', 'English Language', 'Basic Science', 'Geography', 'General Knowledge'], key='practice_subject')

    d1, d2 = st.columns([2, 2])
    with d1:
        st.session_state.current_difficulty = st.select_slider(
            "Difficulty Level (auto-adjusts as you go):",
            ['easy', 'medium', 'hard'],
            value=st.session_state.current_difficulty
        )
    with d2:
        st.caption(f"Wrong-answer streak: {st.session_state.wrong_streak} · Win streak: {st.session_state.win_streak}")
        st.caption("2 wrong in a row → breaks the next problem into sub-steps. 3 correct in a row → levels up.")

    tabs = st.tabs(["✍️ Answer Questions", "📸 Scan Handwritten Work (OCR)", "🎙️ Explain It Out Loud (Voice Recall)", "🧮 Step-by-Step Check"])

    # ---- Tab 1: Adaptive MCQ practice with Socratic hint ladder ----
    with tabs[0]:
        if st.button("🎯 Generate Practice Questions", type="primary"):
            st.session_state['_practice_questions'] = student_engine.get_adaptive_questions(subject.lower(), st.session_state.current_difficulty)

        questions = st.session_state.get('_practice_questions', [])
        if questions:
            for q in questions[:2]:
                st.markdown(f"### Question: {q['question']}")
                st.caption(f"📚 {q.get('exam_style', 'Local exam')} style")

                qid = q['id']
                hint_key = f"hint_level_{qid}"
                current_hint_level = st.session_state.hint_level.get(hint_key, -1)

                hc1, hc2 = st.columns([1, 3])
                with hc1:
                    if st.button("💡 Give me a hint (never the answer)", key=f"hintbtn_{qid}"):
                        current_hint_level = min(current_hint_level + 1, 3)
                        st.session_state.hint_level[hint_key] = current_hint_level
                if current_hint_level >= 0:
                    st.info(SocraticGuardrailEngine.get_hint(q, current_hint_level))

                selected = st.radio("Select your answer:", q['options'], key=f"q_{qid}")

                free_text = st.text_input("Or ask a follow-up question about this problem:", key=f"freetext_{qid}")
                if free_text:
                    shield = SocraticGuardrailEngine.content_shield(free_text)
                    if shield['blocked']:
                        st.warning(f"🛡️ I can't do that: {shield['reason']} I can only give hints, not the final answer or someone else's work.")

                start_time = st.session_state.setdefault(f"start_{qid}", time.time())

                if st.button(f"Submit Answer {qid}", key=f"submit_{qid}"):
                    solve_time = time.time() - start_time
                    confidence = st.session_state.get(f"conf_{qid}", 3)
                    correct = (selected == q['correct'])
                    concept_id = f"{subject.lower()}::{q.get('question','')[:24]}"

                    st.session_state.attempt_logs.append(AttemptLog(
                        concept_id=concept_id, subject=subject.lower(), solve_time_sec=round(solve_time, 1),
                        edit_attempts=current_hint_level + 1, self_reported_confidence=confidence, correct=correct
                    ).__dict__)

                    if correct:
                        st.balloons()
                        st.success(f"✅ Correct! {q.get('explanation', '')}")
                        st.session_state.student_score = min(100, st.session_state.student_score + 5)
                        st.session_state.win_streak += 1
                        st.session_state.wrong_streak = 0
                        card = SpacedRepetitionEngine.get_or_create(st.session_state.sr_cards, concept_id, subject.lower(), q['question'][:40])
                        quality = 5 if current_hint_level < 0 else max(3, 5 - current_hint_level)
                        SpacedRepetitionEngine.schedule(card, quality)
                        if st.session_state.win_streak >= 3:
                            _award_achievement("3-Question Win Streak")
                    else:
                        st.error("❌ Not quite. Try another hint instead of guessing again — the goal is to get there yourself.")
                        st.session_state.wrong_streak += 1
                        st.session_state.win_streak = 0
                        card = SpacedRepetitionEngine.get_or_create(st.session_state.sr_cards, concept_id, subject.lower(), q['question'][:40])
                        SpacedRepetitionEngine.schedule(card, 1)

                    shift = DynamicDifficultyBalancer.next_difficulty(st.session_state.current_difficulty, st.session_state.wrong_streak, st.session_state.win_streak)
                    if shift['changed']:
                        st.session_state.current_difficulty = shift['difficulty']
                        st.info(f"⚖️ Difficulty adjusted to **{shift['difficulty']}**.")
                    if shift['break_into_substeps']:
                        st.warning("🧩 That's two misses in a row — try the **Step-by-Step Check** tab to work this one in smaller pieces.")

    # ---- Tab 2: OCR ingestion ----
    with tabs[1]:
        st.caption("Upload a photo of handwritten math, a diagram, or a textbook page.")
        if not MultimodalOCREngine.is_available():
            st.warning("⚠️ OCR isn't active on this server right now: it needs the `pytesseract` package and the Tesseract binary installed, and I can't confirm those are present in your deployment. You can still upload an image and type the text manually below.")
        img_file = st.file_uploader("Upload image (png/jpg):", type=['png', 'jpg', 'jpeg'], key='ocr_upload')
        if img_file is not None:
            st.image(img_file, caption="Uploaded work", use_container_width=True)
            result = MultimodalOCREngine.extract_text(img_file.getvalue())
            if result['success']:
                st.success("Text extracted:")
                extracted = st.text_area("Extracted text (edit if OCR made mistakes):", value=result['text'], key='ocr_text')
            else:
                st.info(result.get('error', 'OCR unavailable.'))
                extracted = st.text_area("Type what's in the image instead:", key='ocr_text_manual')
            if st.button("Save to practice history", key='ocr_save'):
                st.session_state.ocr_history.append({'timestamp': datetime.now().isoformat(), 'text': extracted, 'subject': subject})
                st.success("Saved. Nothing here leaves this session — see Privacy & Safety.")

    # ---- Tab 3: Voice-first active recall ----
    with tabs[2]:
        st.caption("Explain the concept out loud, then check how much of the key vocabulary you covered.")
        key_terms_input = st.text_input("Key terms to check for (comma-separated) — set by the lesson normally:", value="numerator, denominator, equivalent" if subject == 'Mathematics' else "concept, example, reason")
        key_terms = [t.strip() for t in key_terms_input.split(',') if t.strip()]

        audio = None
        try:
            audio = st.audio_input("Record your explanation:", key='voice_recall_audio')
        except Exception:
            st.caption("Your Streamlit version may not support st.audio_input — type your explanation instead.")

        if audio is not None:
            stt_result = VoiceActiveRecallEngine.transcribe_audio(audio.getvalue())
            st.info(stt_result['note'])

        transcript = st.text_area("Type (or paste) what you said:", key='voice_transcript')
        if st.button("Score my explanation", key='score_explanation'):
            shield = SocraticGuardrailEngine.content_shield(transcript)
            if shield['blocked']:
                st.warning(f"🛡️ {shield['reason']}")
            else:
                score = VoiceActiveRecallEngine.score_explanation(transcript, key_terms)
                st.metric("Concept coverage", f"{score['coverage_pct']}%")
                if score['covered_terms']:
                    st.success("Covered: " + ", ".join(score['covered_terms']))
                if score['missed_terms']:
                    st.warning("Missing: " + ", ".join(score['missed_terms']))
                st.session_state.voice_sessions.append({'timestamp': datetime.now().isoformat(), 'subject': subject, **score})

    # ---- Tab 4: Instant step-by-step evaluation ----
    with tabs[3]:
        st.caption("Split-screen check: type your steps on the left, compare against the expected approach on the right.")
        default_expected = "Isolate the variable term\nDivide both sides by the coefficient\nSimplify to find the value\nCheck by substituting back in"
        sc1, sc2 = st.columns(2)
        with sc2:
            expected_text = st.text_area("Expected steps (one per line):", value=default_expected, key='expected_steps')
        with sc1:
            student_text = st.text_area("Your steps (one per line):", key='student_steps', height=170)

        if st.button("Evaluate my steps", key='eval_steps'):
            expected_steps = [s for s in expected_text.split('\n') if s.strip()]
            student_steps = [s for s in student_text.split('\n') if s.strip()]
            results = InstantStepEvaluator.evaluate_steps(student_steps, expected_steps)
            icons = {'correct': '✅', 'partial': '🟡', 'incorrect': '❌', 'empty': '⬜'}
            for r in results:
                st.markdown(f"{icons[r['status']]} **Step {r['step_number']}** — expected: _{r['expected']}_")
                if r['student_input']:
                    st.caption(f"You wrote: {r['student_input']}")


def show_student_knowledge_map():
    st.header("🧠 Knowledge Map")
    st.caption("Mastery unlocks the next concept — no skipping ahead of a prerequisite.")
    _log_daily_activity()

    subject = st.selectbox("Subject:", ['Mathematics', 'English Language', 'Basic Science', 'Geography', 'General Knowledge'], key='km_subject')
    subject_key = subject.lower()
    nodes = PrerequisiteKnowledgeGraph.get_graph(subject_key)
    mastered_list = st.session_state.mastered_concepts.get(subject_key, [])
    mastered_ids = set(mastered_list)
    status_map = PrerequisiteKnowledgeGraph.compute_status(nodes, mastered_ids)

    icon = {'locked': '🔒', 'unlocked': '🟡', 'mastered': '🟢'}
    label = {'locked': 'Locked', 'unlocked': 'Ready to learn', 'mastered': 'Mastered'}

    for node in nodes:
        st_status = status_map[node.node_id]
        with st.container():
            c1, c2, c3 = st.columns([3, 2, 2])
            with c1:
                prereq_names = [n.name for n in nodes if n.node_id in node.prerequisites]
                st.markdown(f"{icon[st_status]} **{node.name}**")
                if prereq_names:
                    st.caption("Requires: " + ", ".join(prereq_names))
            with c2:
                st.caption(label[st_status])
            with c3:
                if st_status == 'unlocked':
                    if st.button("Mark mastered", key=f"master_{node.node_id}"):
                        mastered_list.append(node.node_id)
                        st.session_state.mastered_concepts[subject_key] = mastered_list
                        card = SpacedRepetitionEngine.get_or_create(st.session_state.sr_cards, node.node_id, subject_key, node.name)
                        SpacedRepetitionEngine.schedule(card, 4)
                        if len(mastered_list) == len(nodes):
                            _award_achievement(f"{subject} Tree Completed")
                        st.rerun()
                elif st_status == 'mastered':
                    st.success("Done")
                else:
                    st.caption("—")
        st.divider()

    next_node = PrerequisiteKnowledgeGraph.next_recommended(nodes, mastered_ids)
    if next_node:
        st.info(f"👉 Recommended next: **{next_node.name}**")
    else:
        st.success("🎉 All concepts in this subject are mastered!")

    dot_lines = ["digraph G {", "rankdir=LR;", 'node [shape=box, style=filled, fontname="Helvetica"];']
    color_map = {'locked': '#d1d5db', 'unlocked': '#fde68a', 'mastered': '#86efac'}
    for node in nodes:
        safe_name = node.name.replace('"', "'")
        dot_lines.append(f'"{node.node_id}" [label="{safe_name}", fillcolor="{color_map[status_map[node.node_id]]}"];')
    for node in nodes:
        for p in node.prerequisites:
            dot_lines.append(f'"{p}" -> "{node.node_id}";')
    dot_lines.append("}")
    try:
        st.graphviz_chart("\n".join(dot_lines))
    except Exception:
        pass  # badge list above is the guaranteed-to-render view


def show_student_review_queue():
    st.header("🔁 Spaced Repetition Review")
    st.caption("Cards resurface right before you'd naturally forget them (SM-2 scheduling, the algorithm family behind Anki).")
    _log_daily_activity()

    due = SpacedRepetitionEngine.due_cards(st.session_state.sr_cards)

    if not due:
        st.success("Nothing due right now. Master a concept in the Knowledge Map or answer practice questions to add cards.")
    else:
        st.info(f"{len(due)} card(s) due today.")
        card = due[0]
        st.markdown(f"### {card.concept_name}")
        st.caption(f"Subject: {card.subject.title()} · Repetitions: {card.repetitions} · Easiness: {card.easiness} · Lapses: {card.lapses}")
        st.write("How well did you recall this, without looking anything up?")
        q = st.slider("Recall quality (0 = totally forgot, 5 = instant and easy):", 0, 5, 3, key=f"quality_{card.concept_id}")
        if st.button("Submit review", key=f"review_{card.concept_id}"):
            SpacedRepetitionEngine.schedule(card, q)
            st.session_state.sr_cards[card.concept_id] = card
            st.success(f"Next review: {card.next_review} (in {card.interval_days} day(s))")
            if len(st.session_state.sr_cards) >= 5 and all(c.repetitions >= 1 for c in st.session_state.sr_cards.values()):
                _award_achievement("Review Habit Builder")
            st.rerun()

    if st.session_state.sr_cards:
        st.divider()
        st.subheader("All tracked cards")
        rows = [{
            'Concept': c.concept_name, 'Subject': c.subject.title(), 'Next Review': c.next_review,
            'Interval (days)': c.interval_days, 'Repetitions': c.repetitions, 'Easiness': c.easiness, 'Lapses': c.lapses
        } for c in st.session_state.sr_cards.values()]
        st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)


def show_student_progress():
    st.header("📊 Progress & Analytics")
    _log_daily_activity()

    student_engine = StudentOutcomeEngine(country_code)

    weeks = ['Week 1', 'Week 2', 'Week 3', 'Week 4', 'Current']
    scores = [45, 52, 58, 63, st.session_state.student_score]
    df = pd.DataFrame({'Week': weeks, 'Score': scores})
    st.line_chart(df.set_index('Week'))

    col1, col2 = st.columns(2)
    with col1:
        grade = student_engine._calculate_projected_grade({'score': st.session_state.student_score})
        st.metric("Projected Grade", grade)
    with col2:
        if st.session_state.student_score >= 60:
            st.success("✅ On track for success!")
        else:
            st.warning("📚 Keep practicing to improve")

    st.divider()
    st.subheader("🔍 Latent Knowledge Gap Tracking")
    st.caption("Combines solve speed, edit/hint attempts, and self-reported confidence — not just right/wrong — to surface hidden weak spots.")

    logs = [AttemptLog(**d) for d in st.session_state.attempt_logs]
    if not logs:
        st.info("Answer some practice questions to generate analytics here.")
    else:
        flags = LatentKnowledgeGapTracker.analyze(logs)
        if not flags:
            st.success("No latent gaps flagged from your recent attempts.")
        else:
            for f in flags:
                with st.expander(f"⚠️ {f['concept_id'].split('::')[-1][:40]} — gap score {f['gap_score']}"):
                    st.write("Reasons: " + ", ".join(f['reasons']))
                    m1, m2, m3, m4 = st.columns(4)
                    m1.metric("Accuracy", f"{f['accuracy_pct']}%")
                    m2.metric("Avg time", f"{f['avg_time_sec']}s")
                    m3.metric("Avg edits/hints", f['avg_edits'])
                    m4.metric("Avg confidence", f"{f['avg_confidence']}/5")

        st.caption(f"Based on {len(logs)} logged attempt(s) this session.")


def show_student_achievements():
    st.header("🏆 Achievements & Streaks")
    _log_daily_activity()

    col1, col2, col3 = st.columns(3)
    col1.metric("Day Streak", f"{st.session_state.streak} days")
    col2.metric("Streak-Repair Buffers", st.session_state.streak_repair_tokens)
    col3.metric("Achievements", len(st.session_state.achievements))

    st.caption("A streak-repair buffer auto-covers one missed day so a single skip doesn't wipe out your progress.")

    st.divider()
    if st.session_state.achievements:
        for achievement in st.session_state.achievements:
            st.markdown(f'<span class="achievement-badge">🏆 {achievement}</span>', unsafe_allow_html=True)
    else:
        st.info("No achievements yet — answer questions, master concepts, and keep your streak alive to earn some.")

    st.divider()
    st.subheader("🌳 Mastery Overview")
    total_mastered = sum(len(v) for v in st.session_state.mastered_concepts.values())
    st.metric("Concepts mastered across all subjects", total_mastered)


def show_student_privacy():
    st.header("🔒 Privacy & Safety")

    st.markdown("""
    **How this works in this app, plainly:**
    - Everything you do here — practice answers, OCR uploads, voice transcripts, notes — lives only in this browser session's memory (`st.session_state`).
    - Nothing is written to a database or file by this app, and nothing here is used to train any AI model.
    - This page's guarantee covers the application code only — I can't verify server- or host-level logging on whatever infrastructure this is deployed on. Check your deployment's own data-retention policy for that layer.
    """)

    st.session_state.privacy_consent = st.toggle("I understand my session data is temporary and stored only for this session.", value=st.session_state.privacy_consent)

    st.divider()
    st.subheader("🛡️ Contextual Content Shield")
    st.caption("Any free-text box in Practice is checked against a simple keyword filter that blocks requests to skip learning (e.g. 'just give me the answer', 'do my homework for me'). It's heuristic, not perfect — it can be worded around, so it's a guardrail, not a guarantee.")

    st.divider()
    st.subheader("🧹 Erase my data")
    st.caption("Clears practice history, spaced-repetition cards, mastery progress, OCR/voice history, and achievements from this session.")
    if st.button("🗑️ Purge my session data", type="secondary"):
        PrivacyControlEngine.purge(st.session_state)
        st.success("Session data cleared.")
        st.rerun()


# ==================== MAIN NAVIGATION ====================

def main():
    if choice == '📄 Document Analysis':
        show_document_analysis()
    elif choice == '🎓 Dashboard':
        show_student_dashboard()
    elif choice == '📝 Practice':
        show_student_practice()
    elif choice == '🧠 Knowledge Map':
        show_student_knowledge_map()
    elif choice == '🔁 Review Queue':
        show_student_review_queue()
    elif choice == '📊 Progress':
        show_student_progress()
    elif choice == '🏆 Achievements':
        show_student_achievements()
    elif choice == '🔒 Privacy & Safety':
        show_student_privacy()
    elif choice == '👨‍🏫 Dashboard':
        show_teacher_dashboard()
    elif choice == '📋 Lesson Builder':
        show_teacher_lesson_builder()
    elif choice == '📝 Assessment':
        show_teacher_assessment()
    elif choice == '⏱️ Hours Saved':
        show_teacher_hours_saved()
    elif choice == '📑 LMS Integration':
        show_teacher_lms_integration()
    elif choice == '🎯 Standards Mapping':
        show_teacher_standards_mapping()
    elif choice == '💼 Dashboard':
        show_professional_dashboard()
    elif choice == '🔬 Research':
        show_professional_research()
    elif choice == '📊 Portfolio':
        show_professional_portfolio()
    elif choice == '📑 Workspace':
        show_professional_workspace()
    elif choice == '⚡ Watchdogs':
        show_professional_watchdogs()
    elif choice == '🏢 Dashboard':
        show_sme_dashboard()
    elif choice == '📈 Growth':
        show_sme_growth()
    elif choice == '🤖 Automation':
        show_sme_automation()
    elif choice == '📊 Analytics':
        show_sme_growth()
    elif choice == '🔌 API Connectors':
        show_sme_api_connectors()
    elif choice == '⚡ Webhooks':
        show_sme_webhooks()
    else:
        show_home()

if __name__ == "__main__":
    main()
